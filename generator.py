import streamlit as st
import json
import csv
import os
import pandas as pd
import s3fs
import numpy as np
import xlsxwriter
import io
from PIL import Image
import plotly.express as px
from st_aggrid import GridOptionsBuilder, AgGrid, GridUpdateMode, DataReturnMode, ColumnsAutoSizeMode
from pyxlsb import open_workbook as open_xlsb
from streamlit_extras.customize_running import center_running
import time
#import python-docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
pd.set_option('display.precision', 1)
from io import BytesIO
from datetime import datetime
import pytz
import time



def app():

    with st.sidebar.expander("ABOUT THE APP"):
        st.markdown("""
                This app generates generic audit reports populating them with data collected in
                prontoforms(truecontext) forms.  
                Depending on the scope, it detects whether EHT circuits, panels, insulation or mix 
                of any of those were examined.  
                Based on that data, most of app fields are automatically populated and only require 
                users confirmation. Those that cannot be directly determined from the forms, need to be
                decited by the user.  
                Anyhow, thorough check-up is highly recommended!  
                When all fields are completed and confirmed, "Generate Report" button will trigger report creation.
                Afterwards, it can be downloaded with the selected name.
                """)

    fs = s3fs.S3FileSystem(anon=False)

    @st.cache_data(ttl=600)
    def list_files(search_item):
        return fs.find(search_item)

    @st.cache_data(ttl=600)
    def open_file(filename):
        return fs.open(filename)
    
    @st.cache_data(ttl=600)
    def read_file(filename):
        with fs.open(filename) as f:
            return f.read().decode("utf-8")

    def get_csvlog(file_name):
        file_content = open_file(file_name)
        data = pd.read_csv(file_content, encoding="cp1252")
        return data

    def save_csv(df,filename):
        with fs.open(filename, 'w') as file:
            return df.to_csv(file, index=False)

    if "load_state" not in st.session_state:
        st.session_state.load_state = False

    @st.cache_data(ttl=1800)
    def get_csvsource(file_name):
        file_content = open_file(file_name)
        data = pd.read_csv(file_content, encoding="cp1252")
        return data


    @st.cache_data(ttl=1800)
    def collect_data(forms):

        #form_df = None

        form_df = pd.DataFrame({'Project':[],'Form':[],'Date':[],'ID':[], 'Section':[], 'Item':[],
                    'Question_label':[],'Question_name':[],
                        'Question_subname':[],'Answer':[],'Status':[],'Comment':[], 'Data_type':[]})

        # for form_name in forms:
        #     searched_dir = f"{directory}{form_name}/"
        #     #list_of_files = fs.find(searched_dir)
        #     list_of_files = list_files(searched_dir)
        #     searched_files = []
        #     unique_ids = []

        #     for s3_file in list_of_files:
        #         if s3_file.endswith('.json'):
        #             searched_files.append(s3_file)

        #     st.write(searched_files)
        id_corr_dict = {}
        id_date_dict = {}
        for file in forms:
            unit_id = file[0:-24]
            if '/Audits/' in unit_id:
                split_1 = unit_id.split('Audits/')[1].split('/',1)
            else:
                split_1 = unit_id.split('Construction/')[1].split('/',1)
            uniq_proj = split_1[0]
            form_pref = split_1[1].split(' - ')[0]    
            idno = split_1[1].split('/',1)[1].split('_',1)[-1].split('_NF')[0]
            unique_id =  uniq_proj + '_' + form_pref + '_' + idno  
            pfdate = file[-24:-5]
            pfdate_format = datetime.strptime(pfdate,'%Y-%m-%d_%H_%M_%S')
            if unique_id not in id_corr_dict.keys():
                id_corr_dict[unique_id] = unit_id
                id_date_dict[unit_id] = pfdate
            else:
                #st.write(id_date_dict)
                #st.write(id_date_dict[unit_id])
                existing_form = id_date_dict[id_corr_dict[unique_id]]
                exist_date = datetime.strptime(existing_form,'%Y-%m-%d_%H_%M_%S')
                if pfdate_format > exist_date:
                    id_date_dict[unit_id] = pfdate
                    id_corr_dict[unique_id] = unit_id
        list_of_ujsons = []
        for unitid, udate in id_date_dict.items():
            list_of_ujsons.append(unitid+udate+'.json') 

        for jejson in list_of_ujsons:
            #st.write(jejson)
            file_content = read_file(jejson)
            data = json.loads(file_content)
            #eht_cct_no = data['pages'][0]['sections'][0]['answers'][4]['values'][0]
            if '/Audits/' in jejson:
                proj_split = jejson.split('Audits/')[1].split('/',1)
            else:
                proj_split = jejson.split('Construction/')[1].split('/',1)   
            project_no = proj_split[0]
            form_split = proj_split[1].split('/',1)
            form_no = form_split[0]
            id_no = form_split[1].split('_',1)[-1].split('_NF')[0]
            date_def = form_split[1].rsplit('.',1)[0].rsplit('_')[-4]
            (project_id,form_id,repdate,temp,weather,unit_id,section_name,question_label,question_name,question_subname,
            answer_value,branch_no,data_type,exception_type,comment) = ([] for i in range(15))

            fault_colors = ['#F6E2DF','#C0392B']
            for answer in data['pages'][0]['sections'][0]['answers']:
                if answer['label']=='WeatherGeneral':
                    try:
                        weather_cond = answer['values'][0]
                    except:
                        weather_cond = 'N/A'
                if answer['label']=='WeatherTemperature':
                    try:
                        weather_temp = answer['values'][0]
                    except:
                        weather_temp = 'N/A'
            for item in data['pages'][0]['sections']:
                #if item['type'] == 'Flow':
                if item['type'] != 'Repeat':
                    for answer in item['answers']:
                        section_name.append(item['label'])
                        question_label.append(answer['label'])
                        question_name.append(answer['question'])
                        question_subname.append('')
                        data_type.append(answer['dataType'])
                        branch_no.append(1)
                        unit_id.append(id_no)
                        project_id.append(project_no)
                        form_id.append(form_no)
                        weather.append(weather_cond)
                        temp.append(weather_temp)
                        try:
                            repdate.append(answer['lastModified']['provided']['time'][:10])
                        except:
                            repdate.append(date_def)
                        try:
                            answer_value.append(answer['values'][0])
                        except:
                            answer_value.append('')
                        try:
                            if answer['valuesMetadata'][0]['exception']['backgroundColor'] in fault_colors:
                                exception_type.append('Fault')
                            else:
                                exception_type.append('OK')  
                        except:
                            exception_type.append('')
                        try:
                            comment.append(answer['comment'])
                        except:
                            comment.append('')
                if item['type'] == 'Repeat':
                    branch_count = 1
                    for branch in item['rows']:
                        for elem in branch['pages'][0]['sections'][0]['answers']:
                            section_name.append(item['label'])
                            question_label.append(elem['label'])
                            question_name.append(item['name'])
                            try:
                                repdate.append(answer['lastModified']['provided']['time'][:10])
                            except:
                                repdate.append(date_def)
                            try:
                                question_subname.append(elem['question'])
                            except:
                                question_subname.append('')
                            data_type.append(elem['dataType'])
                            branch_no.append(branch_count)
                            unit_id.append(id_no)
                            project_id.append(project_no)
                            form_id.append(form_no)
                            weather.append(weather_cond)
                            temp.append(weather_temp)
                            try:
                                answer_value.append(elem['values'][0])
                            except:
                                answer_value.append('')
                            try:
                                if elem['valuesMetadata'][0]['exception']['backgroundColor'] in fault_colors:
                                    exception_type.append('Fault')
                                else:
                                    exception_type.append('OK')
                            except:
                                exception_type.append('')
                            try:
                                comment.append(elem['comment'])
                            except:
                                if elem['label']=='Comments':
                                    try:
                                        comment.append(elem['values'][0])
                                    except:
                                        comment.append('')
                                else:
                                    comment.append('')
                        branch_count += 1


            form_df = pd.concat([form_df,
                                pd.DataFrame({
                                    'Project':project_id,
                                    'Form':form_id,
                                    'Date':repdate,
                                    'Temp':temp,
                                    'Weather':weather,
                                    'ID':unit_id,
                                    'Section':section_name,
                                    'Item':branch_no,
                                    'Question_label':question_label,
                                    'Question_name':question_name,
                                    'Question_subname':question_subname,
                                    'Answer':answer_value,
                                    'Status':exception_type,
                                    'Comment':comment,
                                    'Data_type':data_type
                                            })
                                            ])

        form_df = form_df.astype({"Item": int})
        form_df = form_df[form_df['Data_type'] != 'Image']
        form_df.drop(columns='Data_type',inplace=True)
        
        return form_df

    #@st.cache_data()
    # def to_excel(df):
    #     output = io.BytesIO()
    #     writer = pd.ExcelWriter(output, engine='xlsxwriter')
    #     tab_name = form_select.split('-')[0]
    #     df.to_excel(writer, index=False, sheet_name=tab_name)
    #     workbook = writer.book
    #     worksheet = writer.sheets[tab_name]
    #     format1 = workbook.add_format({'num_format': '0.00'}) 
    #     worksheet.set_column('A:A', None, format1)  
    #     writer.close()
    #     processed_data = output.getvalue()
    #     return processed_data

       
    def to_excel(df,fname):
        output = io.BytesIO()
        # workbook = xlsxwriter.Workbook(output, {'in_memory': True})
        # worksheet = workbook.add_worksheet()
        # worksheet.write('A1', 'Hello')
        # workbook.close()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        #tab_name = form_select.split('-')[0]
        df.to_excel(writer, index=False, sheet_name=fname)
        workbook = writer.book
        worksheet = writer.sheets[fname]
        format1 = workbook.add_format({'num_format': '0.00'}) 
        worksheet.set_column('A:A', None, format1)  
        writer.close()
        processed_data = output.getvalue()
        return processed_data

    #@st.cache_data() 
    def convert_results(df):
        return df.to_csv().encode('utf-8')
    #if st.sidebar.button('Collect Data'):
    with st.form('input'):
        proj_dir = "s3-nvent-prontoforms-data/"
        form_select = ''
        with st.sidebar:
            #category_select = st.sidebar.selectbox('Select Category',('','Audits','Construction'),key='category_selection')
            cat_dir = "s3-nvent-prontoforms-data/Audits/"
            #projects_list = fs.find(cat_dir)
            projects_list = list_files(cat_dir)
            project_nos = []
            for pfile in projects_list:
                try:
                    project_nos.append(pfile.split('/')[2])
                except:
                    continue
            #if category_select:
            project_select = st.sidebar.selectbox('Select a Project',np.unique(project_nos).tolist(),key='project_selection')
            proj_dir = f"{cat_dir}{project_select}/"
            #st.write(proj_dir)
            #forms_list = fs.find(proj_dir)
            forms_list = list_files(proj_dir)
            forms_list_json = []
            for anyfile in forms_list:
                if anyfile.endswith('.json'):
                    forms_list_json.append(anyfile)
            if project_select:            
                generate_button = st.form_submit_button('Generate Report')

        form_nos = []
        audited_373 = 'n'
        audited_374 = 'n'
        audited_375 = 'n'
        audited_373_cnt = 0
        audited_374_cnt = 0
        audited_375_cnt = 0
        audit_lvl_cct = 'None'
        audit_lvl_insul = 'None'
        audit_lvl_panel = 'None'

        for ffile in forms_list_json:
            try:
                form_done = ffile.split('/')[3]
            except:
                continue
            else:
                if form_done.startswith('NF373 -'):
                    form_nos.append(form_done)
                    audited_373_cnt += 1
                    audited_373 = 'y'
                if form_done.startswith('NF374 -'):
                    form_nos.append(form_done)
                    audited_374_cnt += 1
                    audited_374 = 'y'
                if form_done.startswith('NF375 -'):
                    form_nos.append(form_done)
                    audited_375_cnt += 1
                    audited_375 = 'y'

        form_nos_unique = np.unique(form_nos)

        #st.write(forms_list)
        #st.write(forms_list_json)
        #st.write(form_nos)
        #st.write(form_nos_unique)
        # if project_select:
        #     collect_button = st.button('Collect Data')
            

        if project_select:

            #if collect_button:
            center_running()
            time.sleep(2)
        #st.session_state.load_state = True
        #with st.form('input'):
            #generate_button = st.form_submit_button('Generate Report')
            form_df = collect_data(forms_list_json)
            # st.table(form_df[(form_df.Form=="NF374 - EHT Insulation - Audit Inspection Form") &
            #                 (form_df.ID == '1001')
            # ])
            #metric_col1, metric_col2, metric_col3 = st.columns(3)
            #metric_col1.metric('# Projects',form_df['Project'].nunique())
            #metric_col2.metric('# Form Types',form_df['Form'].nunique())
            #metric_col3.metric('# Forms',form_df['ID'].nunique())
            proj_df = get_csvsource('s3-nvent-prontoforms-data/Data_sources/SAP_projects.csv')
            try:
                proj_name = proj_df[proj_df['Project Definition']==project_select]['Project Definition description'].tolist()[0]
            except:
                proj_name = '?'

            #fault_df = form_df[form_df['Status']=='Fault']
            #fault_df = fault_df.groupby(['Project','Form','ID','Question_name']).agg({'Status':'count'}).reset_index()
            #fault_df.rename(columns={'Status':'Quantity'},inplace=True)
            #fault_df_graph = fault_df.groupby(['Project','Form','Question_name']).agg({'Quantity':'sum'}).reset_index()
            #if audited_3743 == 'y':




            st.markdown("""
                    ### General Information
                    """)
            with st.expander("Project Entry Data"):
                col_01, col_02 = st.columns([4,1])
                with col_01:
                    prescreen_lst = ['Compliance and safety',
                                    'Risk mitigation',
                                    'Maintenance insights',
                                    'System retrofit',
                                    'System optimization'
                                    ]

                    prescreen_selection = st.multiselect(
                                                "Reason for Audit",
                                                options=prescreen_lst,default=['Compliance and safety',
                                    'Risk mitigation',
                                    'Maintenance insights',
                                    'System retrofit',
                                    'System optimization'
                                    ]
                                                )

                with col_02:
                    repdate = st.date_input('Report date',key='repdate')

                col_11, col_12, col_13, col_14 = st.columns([1,1,1,1])
                with col_11:
                    projname = st.text_input('Project Name:',value=proj_name,key='projname')
                with col_12:
                    projpono = st.text_input('Project Number:',value=project_select,key='pono')
                with col_13:
                    projsapno = st.text_input('PO Number:',key='sapno',value='1122334455')
                with col_14:
                    compname = st.text_input('Company Name:',key='compname',value='test_company')

                col_21, col_22, col_23, col_24, col_25 = st.columns([1,1,1,1,1])
                with col_21:
                    clientname = st.text_input('Client Name:',key='clientname',value='test_client')
                with col_22:
                    compcountry = st.text_input('Company Country:',key='compcount',value='test_country')
                with col_23:
                    compcity = st.text_input('Company City:',key='compcity',value='test_city')
                with col_24:
                    compstreet = st.text_input('Company Street:',key='compstreet',value='test_city')
                with col_25:
                    compzip = st.text_input('Company ZIP code:',key='compzip',value='test_zipcode')

            with st.expander("Project Controlling Documents"):
                col_docs_1, col_docs_2, col_docs_3, col_docs_4, col_docs_5 = st.columns([1,1,1,1,1])
                with col_docs_1:
                    pid_avail = st.selectbox(
                            "P&IDs available",
                            ("y","n"),
                            index = 0
                            )

                with col_docs_2:
                    iso_avail = st.selectbox(
                            "Isometrics available",
                            ("y","n"),
                            index = 1
                            )

                with col_docs_3:
                    pnl_avail = st.selectbox(
                            "Panel drawings available",
                            ("y","n"),
                            index = 1
                            )

                with col_docs_4:
                    lay_avail = st.selectbox(
                            "Layouts available",
                            ("y","n"),
                            index = 1
                            )

                with col_docs_5:
                    calc_avail = st.selectbox(
                            "Calculation lists available",
                            ("y","n"),
                            index = 1
                            )

                col_attach_1, col_attach_2 = st.columns([1,1])
                with col_attach_1:
                    plotplan_img_attach = st.file_uploader("Overall Site Layout", accept_multiple_files=False,type=['png', 'jpg'] )
                    if plotplan_img_attach is not None:
                        plotplan_img_rep = BytesIO(plotplan_img_attach.read())

            if audited_375 == 'y':


                st.markdown("""
                        ### :blue[EHT Panel Audit]
                        """)

                #Constructs dataframes for checks
                form_df_pnl = form_df[(form_df.Form.str.startswith('NF375'))&(form_df.ID!='')]
                fdf_pnl_vis_fault = form_df_pnl[form_df_pnl.Status=='Fault']         
                fdf_pnl_vis_rep = fdf_pnl_vis_fault[['Date','Temp','Weather','ID','Question_name','Comment']]
                pnl_cnt = form_df_pnl.ID.nunique()
                pnl_cnt_fault = fdf_pnl_vis_fault.ID.nunique()
                # Predefine Audit Scope/Level
                
                if len(form_df_pnl[form_df_pnl.Question_label=='AuditScope']['Answer']) == 0:
                    panel_lvl_index = None
                    audit_lvl_panel = 'None'                    
                else:
                    audit_pnl_lvls = form_df_pnl[form_df_pnl.Question_label=='AuditScope']['Answer'].tolist()
                    resfinal_pnl = 0
                    for lvl_pnl in audit_pnl_lvls:
                        if (lvl_pnl.find('3')!=-1):
                            reslvl_pnl = 2
                        elif (lvl_pnl.find('2')!=-1):
                            reslvl_pnl = 1
                        else:
                            reslvl_pnl = 0
                        resfinal_pnl = max(resfinal_pnl,reslvl_pnl)

                    if resfinal_pnl == 2:
                        audit_lvl_panel = 'Advanced'
                        panel_lvl_index = 2
                    elif resfinal_pnl == 1:
                        audit_lvl_panel = 'Standard'
                        panel_lvl_index = 1
                    else:
                        audit_lvl_panel = 'Basic'
                        panel_lvl_index = 0
                #EHT panels visual - Check for faults and calculate faults factor
                f_pnl_fault = (pnl_cnt_fault / pnl_cnt)

                if f_pnl_fault < 0.5:
                    vispanel_majority_index = 0
                else:
                    vispanel_majority_index = 1

                if len(fdf_pnl_vis_fault) != 0:
                    vispanelfail_gen_index = 0
                    if panel_lvl_index == 2:
                        vispanelfail_ifr_index = 0
                else:
                    vispanelfail_gen_index = 1
                    if panel_lvl_index == 2:
                        vispanelfail_ifr_index = 1

                #RCD testing - button or full measurements
                
                #RCD Button DF
                fdf_rcdbut = form_df_pnl[(form_df_pnl.Section=='RCD_Test_L1') 
                                            & ((form_df_pnl.Question_label=='DeviceTag_L1') | (form_df_pnl.Question_label=='Result_L1'))
                                            & (form_df_pnl.Answer!='')]
                #st.table(fdf_rcdbut)
                fdf_rcdbut_rep = fdf_rcdbut[['ID','Date','Temp','Weather','Question_label','Answer','Item','Status']]
                fdf_rcdbut_rep.reset_index(inplace=True,drop=True)
                panelid_l1 = []
                date_l1 = []
                temp_l1=[]
                weather_l1 = []
                devicetag_l1 = []
                result_l1 = []
                remarks_l1 = []
                cnt_butok = 0
                cnt_butfault = 0
                #st.table(fdf_rcdbut_rep)
                for row in fdf_rcdbut_rep.itertuples():
                    if row[5]=='DeviceTag_L1':
                        rcdbut_panelno = fdf_rcdbut_rep.iloc[row[0], 0]
                        rcdbut_itemno = fdf_rcdbut_rep.iloc[row[0], 6]
                        panelid_l1.append(rcdbut_panelno)
                        date_l1.append(fdf_rcdbut_rep.iloc[row[0], 1])
                        temp_l1.append(fdf_rcdbut_rep.iloc[row[0], 2])
                        weather_l1.append(fdf_rcdbut_rep.iloc[row[0], 3])                        
                        devicetag_l1.append(fdf_rcdbut_rep.iloc[row[0], 5])
                        fdf_rcdbut_result = fdf_rcdbut_rep[(fdf_rcdbut_rep.ID==rcdbut_panelno) 
                                            & (fdf_rcdbut_rep.Item==rcdbut_itemno)
                                            & (fdf_rcdbut_rep.Question_label=='Result_L1')]
                        if len(fdf_rcdbut_result) != 0:
                            result_l1.append(fdf_rcdbut_result['Answer'].tolist()[0])
                            if fdf_rcdbut_result.iloc[0, 7] == 'OK':
                                remarks_l1.append('-')
                                cnt_butok +=1
                            else:
                                remarks_l1.append('Immediate corrective actions required')
                                cnt_butfault +=1
                        else:
                            remarks_l1.append('N/A')
                            result_l1.append('N/A')
                            
                    
                fdf_rcdbut_rep = pd.DataFrame({
                                                'Panel_ID':panelid_l1,
                                                'RCD_ID':devicetag_l1,
                                                'Date':date_l1,
                                                'Temp':temp_l1,
                                                'Weather':weather_l1,
                                                'Status':result_l1,
                                                'Remarks':remarks_l1      
                                                })                
                #st.table(fdf_rcdbut_rep)
                cnt_butall = cnt_butok + cnt_butfault
                if cnt_butall == 0:
                    rcdbut_test_majority_index = 0
                    rcdbut_test_fail_index = 1
                else:
                    f_rcdbutfail = (cnt_butfault / cnt_butall)
                    if cnt_butfault !=0:
                        rcdbut_test_fail_index = 0
                    else:
                        rcdbut_test_fail_index = 1
                    if f_rcdbutfail > 0.5:
                        rcdbut_test_majority_index = 1
                    else:
                        rcdbut_test_majority_index = 0

                if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                #RCD Measurements DF                
                    fdf_rcdmes = form_df_pnl[(form_df_pnl.Section=='RCD_Test_L23') 
                                                & ((form_df_pnl.Question_label=='DeviceTag_L23') | 
                                                    (form_df_pnl.Question_label=='TripCurrent_L23') |
                                                    (form_df_pnl.Question_label=='MeasuredTripTime_L23') |
                                                    (form_df_pnl.Question_label=='TestButton_L23') |
                                                    (form_df_pnl.Question_label=='Result_L23')
                                                    )
                                                & (form_df_pnl.Answer!='')]
                    fdf_rcdmes_rep = fdf_rcdmes[['ID','Date','Temp','Weather','Question_label','Answer','Item','Status']]
                    fdf_rcdmes_rep.reset_index(inplace=True,drop=True)
                    #st.table(fdf_rcdmes_rep)
                    panelid_l23 = []
                    date_l23 = []
                    temp_l23=[]
                    weather_l23 = []
                    devicetag_l23 = []
                    tripcurrent_l23 = []
                    triptime_l23 = []
                    testbutton_l23 = []
                    result_l23 = []
                    remarks_l23 = []
                    cnt_mesok = 0
                    cnt_mesfault = 0

                    for row in fdf_rcdmes_rep.itertuples():
                        if row[5]=='DeviceTag_L23':
                            rcdmes_panelno = fdf_rcdmes_rep.iloc[row[0], 0]
                            rcdmes_itemno = fdf_rcdmes_rep.iloc[row[0], 6]
                            panelid_l23.append(rcdmes_panelno)
                            date_l23.append(fdf_rcdmes_rep.iloc[row[0], 1])
                            temp_l23.append(fdf_rcdmes_rep.iloc[row[0], 2])
                            weather_l23.append(fdf_rcdmes_rep.iloc[row[0], 3])                        
                            devicetag_l23.append(fdf_rcdmes_rep.iloc[row[0], 5])
                            fdf_rcdmes_curr = fdf_rcdmes_rep[(fdf_rcdmes_rep.ID==rcdmes_panelno)
                                                & (fdf_rcdmes_rep.Item==rcdmes_itemno)
                                                & (fdf_rcdmes_rep.Question_label=='TripCurrent_L23')]
                            fdf_rcdmes_time = fdf_rcdmes_rep[(fdf_rcdmes_rep.ID==rcdmes_panelno)
                                                & (fdf_rcdmes_rep.Item==rcdmes_itemno)
                                                & (fdf_rcdmes_rep.Question_label=='MeasuredTripTime_L23')]
                            fdf_rcdmes_but = fdf_rcdmes_rep[(fdf_rcdmes_rep.ID==rcdmes_panelno)
                                                & (fdf_rcdmes_rep.Item==rcdmes_itemno)
                                                & (fdf_rcdmes_rep.Question_label=='TestButton_L23')]
                            fdf_rcdmes_result = fdf_rcdmes_rep[(fdf_rcdmes_rep.ID==rcdmes_panelno)
                                                & (fdf_rcdmes_rep.Item==rcdmes_itemno)
                                                & (fdf_rcdmes_rep.Question_label=='Result_L23')]

                            if len(fdf_rcdmes_curr) != 0:
                                tripcurrent_l23.append(fdf_rcdmes_curr['Answer'].tolist()[0])
                            else:
                                tripcurrent_l23.append('N/A')

                            if len(fdf_rcdmes_time) != 0:
                                triptime_l23.append(fdf_rcdmes_time['Answer'].tolist()[0])
                            else:
                                triptime_l23.append('N/A')

                            if len(fdf_rcdmes_but) != 0:
                                testbutton_l23.append(fdf_rcdmes_but['Answer'].tolist()[0])
                            else:
                                testbutton_l23.append('N/A')

                            if len(fdf_rcdmes_result) != 0:
                                result_l23.append(fdf_rcdmes_result['Answer'].tolist()[0])
                                if fdf_rcdmes_result.iloc[0, 7] == 'OK':
                                    remarks_l23.append('-')
                                    cnt_mesok +=1
                                else:
                                    remarks_l23.append('Immediate corrective actions required')
                                    cnt_mesfault +=1
                            else:
                                remarks_l23.append('N/A')
                                result_l23.append('N/A')

                    fdf_rcdmes_rep = pd.DataFrame({
                                                    'Panel_ID':panelid_l23,
                                                    'RCD_ID':devicetag_l23,
                                                    'Date':date_l23,
                                                    'Temp':temp_l23,
                                                    'Weather':weather_l23,
                                                    'Button_test':testbutton_l23,
                                                    'Trip_time':triptime_l23,
                                                    'Trip_current':tripcurrent_l23,
                                                    'Status':result_l23,
                                                    'Remarks':remarks_l23       
                                                    })           
                    #st.table(fdf_rcdmes_rep)

                    cnt_mesall = cnt_mesok + cnt_mesfault
                    if cnt_mesall == 0:
                        rcdmes_test_majority_index = 0
                        rcdmes_test_fail_index = 1
                    else:
                        f_rcdmesfail = (cnt_mesfault / cnt_mesall)
                        if cnt_mesfault !=0:
                            rcdmes_test_fail_index = 0
                        else:
                            rcdmes_test_fail_index = 1
                        if f_rcdmesfail > 0.5:
                            rcdmes_test_majority_index = 1
                        else:
                            rcdmes_test_majority_index = 0

                with st.expander(":blue[General information]"):
                    #DATA ENTRY FIELDS
                    # st.write("""
                    #         :orange[General information]
                    #         """)
                    col_pnlinfo_1, col_pnlinfo_2, col_pnlinfo_3, col_pnlinfo_4 = st.columns([1,1,1,1])

                    with col_pnlinfo_1:
                        audit_lvl_panel = st.selectbox(
                            ":red[EHT Panel Audit Scope]",
                            ("Basic", "Standard", "Advanced"),
                            index=panel_lvl_index,
                            placeholder='"MANUAL INPUT REQ."'
                            )

                    with col_pnlinfo_2:
                        EHT_pnl_no = st.text_input('#Panels',value = str(pnl_cnt), key='pnlqty')

                with st.expander(":blue[EHT Panel Visual Inspection]"):
                    #DATA ENTRY FIELDS
                    # st.write("""
                    #         :orange[General information]
                    #         """)
                    col_pnlvis_1, col_pnlvis_2, col_pnlvis_3, col_pnlvis_4 = st.columns([1,1,1,1])

                    with col_pnlvis_1:
                        vispanel_majority = st.selectbox(
                            "EHT panels Majority-Status",
                            ("correct", "incorrect"),
                            index = vispanel_majority_index,
                            help="Majority - more than 50%",
                            key='vispanel_majority'
                        )
                    with col_pnlvis_2:
                        vispanelfail_gen = st.selectbox(
                            "General reported faults",
                            ("y", "n"),
                            index = vispanelfail_gen_index,
                            help="Are there any faults detected with standard inspection?",
                            key = 'vispanelfail_gen'
                    )
                    if panel_lvl_index == 2:
                        with col_pnlvis_3:
                            vispanelfail_ifr = st.selectbox(
                                ":red[Infrared reported faults]",
                                ("y", "n"),
                                index = vispanelfail_ifr_index,
                                help="Are there any faults detected with infrared camera?",
                                key = 'vispanelfail_ifr'
                        )                       
                #if audit_lvl_panel == 'Basic':
                with st.expander(":blue[RCD button test]"):
                    col_rcdbut_1, col_rcdbut_2, col_rcdbut_3, col_rcdbut_4 = st.columns([1,1,1,1])

                    with col_rcdbut_1:
                        rcdbut_test_majority =  st.selectbox(
                                "Majority of RCD button tests-Status",
                                ("correct", "incorrect"),
                                index = rcdbut_test_majority_index,
                                help = 'Majority is > 50%'
                        )
                    with col_rcdbut_2:
                        rcdbut_test_fail =  st.selectbox(
                                "RCD button test faults",
                                ("y", "n"),
                                index = rcdbut_test_fail_index,
                                help = 'Did any of RCDs fail button test?'
                    )                        
                if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                    with st.expander(":blue[RCD measurements]"):
                        col_rcdmes_1, col_rcdmes_2, col_rcdmes_3, col_rcdmes_4 = st.columns([1,1,1,1])

                        with col_rcdmes_1:
                            rcdmes_test_majority =  st.selectbox(
                                    "Majority of RCD extended tests-Status",
                                    ("correct", "incorrect"),
                                    index = rcdmes_test_majority_index,
                                    help = 'Majority is > 50%'
                            )
                        with col_rcdmes_2:
                            rcdmes_test_fail =  st.selectbox(
                                    "RCD extended test faults",
                                    ("y", "n"),
                                    index = rcdmes_test_fail_index,
                                    help = 'Did any of RCDs fail extended test?'
                        )    


            if audited_373 == 'y':

                st.markdown("""
                        ### :orange[EHT Circuits Audit]
                        """)

                #Constructs dataframes for checks
                form_df_cct = form_df[(form_df.Form.str.startswith('NF373'))&(form_df.ID!='')]
                fdf_cct_vis = form_df_cct[form_df_cct.Section!='FieldThermostatInspe']
                fdf_cct_vis_fault = fdf_cct_vis[fdf_cct_vis.Status=='Fault']
                #st.table(form_df_cct)
                fdf_cct_th = form_df_cct[form_df_cct.Section=='FieldThermostatInspe']
                fdf_cct_th_fault = fdf_cct_th[fdf_cct_th.Status=='Fault']
                fdf_th_setmain_fault = fdf_cct_th_fault[fdf_cct_th_fault.Question_label=="Ins_ThMSetpoint"]
                fdf_th_setlim_fault = fdf_cct_th_fault[fdf_cct_th_fault.Question_label=="Ins_ThLSetpoint"]
                fdf_th_pw_fault = fdf_cct_th_fault[fdf_cct_th_fault.Question_label=="Ins_ThFieldPWOut"]
                fdf_th_rtd_fault = fdf_cct_th_fault[fdf_cct_th_fault.Question_label=="Ins_ThFieldPt100"]
                    #Table for Cct visual inspection:
                fdf_cct_vis_rep = fdf_cct_vis_fault[['Date','Temp','Weather','ID','Question_name','Comment']]
                #st.table(fdf_cct_vis_rep)
                    #Table for Controller visual inspection
                fdf_cct_th_rep = fdf_cct_th_fault[['Date','Temp','Weather','ID','Question_name','Comment']]
                #st.table(fdf_cct_vis_rep)
                #st.table(fdf_cct_th_rep)
                cct_cnt = form_df_cct.ID.nunique()
                cct_cnt_fault = fdf_cct_vis_fault.ID.nunique()
                th_cnt = fdf_cct_th.ID.nunique()
                th_cnt_fault = fdf_cct_th_fault.ID.nunique()


                #EHT circuits visual - Check for faults and calculate faults factor
                f_cct_fault = (cct_cnt_fault / cct_cnt)

                if f_cct_fault < 0.5:
                    viscct_majority_index = 0
                else:
                    viscct_majority_index = 1

                if len(fdf_cct_vis_fault) != 0:
                    viscctfail_gen_index = 0
                else:
                    viscctfail_gen_index = 1

                #Controllers - Check for faults  and calculate faults factor
                f_th_fault = (th_cnt_fault / th_cnt)

                if f_th_fault < 0.5:
                    ctrl_majority_index = 0
                else:
                    ctrl_majority_index = 1                   

                if len(fdf_th_setmain_fault) != 0:
                    ctrlfail_mtemp_index = 0
                else:
                    ctrlfail_mtemp_index = 1

                if len(fdf_th_setlim_fault) != 0:
                    ctrlfail_ltemp_index = 0
                else:
                    ctrlfail_ltemp_index = 1                     
                
                if len(fdf_th_pw_fault) != 0:
                    ctrlfail_pw_index = 0
                else:
                    ctrlfail_pw_index = 1    

                if len(fdf_th_rtd_fault) != 0:
                    ctrlfail_sens_index = 0
                else:
                    ctrlfail_sens_index = 1                       


                #Controllers - Control method calculation
                ctrlmet_df = form_df_cct[form_df_cct.Question_label=='ThCtrlMethodDoc']
                ctrlmet_df_len = len(ctrlmet_df)
                ctrlmet_unctrl = 0
                ctrlmet_other = 0
                ctrl_loc_field = 0
                ctrl_loc_panel = 0
                for index,row in ctrlmet_df.iterrows():
                    if row['Answer'] == 'Ambient Sensing - Field Controller':
                        ctrl_loc_field += 1
                    if row['Answer'] == 'Ambient Sensing - Panel Controller':
                        ctrl_loc_panel += 1
                    if row['Answer'] == 'Line Sensing - Field Controller':
                        ctrl_loc_field += 1
                    if row['Answer'] == 'Line Sensing - Panel Controller with Field Sensor':
                        ctrl_loc_panel += 1
                    if row['Answer'] == 'Uncontrolled':
                        ctrlmet_unctrl +=1
                    if row['Answer'] == 'Other':
                        ctrlmet_other +=1
                ctrlmet_ctrl = ctrl_loc_field + ctrl_loc_panel

                if ctrlmet_unctrl > 0:
                    ctrl_not_index = 0
                else:
                    ctrl_not_index = 1

                if (ctrl_loc_field != 0 and ctrl_loc_panel == 0):
                    ctrl_mech_index = 0
                    ctrl_el_index = 0
                    ctrl_loc_index = 1
                if (ctrl_loc_field != 0 and ctrl_loc_panel != 0):
                    ctrl_mech_index = 0
                    ctrl_el_index = 0
                    ctrl_loc_index = 2
                elif (ctrl_loc_field == 0 and ctrl_loc_panel == 0):
                    ctrl_mech_index = 1
                    ctrl_el_index = 1
                    ctrl_loc_index = 3
                elif (ctrl_loc_field == 0 and ctrl_loc_panel != 0):
                    ctrl_mech_index = 1
                    ctrl_el_index = 0
                    ctrl_loc_index = 0

                
                # Predefine Audit Scope/Level
                if len(form_df_cct[form_df_cct.Question_label=='AuditScope']['Answer']) == 0:
                    cct_lvl_index = None
                    audit_lvl_cct = 'None'  
                else:
                    audit_cct_lvls = form_df_cct[form_df_cct.Question_label=='AuditScope']['Answer'].tolist()
                    resfinal_cct = 0
                    for lvl_cct in audit_cct_lvls:
                        if (lvl_cct.find('3')!=-1):
                            reslvl_cct = 2
                        elif (lvl_cct.find('2')!=-1):
                            reslvl_cct = 1
                        else:
                            reslvl_cct = 0
                        resfinal_cct = max(resfinal_cct,reslvl_cct)

                    if resfinal_cct == 2:
                        audit_lvl_cct = 'Advanced'
                        cct_lvl_index = 2
                    elif resfinal_cct == 1:
                        audit_lvl_cct = 'Standard'
                        cct_lvl_index = 1
                    else:
                        audit_lvl_cct = 'Basic'
                        cct_lvl_index = 0

                #IR test values
                fdf_irtest = form_df_cct[(
                            (form_df_cct.Question_label.str.startswith('El_1phMegger')) | 
                            (form_df_cct.Question_label.str.startswith('El_3phdMegger')) |
                            (form_df_cct.Question_label.str.startswith('El_3phsMegger'))
                            ) & (form_df_cct.Answer!='')]
                
                fdf_irtest_rep = fdf_irtest[['ID','Date','Temp','Weather','Question_label','Answer']]
                fdf_irtest_rep.rename(columns={'Answer':'Result'},inplace=True)
                fdf_irtest_rep['Result']=fdf_irtest_rep['Result'].astype(float)
                
                    #Table for check if MI
                fdf_ifMI = form_df_cct[form_df_cct.Question_label=='CheckIfMI']
                fdf_ifMI['Answer']=fdf_ifMI['Answer'].astype(int)
                fdf_ifMI = fdf_ifMI.groupby('ID').agg({'Answer':'sum'})
                fdf_ifMI.loc[fdf_ifMI.Answer >= 0, 'Answer'] = 1000
                fdf_ifMI.loc[fdf_ifMI.Answer < 0, 'Answer'] = 2500
                fdf_ifMI.rename(columns={'Answer':'Voltage'},inplace=True)

                fdf_irtest_rep = fdf_irtest_rep.merge(fdf_ifMI,on='ID',how='left')
                fdf_irtest_rep.reset_index(inplace=True)
                fdf_irtest_rep['L1-PE'] = np.nan
                #fdf_irtest_rep['L1-PE'] = fdf_irtest_rep['L1-PE'].astype(float)
                fdf_irtest_rep['L2-PE'] = np.nan
                #fdf_irtest_rep['L2-PE'] = fdf_irtest_rep['L2-PE'].astype(float)
                fdf_irtest_rep['L3-PE'] = np.nan
                #fdf_irtest_rep['L3-PE'] = fdf_irtest_rep['L3-PE'].astype(float)
                fdf_irtest_rep['Status'] = '-'
                fdf_irtest_rep['Remarks'] = '-'
                cnt_irabove100m = 0
                cnt_irbelow100m = 0
                cnt_irbelow10m = 0
                cnt_irbelow1kpv = 0

                for row in fdf_irtest_rep.itertuples():
                    if row[6]=='El_1phMegger':
                        fdf_irtest_rep.at[row[0], 'L1-PE'] = fdf_irtest_rep.iloc[row[0], 6]
                        fdf_irtest_rep.at[row[0], 'L2-PE'] = "N/A"
                        fdf_irtest_rep.at[row[0], 'L3-PE'] = "N/A"
                        if row[7]>=100:
                            cnt_irabove100m +=1
                            fdf_irtest_rep.at[row[0], 'Status'] = "Perfect"
                        elif 10<=row[7]<100:
                            cnt_irbelow100m +=1
                            fdf_irtest_rep.at[row[0], 'Status'] = "Good"
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Trend analysis advised"
                        elif 0.23<row[7]<10:
                            cnt_irbelow10m +=1
                            fdf_irtest_rep.at[row[0], 'Status'] = "Poor"
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Low IR to be investigated"
                        elif row[7]<=0.23:
                            cnt_irbelow1kpv +=1
                            fdf_irtest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Immediate corrective actions required"

                    if (row[6]=='El_3phdMeggerL1' or row[6]=='El_3phsMeggerL1'):
                        fdf_irtest_rep.at[row[0], 'L1-PE'] = fdf_irtest_rep.iloc[row[0], 6]
                        fdf_irtest_rep.at[row[0], 'L2-PE'] = fdf_irtest_rep.iloc[row[0]+1, 6]
                        fdf_irtest_rep.at[row[0], 'L3-PE'] = fdf_irtest_rep.iloc[row[0]+2, 6]
                        for row_phase in [0, 1, 2]:
                            if fdf_irtest_rep.iloc[row[0] + row_phase, 6]>=100:
                                cnt_irabove100m +=1
                            elif 10<=fdf_irtest_rep.iloc[row[0] + row_phase, 6]<100:
                                cnt_irbelow100m +=1
                            elif 0.4<fdf_irtest_rep.iloc[row[0] + row_phase, 6]<10:
                                cnt_irbelow10m +=1
                            elif fdf_irtest_rep.iloc[row[0] + row_phase, 6]<=0.4:
                                cnt_irbelow1kpv +=1

                        if (fdf_irtest_rep.iloc[row[0], 6]>=100 and fdf_irtest_rep.iloc[row[0]+1, 6]>=100 and fdf_irtest_rep.iloc[row[0]+2, 6]>=100):
                            fdf_irtest_rep.at[row[0], 'Status'] = "Perfect"
                        if (10<=fdf_irtest_rep.iloc[row[0], 6]<100 or 10<=fdf_irtest_rep.iloc[row[0]+1, 6]<100 or 10<=fdf_irtest_rep.iloc[row[0]+2, 6]<100):
                            fdf_irtest_rep.at[row[0], 'Status'] = "Good" 
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Trend analysis advised"
                        if (0.4<fdf_irtest_rep.iloc[row[0], 6]<10 or 0.4<fdf_irtest_rep.iloc[row[0]+1, 6]<10 or 0.4<fdf_irtest_rep.iloc[row[0]+2, 6]<10):
                            fdf_irtest_rep.at[row[0], 'Status'] = "Poor"
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Low IR to be investigated"
                        if (fdf_irtest_rep.iloc[row[0], 6]<=0.4 or fdf_irtest_rep.iloc[row[0]+1, 6]<=0.4 or fdf_irtest_rep.iloc[row[0]+2, 6]<=0.4):
                            fdf_irtest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_irtest_rep.at[row[0], 'Remarks'] = "Immediate corrective actions required"

                fdf_irtest_rep=fdf_irtest_rep.dropna(axis = 0, how = 'any')
                fdf_irtest_rep.drop(columns='Result',inplace=True)
                #fdf_irtest_rep['L1-PE'] = fdf_irtest_rep['L1-PE'].apply(lambda x: round(x, 2))
                #fdf_irtest_rep[['L1-PE','L2-PE','L3-PE']] = fdf_irtest_rep[['L1-PE','L2-PE','L3-PE']].round(1)
                #st.table(fdf_irtest_rep)

                cnt_irall = len(fdf_irtest)
                if cnt_irall == 0:
                    ir_test_majority_index = 0
                    ir_test_below10_index = 1
                    ir_test_fault_index = 1
                else:
                    f_irabove100m = (cnt_irabove100m / cnt_irall)
                    f_irbelow100m = (cnt_irbelow100m / cnt_irall)
                    f_irbelow10m = (cnt_irbelow10m / cnt_irall)
                    f_irbelow1kpv = (cnt_irbelow1kpv / cnt_irall)

                    if f_irabove100m > 0.5:
                        ir_test_majority_index = 0
                    elif (f_irbelow100m + f_irbelow100m) > 0.5:
                        ir_test_majority_index = 1
                    elif (f_irbelow100m + f_irbelow100m + f_irbelow10m) > 0.5:
                        ir_test_majority_index = 2
                    else:
                        ir_test_majority_index = 3                                                                 

                ir_test_below10_index = 1
                if cnt_irbelow10m != 0:
                    ir_test_below10_index = 0

                ir_test_fault_index = 1
                if cnt_irbelow1kpv != 0:
                    ir_test_fault_index = 0

                #Continuity test values
                fdf_conttest = form_df_cct[(
                            (form_df_cct.Question_label.str.startswith('El_1phResistance')) | 
                            (form_df_cct.Question_label.str.startswith('El_3phdResistance')) |
                            (form_df_cct.Question_label.str.startswith('El_3phsResistance'))
                            ) & (form_df_cct.Answer!='')]
                
                fdf_conttest_rep = fdf_conttest[['ID','Date','Temp','Weather','Question_label','Answer']]
                fdf_conttest_rep.rename(columns={'Answer':'Result'},inplace=True)
                fdf_conttest_rep['Result']=fdf_conttest_rep['Result'].astype(float)
                fdf_conttest_rep.reset_index(inplace=True,drop=True)
                fdf_conttest_rep['L-N'] = np.nan
                fdf_conttest_rep['L1-L2'] = np.nan
                fdf_conttest_rep['L1-L3'] = np.nan
                fdf_conttest_rep['L2-L3'] = np.nan
                fdf_conttest_rep['Status'] = '-'
                fdf_conttest_rep['Remarks'] = '-'
                cnt_contshort = 0
                cnt_contbreak = 0


                for row in fdf_conttest_rep.itertuples():
                    if row[5]=='El_1phResistance':
                        fdf_conttest_rep.at[row[0], 'L-N'] = fdf_conttest_rep.iloc[row[0], 5]
                        fdf_conttest_rep.at[row[0], 'L1-L2'] = "N/A"
                        fdf_conttest_rep.at[row[0], 'L1-L3'] = "N/A"
                        fdf_conttest_rep.at[row[0], 'L2-L3'] = "N/A"
                        if row[6]>=100000:
                            cnt_contbreak +=1
                            fdf_conttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Disturbed current path. Action required."
                        elif row[6]==0:
                            cnt_contshort +=1
                            fdf_conttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Short-circuit. Action required."
                        else:
                            fdf_conttest_rep.at[row[0], 'Status'] = "Tentatively acceptable"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Trend analysis required"

                    cnt_contbreak_cct = 0
                    cnt_contshort_cct = 0                       
                    if (row[5]=='El_3phdResistance12' or row[5]=='El_3phsResistance12'):
                        fdf_conttest_rep.at[row[0], 'L-N'] = "N/A"                            
                        fdf_conttest_rep.at[row[0], 'L1-L2'] = fdf_conttest_rep.iloc[row[0], 5]
                        fdf_conttest_rep.at[row[0], 'L1-L3'] = fdf_conttest_rep.iloc[row[0]+1, 5]
                        fdf_conttest_rep.at[row[0], 'L2-L3'] = fdf_conttest_rep.iloc[row[0]+2, 5]
                        for row_phase in [0, 1, 2]:
                            if fdf_conttest_rep.iloc[row[0] + row_phase, 5]>=100000:
                                cnt_contbreak +=1
                                cnt_contbreak_cct +=1
                            elif fdf_conttest_rep.iloc[row[0] + row_phase, 5]==0:
                                cnt_contshort +=1
                                cnt_contshort_cct +=1

                        if (cnt_contbreak_cct == 0 and cnt_contshort_cct == 0):
                            fdf_conttest_rep.at[row[0], 'Status'] = "Tentatively acceptable"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Trend analysis required"
                        elif (cnt_contbreak_cct != 0 and cnt_contshort_cct == 0):
                            fdf_conttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Disturbed current path. Action required."                         
                        elif (cnt_contbreak_cct == 0 and cnt_contshort_cct != 0):
                            fdf_conttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Short-circuit. Action required." 
                        elif (cnt_contbreak_cct != 0 and cnt_contshort_cct != 0):
                            fdf_conttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_conttest_rep.at[row[0], 'Remarks'] = "Disturbed current path & Short-circuit. Actions required!" 

                fdf_conttest_rep=fdf_conttest_rep.dropna(axis = 0, how = 'any')
                fdf_conttest_rep.drop(columns='Result',inplace=True)
                #st.table(fdf_conttest_rep)

                cnt_contall = len(fdf_conttest)
                if cnt_contall == 0:
                    cont_major_index = 0
                    cont_short_index = 1
                    cont_broke_index = 1
                else:
                    f_contbreak = (cnt_contbreak / cnt_contall)
                    f_contshort = (cnt_contshort / cnt_contall)

                    if (f_contbreak + f_contshort) <= 0.5:
                        cont_major_index = 0
                    else:
                        cont_major_index = 1                                                              

                cont_broke_index = 1
                if cnt_contbreak != 0:
                    cont_broke_index = 0

                cont_short_index = 1
                if cnt_contshort != 0:
                    cont_short_index = 0

                #Voltage test values
                fdf_volttest = form_df_cct[(
                            (form_df_cct.Question_label.str.startswith('El_1phVoltage')) | 
                            (form_df_cct.Question_label.str.startswith('El_3phdVoltage')) |
                            (form_df_cct.Question_label.str.startswith('El_3phsVoltage'))
                            ) & (form_df_cct.Answer!='')]

                fdf_volttest_rep = fdf_volttest[['ID','Date','Temp','Weather','Question_label','Answer']]
                fdf_volttest_rep.rename(columns={'Answer':'Result'},inplace=True)
                fdf_volttest_rep['Result']=fdf_volttest_rep['Result'].astype(float)
                fdf_volttest_rep.reset_index(inplace=True,drop=True)
                fdf_volttest_rep['L-N'] = np.nan
                fdf_volttest_rep['L1-L2'] = np.nan
                fdf_volttest_rep['L1-L3'] = np.nan
                fdf_volttest_rep['L2-L3'] = np.nan
                fdf_volttest_rep['Status'] = '-'
                fdf_volttest_rep['Remarks'] = '-'
                cnt_volt_ok = 0
                cnt_volt_out = 0
                cnt_volt_zero = 0

                for row in fdf_volttest_rep.itertuples():
                    if row[5]=='El_1phVoltage':
                        fdf_volttest_rep.at[row[0], 'L-N'] = fdf_volttest_rep.iloc[row[0], 5]
                        fdf_volttest_rep.at[row[0], 'L1-L2'] = "N/A"
                        fdf_volttest_rep.at[row[0], 'L1-L3'] = "N/A"
                        fdf_volttest_rep.at[row[0], 'L2-L3'] = "N/A"
                        if 207<=row[6]<=253:
                            cnt_volt_ok +=1
                            fdf_volttest_rep.at[row[0], 'Status'] = "Correct"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Value in acceptable range"
                        elif row[6]==0:
                            cnt_volt_zero +=1
                            fdf_volttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "No power supplied. Action required."
                        else:
                            cnt_volt_out +=1
                            fdf_volttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Voltage out of tolerance level"

                    cnt_volt_ok_cct = 0
                    cnt_volt_out_cct = 0
                    cnt_volt_zero_cct = 0
                    if (row[5]=='El_3phdVoltage12' or row[5]=='El_3psdVoltage12'):
                        fdf_volttest_rep.at[row[0], 'L-N'] = "N/A"                            
                        fdf_volttest_rep.at[row[0], 'L1-L2'] = fdf_volttest_rep.iloc[row[0], 5]
                        fdf_volttest_rep.at[row[0], 'L1-L3'] = fdf_volttest_rep.iloc[row[0]+1, 5]
                        fdf_volttest_rep.at[row[0], 'L2-L3'] = fdf_volttest_rep.iloc[row[0]+2, 5]
                        for row_phase in [0, 1, 2]:
                            if 207<=fdf_volttest_rep.iloc[row[0] + row_phase, 5]<=253:
                                cnt_volt_ok +=1
                                cnt_volt_ok_cct +=1
                            elif fdf_volttest_rep.iloc[row[0] + row_phase, 5]==0:
                                cnt_volt_zero +=1
                                cnt_volt_zero_cct +=1
                        cnt_volt_out_cct = cnt_volt_ok_cct - cnt_volt_zero_cct
                        cnt_volt_out += cnt_volt_out_cct

                        if (cnt_volt_out_cct == 0 and cnt_volt_zero_cct == 0):
                            fdf_volttest_rep.at[row[0], 'Status'] = "Correct"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Values in acceptable range"
                        elif (cnt_volt_out_cct != 0 and cnt_volt_zero_cct == 0):
                            fdf_volttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Voltage out of tolerance level. Action required."                         
                        elif (cnt_volt_out_cct == 0 and cnt_volt_zero_cct != 0):
                            fdf_volttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Power supply issue. Action required." 
                        elif (cnt_volt_out_cct != 0 and cnt_volt_zero_cct != 0):
                            fdf_volttest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_volttest_rep.at[row[0], 'Remarks'] = "Power supply issues. Actions required!" 

                fdf_volttest_rep=fdf_volttest_rep.dropna(axis = 0, how = 'any')
                fdf_volttest_rep.drop(columns='Result',inplace=True)

                cnt_voltall = len(fdf_volttest)
                if cnt_voltall == 0:
                    volt_major_index = 0
                    volt_out_index = 1
                    volt_zero_index = 1
                else:
                    f_voltout = (cnt_volt_out / cnt_voltall)
                    f_voltzero = (cnt_volt_zero / cnt_voltall)
                    f_voltok = (cnt_volt_ok / cnt_voltall)

                    if f_voltok > 0.5:
                        volt_major_index = 0
                    else:
                        volt_major_index = 1                                                              

                volt_out_index = 1
                if cnt_volt_out != 0:
                    volt_out_index = 0

                volt_zero_index = 1
                if cnt_volt_zero != 0:
                    volt_zero_index = 0

                #Current test values
                fdf_currtest = form_df_cct[(
                            (form_df_cct.Question_label.str.startswith('El_1phCurrent')) | 
                            (form_df_cct.Question_label.str.startswith('El_3phdCurrent')) |
                            (form_df_cct.Question_label.str.startswith('El_3phsCurrent'))
                            ) & (form_df_cct.Answer!='')]

                fdf_currtest_rep = fdf_currtest[['ID','Date','Temp','Weather','Question_label','Answer']]
                fdf_currtest_rep.rename(columns={'Answer':'Result'},inplace=True)
                fdf_currtest_rep['Result']=fdf_currtest_rep['Result'].astype(float)
                fdf_currtest_rep.reset_index(inplace=True,drop=True)
                fdf_currtest_rep['L1'] = np.nan
                fdf_currtest_rep['L2'] = np.nan
                fdf_currtest_rep['L3'] = np.nan
                fdf_currtest_rep['Status'] = '-'
                fdf_currtest_rep['Remarks'] = '-'
                cnt_currzero = 0
                #cnt_currtrip = 0

                for row in fdf_currtest_rep.itertuples():
                    if row[5]=='El_1phCurrent 1':
                        fdf_currtest_rep.at[row[0], 'L1'] = fdf_currtest_rep.iloc[row[0], 5]
                        fdf_currtest_rep.at[row[0], 'L2'] = "N/A"
                        fdf_currtest_rep.at[row[0], 'L3'] = "N/A"
                        if row[6]==0:
                            cnt_currzero +=1
                            fdf_currtest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_currtest_rep.at[row[0], 'Remarks'] = "No power drawn. Inspection required."
                        else:
                            fdf_currtest_rep.at[row[0], 'Status'] = "Tentatively acceptable"
                            fdf_currtest_rep.at[row[0], 'Remarks'] = "Trend analysis required"

                    #cnt_currok_cct = 0
                    #cnt_currtrip_cct = 0
                    cnt_currzero_cct = 0

                    if (row[5]=='El_3phdCurrent1' or row[5]=='El_3phsCurrent1'):                        
                        fdf_currtest_rep.at[row[0], 'L1'] = fdf_currtest_rep.iloc[row[0], 5]
                        fdf_currtest_rep.at[row[0], 'L2'] = fdf_currtest_rep.iloc[row[0]+1, 5]
                        fdf_currtest_rep.at[row[0], 'L3'] = fdf_currtest_rep.iloc[row[0]+2, 5]
                        for row_phase in [0, 1, 2]:
                            if fdf_currtest_rep.iloc[row[0] + row_phase, 5]==0:
                                cnt_currzero +=1
                                cnt_currzero_cct +=1

                        if cnt_currzero_cct == 0:
                            fdf_currtest_rep.at[row[0], 'Status'] = "Tentatively acceptable"
                            fdf_currtest_rep.at[row[0], 'Remarks'] = "Trend analysis required"
                        else:
                            fdf_currtest_rep.at[row[0], 'Status'] = "Fault"
                            fdf_currtest_rep.at[row[0], 'Remarks'] = "No power drawn. Inspection required."                             

                fdf_currtest_rep=fdf_currtest_rep.dropna(axis = 0, how = 'any')
                fdf_currtest_rep.drop(columns='Result',inplace=True)
                #st.table(fdf_currtest_rep)
                cnt_currall = len(fdf_currtest)
                if cnt_currall == 0:
                    curr_major_index = 0
                    curr_trip_index = 1
                    curr_zero_index = 1
                else:
                    f_currzero = (cnt_currzero / cnt_currall)

                    if f_currzero <= 0.5:
                        curr_major_index = 0
                    else:
                        curr_major_index = 1                                                              

                curr_zero_index = 1
                if cnt_currzero != 0:
                    curr_zero_index = 0

                curr_trip_index=None


                #DATA ENTRY FIELDS
                with st.expander(":orange[General Information]"):
                # st.write("""
                #         :orange[General information]
                #         """)
                    col_cctinfo_1, col_cctinfo_2, col_cctinfo_3, col_cctinfo_4 = st.columns([1,1,1,1])

                    with col_cctinfo_1:
                        audit_lvl_cct = st.selectbox(
                            ":orange[EHT Circuit Audit Scope]",
                            ("Basic", "Standard", "Advanced"),
                            index=cct_lvl_index,
                            placeholder='"MANUAL INPUT REQ."'
                            )

                    with col_cctinfo_2:
                        EHT_cct_no = st.text_input('#Circuits',value = str(cct_cnt), key='cctqty')

                #st.write(f"Controlled:{ctrlmet_ctrl}, Uncontrolled:{ctrlmet_unctrl}, Other:{ctrlmet_other}")
                # st.write("""
                #         :orange[Control&monitoring info]
                #         """)
                with st.expander(":orange[Control&monitoring info]"):
                    col_ctrlinfo_1, col_ctrlinfo_2, col_ctrlinfo_3, col_ctrlinfo_4 = st.columns([1,1,1,1])                   
                    with col_ctrlinfo_1:
                        ctrl_el = st.selectbox(
                            ":red[Electronic controllers]",
                            ("y","n"),
                            index=ctrl_el_index,
                            help="Preselected: positive if ANY control method present"
                        )
                    with col_ctrlinfo_2:
                        ctrl_mech = st.selectbox(
                            ":red[Mechanical controllers]",
                            ("y","n"),
                            index=ctrl_mech_index,
                            help="Preselected: positive if FIELD control method present"
                        )
                    with col_ctrlinfo_3:
                        ctrl_not = st.selectbox(
                            "Uncontrolled circuits",
                            ("y","n"),
                            index=ctrl_not_index,
                            help="Preselected: positive if any uncontrolled ccts present"
                        )
                    with col_ctrlinfo_4:
                        ctrl_loc = st.selectbox(
                            "Thermostats location",
                            ("panel","field","both",'N/A'),
                            index=ctrl_loc_index,
                            help="Preselected: based on form panel/field information"
                        )
                with st.expander(":orange[EHT Circuit Visual Inspection]"):
                # st.write("""
                #         :orange[EHT Circuit Visual Inspection]
                #         """)
                    col_cctstat_1, col_cctstat_2, col_cctstat_3, col_cctstat_4 = st.columns([1,1,1,1])
                    with col_cctstat_1:
                        viscct_majority = st.selectbox(
                            "EHT circuits Majority-Status",
                            ("correct", "incorrect"),
                            index = viscct_majority_index,
                            help="Majority - more than 50%"
                        )
                    with col_cctstat_2:
                        viscctfail_gen = st.selectbox(
                            "Reported faults",
                            ("y", "n"),
                            index = viscctfail_gen_index,
                            help="Are there any faults reported?"
                    )
                with st.expander(":orange[Controllers Inspection]"):
                # st.write("""
                #         :orange[Controllers Inspection]
                #         """)
                    col_ctrlstat_1, col_ctrlstat_2, col_ctrlstat_3, col_ctrlstat_4, col_ctrlstat_5 = st.columns([1,1,1,1,1])                   
                    with col_ctrlstat_1:
                        ctrl_majority =  st.selectbox(
                                "Majority of controllers-Status",
                                ("correct", "incorrect"),
                                index = ctrl_majority_index
                        )
                    with col_ctrlstat_2:
                        ctrlfail_mtemp =  st.selectbox(
                                "Tm setpoint errors",
                                ("y", "n"),
                                index = ctrlfail_mtemp_index
                        )
                    with col_ctrlstat_3:
                        ctrlfail_ltemp =  st.selectbox(
                                "Limiter setpoint errors",
                                ("y", "n"),
                                index = ctrlfail_ltemp_index
                        )
                    with col_ctrlstat_4:
                        ctrlfail_pw =  st.selectbox(
                                "Power switching faults",
                                ("y", "n"),
                                index = ctrlfail_pw_index
                        )
                    with col_ctrlstat_5:
                        ctrlfail_sens =  st.selectbox(
                                "Sensor faults",
                                ("y", "n"),
                                index = ctrlfail_sens_index
                        )

                with st.expander(":orange[Measurements: Insulation Resistance]"):
                # st.write("""
                #         :orange[Measurements: Insulation Resistance]
                #         """)
                    col_irtest_1, col_irtest_2, col_irtest_3, col_irtest_4 = st.columns([1,1,1,1])
                    with col_irtest_1:
                        ir_test_majority =  st.selectbox(
                                "Majority of IR tests - Status",
                                ("perfect", "good", "poor", "fault"),
                                index = ir_test_majority_index,
                                help = "Perfect:≥100MΩ, Good:≥10MΩ & ≤100MΩ, Poor:≤10MΩ, Fault:≤1kΩ per Volt of operating voltage (230V=230kΩ)"
                                )
                    with col_irtest_2:
                        ir_test_below10 =  st.selectbox(
                                "IR tests <10MΩ",
                                ("y", "n"),
                                index = ir_test_below10_index,
                                help = "Are there any results below 10MΩ?"
                                )
                    with col_irtest_3:
                        ir_test_fault =  st.selectbox(
                                "IR faulty tests",
                                ("y", "n"),
                                index = ir_test_fault_index,
                                help = "Did any test record faulty value?"
                                )

                with st.expander(":orange[Measurements: Continuity]"):
                # st.write("""
                #         :orange[Measurements: Continuity]
                #         """)
                    col_conttest_1, col_conttest_2, col_conttest_3, col_conttest_4 = st.columns([1,1,1,1])
                    with col_conttest_1:
                        cont_major =  st.selectbox(
                                ":red[Majority of Continuity tests - Status]",
                                ("good", "poor"),
                                index = cont_major_index,
                                help = "Based only on counted faults. Please verify!"
                                )
                    with col_conttest_2:
                        cont_short =  st.selectbox(
                                "Short-circuit results",
                                ("y", "n"),
                                index = cont_short_index,
                                help = "Are there any short-circuit measurements?"
                                )
                    with col_conttest_3:
                        cont_broke =  st.selectbox(
                                "Distrubed path results",
                                ("y", "n"),
                                index = cont_broke_index,
                                help = "Did any test record unusually high resistance (>100kΩ)?"
                                )

                with st.expander(":orange[Measurements: Voltage]"):
                # st.write("""
                #         :orange[Measurements: Voltage]
                #         """)
                    col_volttest_1, col_volttest_2, col_volttest_3, col_volttest_4 = st.columns([1,1,1,1])
                    with col_volttest_1:
                        volt_major =  st.selectbox(
                                "Majority of Voltage tests - Status",
                                ("in-range", "out-of-range"),
                                index = cont_major_index
                                )
                    with col_volttest_2:
                        volt_out =  st.selectbox(
                                "Out-of-range results",
                                ("y", "n"),
                                index = volt_out_index,
                                help = "Are there any out-of-range measurements?"
                                )
                    with col_volttest_3:
                        volt_zero =  st.selectbox(
                                "No voltage results",
                                ("y", "n"),
                                index = volt_zero_index,
                                help = "Did any test resulted in no voltage?"
                                )

                with st.expander(":orange[Measurements: Current]"):
                # st.write("""
                #         :orange[Measurements: Current]
                #         """)
                    col_currtest_1, col_currtest_2, col_currtest_3, col_currtest_4 = st.columns([1,1,1,1])
                    with col_currtest_1:
                        curr_major =  st.selectbox(
                                ":red[Majority of Current checks - Status]",
                                ("in-range", "out-of-range"),
                                index = curr_major_index,
                                help = "Based only on counted faults. Please verify!"
                                )
                    with col_currtest_2:
                        curr_trip =  st.selectbox(
                                ":red[CB trips]",
                                ("y", "n"),
                                index = curr_trip_index,
                                help = "Did any CBs trip during the test?",
                                placeholder="MANUAL INPUT REQ."
                                )
                    with col_currtest_3:
                        curr_zero =  st.selectbox(
                                "Zero amps readings",
                                ("y", "n"),
                                index = curr_zero_index,
                                help = "Did any test record zero amps?"
                                )

                if (ctrlfail_mtemp == 'y' or ctrlfail_ltemp == 'y'):
                    ctrlfail_temp = 'y'
                else:
                    ctrlfail_temp = 'n'

            if audited_374 == 'y':

                st.markdown("""
                        ### :green[EHT Insulation Audit]
                        """)

                #Constructs dataframes for checks
                map_insulation = {'Code 1 Count':'Missing/Damaged Sealant',
                'Code 2 Count':'Damaged Insulation/Cladding',
                'Code 3 Count':'Missing Cladding',
                'Code 4 Count':'Missing Insulation and Cladding and/or Blankets',
                'Code 5 Count':'Other'
                }
                form_df_ins = form_df[(form_df.Form.str.startswith('NF374'))&(form_df.ID!='')]
                form_df_ins.reset_index(inplace=True,drop=True)
                for row in form_df_ins.itertuples():
                    if (form_df_ins.iloc[row[0],6]=='Comments' and form_df_ins.iloc[row[0]+1,6]=='C1_Count'):
                        for cnt_row in range(1,6):
                            form_df_ins.at[row[0]+cnt_row,'Comment'] = form_df_ins.iloc[row[0],11]
                #st.table(form_df_ins)
                fdf_ins_vis_fault = form_df_ins[(form_df_ins.Status =='Fault') &
                (form_df_ins.Section !='InsulationParameters')]
                fdf_ins_vis_fault.reset_index(inplace=True,drop=True)
                #st.table(fdf_ins_vis_fault)
                for row in fdf_ins_vis_fault.itertuples():
                    if row[7].endswith('Count'):
                          fdf_ins_vis_fault.at[row[0],'Question_name'] = fdf_ins_vis_fault.iloc[row[0],8]
                fdf_ins_vis_fault['Section'] = np.arange(1,len(fdf_ins_vis_fault)+1)
                fdf_ins_vis_rep = fdf_ins_vis_fault[['ID','Section','Date','Temp','Weather','Question_name','Comment']]               
                fdf_ins_vis_rep['Question_name'] = fdf_ins_vis_rep['Question_name'].replace(map_insulation)
                #st.table(fdf_ins_vis_rep)
                ins_cnt = form_df_ins.ID.nunique()
                ins_cnt_fault = fdf_ins_vis_fault.ID.nunique()

                # Predefine Audit Scope/Level

                if len(form_df_ins[form_df_ins.Question_label=='AuditScope']['Answer']) == 0:
                    insul_lvl_index = None
                    audit_lvl_insul = 'None'                    
                else:
                    audit_ins_lvls = form_df_ins[form_df_ins.Question_label=='AuditScope']['Answer'].tolist()
                    resfinal_ins = 0
                    for lvl_ins in audit_ins_lvls:
                        if (lvl_ins.find('3')!=-1):
                            reslvl_ins = 2
                        elif (lvl_ins.find('2')!=-1):
                            reslvl_ins = 1
                        else:
                            reslvl_ins = 0
                        resfinal_ins = max(resfinal_ins,reslvl_ins)

                    if resfinal_ins == 2:
                        audit_lvl_insul = 'Advanced'
                        insul_lvl_index = 2
                    elif audit_lvl_insul == 1:
                        audit_lvl_insul = 'Standard'
                        insul_lvl_index = 1
                    else:
                        audit_lvl_insul = 'Basic'
                        insul_lvl_index = 0                

                #EHT insulation visual - Check for faults and calculate faults factor
                f_ins_fault = (ins_cnt_fault / ins_cnt)

                if f_ins_fault < 0.5:
                    visinsul_majority_index = 0
                else:
                    visinsul_majority_index = 1

                if len(fdf_ins_vis_fault) != 0:
                    visinsulfail_gen_index = 0
                    if (insul_lvl_index == 1 or insul_lvl_index == 2):
                        visinsulfail_ifr_index = 0
                else:
                    visinsulfail_gen_index = 1
                    if (insul_lvl_index == 1 or insul_lvl_index == 2):
                        visinsulfail_ifr_index = 1

                #Insulation Parameters DF
                if insul_lvl_index == 2:
                    fdf_insparam = form_df_ins[form_df_ins.Section=='InsulationParameters']
                    fdf_insparam_pvt = fdf_insparam.pivot(index=['ID','Item','Date','Temp','Weather'], columns='Question_label', values='Answer')
                    fdf_insparam_pvt.reset_index(inplace=True)
                    fdf_insparam_pvt['Section'] = np.arange(1,len(fdf_insparam_pvt)+1)
                    fdf_insparam_rep = fdf_insparam_pvt[['ID','Section','Date','Temp','Weather',
                                                            'PipeDiameter','InsulationType',
                                                            'InsulationThickness','UoM','AlignmentDoc']]
                    cnt_inparam_ok = len(fdf_insparam_rep[fdf_insparam_rep.AlignmentDoc=='Correct'])
                    cnt_insparam_fault = len(fdf_insparam_rep[fdf_insparam_rep.AlignmentDoc=='Incorrect'])
                    cnt_insparam_all = cnt_inparam_ok + cnt_insparam_fault
                    
                    if cnt_insparam_all == 0:
                        inscheck_majority_index = 0
                        incheck_some_index = 1
                        f_insparam_fault = 0
                    else:
                        f_insparam_fault = (cnt_insparam_fault / cnt_insparam_all)
                        if cnt_insparam_fault !=0:
                            incheck_some_index = 0
                        else:
                            incheck_some_index = 1
                        if f_insparam_fault > 0.5:
                            inscheck_majority_index = 1
                        else:
                            inscheck_majority_index = 0

                    #st.table(fdf_insparam_rep)
                #FORM DISPLAY
                with st.expander(":green[General information]"):
                    col_insinfo_1, col_insinfo_2, col_insinfo_3, col_insinfo_4 = st.columns([1,1,1,1])

                    with col_insinfo_1:
                        audit_lvl_insul = st.selectbox(
                            ":red[EHT Insulation Audit Scope]",
                            ("Basic", "Standard", "Advanced"),
                            index=insul_lvl_index,
                            placeholder='"MANUAL INPUT REQ."'
                            )

                    with col_insinfo_2:
                        EHT_ins_no = st.text_input('#Circuits',value = str(ins_cnt), key='insqty')

                with st.expander(":green[EHT Insulation Visual Inspection]"):
                    col_insvis_1, col_insvis_2, col_insvis_3, col_insvis_4 = st.columns([1,1,1,1])

                    with col_insvis_1:
                        visinsul_majority = st.selectbox(
                            "EHT insulation Majority-Status",
                            ("correct", "incorrect"),
                            index = visinsul_majority_index,
                            help="Majority - more than 50%",
                            key='visinsul_majority'
                        )
                    with col_insvis_2:
                        visinsulfail_gen = st.selectbox(
                            "General reported faults",
                            ("y", "n"),
                            index = visinsulfail_gen_index,
                            help="Are there any faults detected with standard inspection?",
                            key = 'visinsulfail_gen'
                    )
                    if (insul_lvl_index == 1 or insul_lvl_index == 2):
                        with col_insvis_3:
                            visinsulfail_ifr = st.selectbox(
                                ":red[Infrared reported faults]",
                                ("y", "n"),
                                index = visinsulfail_ifr_index,
                                help="Are there any faults detected with infrared camera?",
                                key = 'visinsulfail_ifr'
                        )

                if insul_lvl_index == 2:
                    with st.expander(":green[EHT Insulation Parameters Check]"):
                        col_insparam_1, col_insparam_2, col_insparam_3, col_insparam_4 = st.columns([1,1,1,1])

                        with col_insparam_1:
                            inscheck_majority = st.selectbox(
                                "EHT insulation Parameters Majority-Status",
                                ("correct", "incorrect"),
                                index = inscheck_majority_index,
                                help="Majority - more than 50%",
                                key='insulparam_majority'
                            )
                        with col_insparam_2:
                            incheck_some = st.selectbox(
                                "Reported discrepancies",
                                ("y", "n"),
                                index = incheck_some_index,
                                help="Are there any inconsistencies between reality and documentation?",
                                key = 'insulparam_fault'
                        )                    
#####################################################################################
#####################################################################################
#       GENERATING REPORT       #

            reason_comp = 'n'
            reason_risk = 'n'
            reason_maint = 'n'
            reason_retro = 'n'
            reason_optym = 'n'

            if 'Compliance and safety' in prescreen_selection:
                reason_comp = 'y'
            if 'Risk mitigation' in prescreen_selection:
                reason_risk = 'y'
            if 'Maintenance insights' in prescreen_selection:
                reason_maint = 'y'
            if 'System retrofit' in prescreen_selection:
                reason_retro = 'y'
            if 'System optimization' in prescreen_selection:
                reason_optym = 'y'



            if generate_button:
                #st.sidebar.write('Generate!')
                #st.sidebar.success("Form Dispatched!")
                if (not prescreen_selection):
                    st.sidebar.warning("Please enter reasons for this audit")
                elif (lay_avail == 'y' and plotplan_img_attach is None):
                    st.sidebar.warning("Please attach available layout")
                elif (lay_avail == 'n' and plotplan_img_attach is not None):
                    st.sidebar.warning("Since layout was attached, it is clearly available")
                elif (audited_375 == 'y' and audit_lvl_panel == None):
                    st.sidebar.warning("Please choose Panel Audit Scope")
                elif (audited_373 == 'y' and audit_lvl_cct == None):
                    st.sidebar.warning("Please choose Circuit Audit Scope")
                elif (audited_374 == 'y' and audit_lvl_insul == None):
                    st.sidebar.warning("Please choose Insulation Audit Scope")
                elif (audited_373 == 'y' and curr_trip==None):
                    st.sidebar.warning("Please provide info about CB trips")
                else:
                    center_running()
                    time.sleep(2)
                    st.session_state.load_state = True
                    document = Document("docx_template.docx")
                                
                                
                    a_df_file_content = open_file('s3-nvent-prontoforms-data/Data_sources/audit_rep_text.csv')
                    a_df = pd.read_csv(a_df_file_content, encoding="utf-8")
                    a_df.set_index('Phrase',inplace=True)
                    lang = 'EN'
    #@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@               
                    default_font = "Arial"
                    header = document.sections[0].header
                    header_par = header.paragraphs[0]
                    header_par.text = f'\t\t{compname}\n \t\t{projname}\n \t\t{projpono}'
                    header_par_style = header_par.style
                    header_par_style.font.name = default_font
                    header_par_style.font.size = Pt(9)
                    header_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    footer_par = document.sections[0].footer.paragraphs[0]
                    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    report_title = document.add_paragraph()
                    text_df = str(a_df.loc['report_title',lang]).replace('\\n', '\n')
                    report_title_run = report_title.add_run(text_df)
                    #report_title_run = report_title.add_run("ELECTRICAL HEAT TRACING\n AUDIT REPORT")
                    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    report_title_run.font.name = default_font
                    report_title_run.font.size = Pt(24)
                    report_title_run.font.color.rgb = RGBColor(196, 38, 46)
                    report_title.paragraph_format.space_after = Pt(36)

                    nvent_logo = document.add_picture('nvent_raychem.png')
                    nvent_logo_par = document.paragraphs[-1]
                    nvent_logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    nvent_logo_par.paragraph_format.space_after = Pt(24)

                    nvent_industrial = document.add_picture('industrial_red.png',width=Inches(6.5),height=Inches(1.8))
                    nvent_industrial_par = document.paragraphs[-1]
                    nvent_industrial_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    nvent_industrial_par.paragraph_format.space_after = Pt(24)

                    company_details = document.add_paragraph()
                    company_details_run_1 = company_details.add_run(f'\t\t{compname}\n \t\t{clientname}')
                    company_details_run_1.bold = True
                    company_details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    company_details_run_1.font.name = default_font
                    company_details_run_1.font.size = Pt(12)
                    company_details.paragraph_format.space_after = Pt(36)

                    company_address = document.add_paragraph()
                    company_address_run_1 = company_address.add_run(f'\t\t{compstreet}\n \t\t{compzip} {compcity}\n \t\t{compcountry}')
                    company_address.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    company_address_run_1.font.name = default_font
                    company_address_run_1.font.size = Pt(12)
                    company_address.paragraph_format.space_after = Pt(36)

                    nvent_ref = document.add_paragraph()
                    text_df1 = a_df.loc['nvent_ref',lang]
                    text_df2 = a_df.loc['client_ref',lang]
                    nvent_ref_run_1 = nvent_ref.add_run(f'\t\t{text_df1}{projpono}\n \t{text_df2}{projsapno}\n\n \t\t{repdate}')
                    #nvent_ref_run_1 = nvent_ref.add_run(f'\t\tnVent Ref:{projpono}\n \tClient Ref:{projsapno}\n\n \t\t{repdate}')
                    nvent_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    nvent_ref_run_1.font.name = default_font
                    nvent_ref_run_1.font.size = Pt(12)

                    ###################################
                    ### SECTION - TABLE OF CONTENTS ###
                    ###################################

                    section_toc = document.add_section()
                    document.add_page_break()
                    ##########################
                    ### SECTION - CONTACTS ###
                    ##########################

                    section_contact = document.add_section()


                    contact_heading = document.add_heading('',level=1)
                    contact_heading_run = contact_heading.add_run(f"1. {a_df.loc['contact_details',lang]}")
                    contact_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    contact_heading_run.font.name = default_font
                    contact_heading_run.font.size = Pt(18)
                    contact_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    contact_heading.paragraph_format.space_after = Pt(10)

                        ## TABLE - PROJECT INFO or PLACE

                    project_records = (
                        (a_df.loc['company',lang], compname),
                        (a_df.loc['address',lang], f"{compstreet} {compcity} {compcountry}" ),
                        (a_df.loc['zip',lang], compzip),
                        (a_df.loc['place',lang], '')
                    )

                    project_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    prj_cells = project_table.rows[0].cells
                    prj_cells[0].text = a_df.loc['project',lang]
                    prj_cells[1].text = ''
                    for descr, detail in project_records:
                        prj_cells = project_table.add_row().cells
                        prj_cells[0].text = descr
                        prj_cells[1].text = detail
                    for cell in project_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in project_table.columns[1].cells:
                        cell.width = Cm(10)

                    project_table_space = document.add_paragraph()
                    project_table_space.paragraph_format.space_after = Pt(10)

        ## TABLE - CUSTOMER CONTACT PERSON 1

                    cust1_records = (
                        (a_df.loc['client_name',lang], clientname),
                        (a_df.loc['client_phone',lang], ''),
                        (a_df.loc['client_email',lang], ''),
                        (a_df.loc['client_function',lang], '')
                    )

                    cust1_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    cust1_cells = cust1_table.rows[0].cells
                    cust1_cells[0].text = a_df.loc['client_rep',lang]
                    cust1_cells[1].text = ''
                    for descr, detail in cust1_records:
                        cust1_cells = cust1_table.add_row().cells
                        cust1_cells[0].text = descr
                        cust1_cells[1].text = detail
                    for cell in cust1_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in cust1_table.columns[1].cells:
                        cell.width = Cm(10)
                    cust1_table_space = document.add_paragraph()
                    cust1_table_space.paragraph_format.space_after = Pt(10)

                    pm_records = (
                        (a_df.loc['pm_name',lang], ''),
                        (a_df.loc['pm_phone',lang], ''),
                        (a_df.loc['pm_email',lang], '')
                    )

                    pm_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    pm_cells = pm_table.rows[0].cells
                    pm_cells[0].text = a_df.loc['nvent_pm',lang]
                    pm_cells[1].text = ''
                    for descr, detail in pm_records:
                        pm_cells = pm_table.add_row().cells
                        pm_cells[0].text = descr
                        pm_cells[1].text = detail
                    for cell in pm_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in pm_table.columns[1].cells:
                        cell.width = Cm(10)
                    pm_table_space = document.add_paragraph()
                    pm_table_space.paragraph_format.space_after = Pt(10)

                        ## TABLE - nVent PSC

                    psc_records = (
                        (a_df.loc['psc_name',lang], ''),
                        (a_df.loc['psc_phone',lang], ''),
                        (a_df.loc['psc_email',lang], '')
                    )

                    psc_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    psc_cells = psc_table.rows[0].cells
                    psc_cells[0].text = a_df.loc['nvent_psc',lang]
                    psc_cells[1].text = ''
                    for descr, detail in psc_records:
                        psc_cells = psc_table.add_row().cells
                        psc_cells[0].text = descr
                        psc_cells[1].text = detail
                    for cell in psc_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in psc_table.columns[1].cells:
                        cell.width = Cm(10)
                    psc_table_space = document.add_paragraph()
                    psc_table_space.paragraph_format.space_after = Pt(10)

                        ## TABLE - nVent Auditor

                    auditor_records = (
                        (a_df.loc['auditor_name',lang], ''),
                        (a_df.loc['auditor_phone',lang], ''),
                        (a_df.loc['auditor_email',lang], '')
                    )

                    auditor_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    auditor_cells = auditor_table.rows[0].cells
                    auditor_cells[0].text = a_df.loc['nvent_auditor',lang]
                    auditor_cells[1].text = ''
                    for descr, detail in psc_records:
                        auditor_cells = auditor_table.add_row().cells
                        auditor_cells[0].text = descr
                        auditor_cells[1].text = detail
                    for cell in auditor_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in auditor_table.columns[1].cells:
                        cell.width = Cm(10)
                    auditor_table_space = document.add_paragraph()
                    auditor_table_space.paragraph_format.space_after = Pt(10)

                        ## TABLE - nVent Sales Rep

                    seller_records = (
                        (a_df.loc['sales_name',lang], ''),
                        (a_df.loc['sales_phone',lang], ''),
                        (a_df.loc['sales_email',lang], '')
                    )

                    seller_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    seller_cells = seller_table.rows[0].cells
                    seller_cells[0].text = a_df.loc['nvent_sales',lang]
                    seller_cells[1].text = ''
                    for descr, detail in seller_records:
                        seller_cells = seller_table.add_row().cells
                        seller_cells[0].text = descr
                        seller_cells[1].text = detail
                    for cell in seller_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in seller_table.columns[1].cells:
                        cell.width = Cm(10)
                    seller_table_space = document.add_paragraph()
                    seller_table_space.paragraph_format.space_after = Pt(10)

                        ## TABLE - nVent Office

                    office_records = (
                        (a_df.loc['office_street',lang], ''),
                        (a_df.loc['office_zip',lang], ''),
                        (a_df.loc['office_country',lang], ''),
                        (a_df.loc['office_phone',lang], ''),
                        (a_df.loc['office_website',lang], '')
                    )

                    office_table = document.add_table(rows=1, cols=2,style='GridTable3-Accent3')
                    office_cells = office_table.rows[0].cells
                    office_cells[0].text = a_df.loc['nvent_office',lang]
                    office_cells[1].text = ''
                    for descr, detail in office_records:
                        office_cells = office_table.add_row().cells
                        office_cells[0].text = descr
                        office_cells[1].text = detail
                    for cell in office_table.columns[0].cells:
                        cell.width = Cm(5)
                    for cell in office_table.columns[1].cells:
                        cell.width = Cm(10)

                    ##################################
                    ### SECTION - SCOPE DEFINITION ###
                    ##################################

                    section_scope = document.add_section()
                    scope_heading = document.add_heading('',level=1)
                    scope_heading_run = scope_heading.add_run(f"2. {a_df.loc['scope_section',lang]}")
                    scope_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    scope_heading_run.font.name = default_font
                    scope_heading_run.font.size = Pt(18)
                    scope_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    scope_heading.paragraph_format.space_after = Pt(10)

                        ## Description paragraph
                    paragh_scope_def1 = document.add_paragraph()
                    paragh_scope_def1.add_run(a_df.loc['scopedef_p1',lang])
                    paragh_scope_def1.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    paragh_scope_panel = document.add_paragraph(a_df.loc['scopedef_p1_b1',lang],style='List Bullet')
                    paragh_scope_panel.paragraph_format.left_indent = Inches(0.5)
                    paragh_scope_cct = document.add_paragraph(a_df.loc['scopedef_p1_b2',lang],style='List Bullet')
                    paragh_scope_cct.paragraph_format.left_indent = Inches(0.5)
                    paragh_scope_ins = document.add_paragraph(a_df.loc['scopedef_p1_b3',lang],style='List Bullet')
                    paragh_scope_ins.paragraph_format.left_indent = Inches(0.5)

                    paragh_scope_def2 = document.add_paragraph()
                    paragh_scope_def2.add_run(a_df.loc['scopedef_p2',lang])
                    paragh_scope_def2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    audit_lvls = document.add_picture('3_levels.png')
                    audit_lvls_par = document.paragraphs[-1]
                    audit_lvls_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    audit_lvls_par.paragraph_format.space_after = Pt(24)

                        ## TABLE - EHT PANEL AUDIT SCOPE: ACTIVITIES

                    ehtpanelact_records = (
                        (a_df.loc['scopepnla_t11',lang], str(a_df.loc['scopepnla_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopepnla_t21',lang], str(a_df.loc['scopepnla_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopepnla_t31',lang], str(a_df.loc['scopepnla_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtpanelact_table_desc = document.add_paragraph(a_df.loc['scopepnla_tdesc',lang])
                    ehtpanelact_table_desc.paragraph_format.space_after = Pt(1)
                    ehtpanelact_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtpanelact_cells = ehtpanelact_table.rows[0].cells
                    ehtpanelact_cells[0].text = a_df.loc['scopepnla_tc1',lang]
                    ehtpanelact_cells[1].text = a_df.loc['scopepnla_tc2',lang]
                    for descr, detail in ehtpanelact_records:
                        ehtpanelact_cells = ehtpanelact_table.add_row().cells
                        ehtpanelact_cells[0].text = descr
                        ehtpanelact_cells[1].text = detail
                    for cell in ehtpanelact_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtpanelact_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtpanelact_table_space = document.add_paragraph()

                    document.add_page_break()

                        ## TABLE - EHT PANEL AUDIT SCOPE: DELIVERABLES

                    ehtpaneldel_records = (
                        (a_df.loc['scopepnld_t11',lang], str(a_df.loc['scopepnld_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopepnld_t21',lang], str(a_df.loc['scopepnld_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopepnld_t31',lang], str(a_df.loc['scopepnld_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtpaneldel_table_desc = document.add_paragraph(a_df.loc['scopepnld_tdesc',lang])
                    ehtpaneldel_table_desc.paragraph_format.space_after = Pt(1)
                    ehtpaneldel_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtpaneldel_cells = ehtpaneldel_table.rows[0].cells
                    ehtpaneldel_cells[0].text = a_df.loc['scopepnld_tc1',lang]
                    ehtpaneldel_cells[1].text = a_df.loc['scopepnld_tc1',lang]
                    for descr, detail in ehtpaneldel_records:
                        ehtpaneldel_cells = ehtpaneldel_table.add_row().cells
                        ehtpaneldel_cells[0].text = descr
                        ehtpaneldel_cells[1].text = detail
                    for cell in ehtpaneldel_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtpaneldel_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtpaneldel_table_space = document.add_paragraph()
                    ehtpaneldel_table_space.paragraph_format.space_after = Pt(20)

        ## TABLE - EHT CIRCUIT AUDIT SCOPE: ACTIVITIES

                    ehtcctact_records = (
                        (a_df.loc['scopeccta_t11',lang], str(a_df.loc['scopeccta_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeccta_t21',lang], str(a_df.loc['scopeccta_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeccta_t31',lang], str(a_df.loc['scopeccta_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtcctact_table_desc = document.add_paragraph(a_df.loc['scopeccta_tdesc',lang])
                    ehtcctact_table_desc.paragraph_format.space_after = Pt(1)
                    ehtcctact_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtcctact_cells = ehtcctact_table.rows[0].cells
                    ehtcctact_cells[0].text = a_df.loc['scopeccta_tc1',lang]
                    ehtcctact_cells[1].text = a_df.loc['scopeccta_tc2',lang]
                    for descr, detail in ehtcctact_records:
                        ehtcctact_cells = ehtcctact_table.add_row().cells
                        ehtcctact_cells[0].text = descr
                        ehtcctact_cells[1].text = detail
                    for cell in ehtcctact_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtcctact_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtcctact_table_space = document.add_paragraph()
                    ehtcctact_table_space.paragraph_format.space_after = Pt(5)

                        ## TABLE - EHT CIRCUIT AUDIT SCOPE: DELIVERABLES

                    ehtcctdel_records = (
                        (a_df.loc['scopecctd_t11',lang], str(a_df.loc['scopecctd_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopecctd_t21',lang], str(a_df.loc['scopecctd_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopecctd_t31',lang], str(a_df.loc['scopecctd_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtcctdel_table_desc = document.add_paragraph(a_df.loc['scopecctd_tdesc',lang])
                    ehtcctdel_table_desc.paragraph_format.space_after = Pt(1)
                    ehtcctdel_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtcctdel_cells = ehtcctdel_table.rows[0].cells
                    ehtcctdel_cells[0].text = a_df.loc['scopecctd_tc1',lang]
                    ehtcctdel_cells[1].text = a_df.loc['scopecctd_tc2',lang]
                    for descr, detail in ehtcctdel_records:
                        ehtcctdel_cells = ehtcctdel_table.add_row().cells
                        ehtcctdel_cells[0].text = descr
                        ehtcctdel_cells[1].text = detail
                    for cell in ehtcctdel_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtcctdel_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtcctdel_table_space = document.add_paragraph()
                    ehtcctdel_table_space.paragraph_format.space_after = Pt(20)

                        ## TABLE - EHT INSULATION AUDIT SCOPE: ACTIVITIES

                    ehtinsulact_records = (
                        (a_df.loc['scopeinsa_t11',lang], str(a_df.loc['scopeinsa_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeinsa_t21',lang], str(a_df.loc['scopeinsa_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeinsa_t31',lang], str(a_df.loc['scopeinsa_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtinsulact_table_desc = document.add_paragraph(a_df.loc['scopeinsa_tdesc',lang])
                    ehtinsulact_table_desc.paragraph_format.space_after = Pt(1)
                    ehtinsulact_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtinsulact_cells = ehtinsulact_table.rows[0].cells
                    ehtinsulact_cells[0].text = a_df.loc['scopeinsa_tc1',lang]
                    ehtinsulact_cells[1].text = a_df.loc['scopeinsa_tc2',lang]
                    for descr, detail in ehtinsulact_records:
                        ehtinsulact_cells = ehtinsulact_table.add_row().cells
                        ehtinsulact_cells[0].text = descr
                        ehtinsulact_cells[1].text = detail
                    for cell in ehtinsulact_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtinsulact_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtinsulact_table_space = document.add_paragraph()

                    document.add_page_break()

                        ## TABLE - EHT INSULATION AUDIT SCOPE: DELIVERABLES

                    ehtinsuldel_records = (
                        (a_df.loc['scopeinsd_t11',lang], str(a_df.loc['scopeinsd_t12',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeinsd_t21',lang], str(a_df.loc['scopeinsd_t22',lang]).replace('\\n', '\n')),
                        (a_df.loc['scopeinsd_t31',lang], str(a_df.loc['scopeinsd_t32',lang]).replace('\\n', '\n'))
                    )

                    ehtinsuldel_table_desc = document.add_paragraph(a_df.loc['scopeinsd_tdesc',lang])
                    ehtinsuldel_table_desc.paragraph_format.space_after = Pt(1)
                    ehtinsuldel_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    ehtinsuldel_cells = ehtinsuldel_table.rows[0].cells
                    ehtinsuldel_cells[0].text = a_df.loc['scopeinsd_tc1',lang]
                    ehtinsuldel_cells[1].text = a_df.loc['scopeinsd_tc2',lang]
                    for descr, detail in ehtinsuldel_records:
                        ehtinsuldel_cells = ehtinsuldel_table.add_row().cells
                        ehtinsuldel_cells[0].text = descr
                        ehtinsuldel_cells[1].text = detail
                    for cell in ehtinsuldel_table.columns[0].cells:
                        cell.width = Cm(2.5)
                    for cell in ehtinsuldel_table.columns[1].cells:
                        cell.width = Cm(12.5)
                    ehtinsuldel_table_space = document.add_paragraph()
                    #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                    ehtinsuldel_table_space.paragraph_format.space_after = Pt(20)

                    paragh_audit_select = document.add_paragraph()
                    paragh_audit_select.add_run(
                            a_df.loc['scope_choice_p1',lang].format(compname))
                    # paragh_audit_select.add_run(
                    #         f"For this specific Audit {compname} have chosen:")
                    paragh_audit_select.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragh_audit_select.paragraph_format.space_after = Pt(1)

                    selection_records = (
                        (a_df.loc['scope_choice_t1',lang], audit_lvl_panel),
                        (a_df.loc['scope_choice_t2',lang], audit_lvl_cct),
                        (a_df.loc['scope_choice_t3',lang], audit_lvl_insul)
                    )

                    audit_select_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                    audit_select_cells = audit_select_table.rows[0].cells
                    audit_select_cells[0].text = a_df.loc['scope_choice_c1',lang]
                    audit_select_cells[1].text = a_df.loc['scope_choice_c2',lang]
                    for descr, detail in selection_records:
                        audit_select_cells = audit_select_table.add_row().cells
                        audit_select_cells[0].text = descr
                        audit_select_cells[1].text = detail
                    for cell in audit_select_table.columns[0].cells:
                        cell.width = Cm(4)
                    for cell in audit_select_table.columns[1].cells:
                        cell.width = Cm(3)
                    audit_select_table_space = document.add_paragraph()
                    audit_select_table_space.paragraph_format.space_after = Pt(10)

                    ##############################################
                    ### SECTION - EHT INSTALLATION DESCRIPTION ###
                    ##############################################

                    section_eht_descr = document.add_section()
                    eht_descr_heading = document.add_heading('',level=1)
                    eht_descr_heading_run = eht_descr_heading.add_run(f"3. {a_df.loc['descr_section',lang]}")
                    eht_descr_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    eht_descr_heading_run.font.name = default_font
                    eht_descr_heading_run.font.size = Pt(18)
                    eht_descr_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    eht_descr_heading.paragraph_format.space_after = Pt(10)

                    paragh_eht_descr1 = document.add_paragraph()
                    paragh_eht_descr1_run = paragh_eht_descr1.add_run(a_df.loc['descrinst_p1',lang])
                    paragh_eht_descr1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragh_eht_descr1.paragraph_format.space_after = Pt(5)

                    if reason_comp == 'y':
                        paragh_ehtdescr_compl = document.add_paragraph(a_df.loc['descrinst_p1_b1',lang],style='List Bullet')
                        paragh_ehtdescr_compl.paragraph_format.left_indent = Inches(0.5)
                    if reason_risk == 'y':
                        paragh_ehtdescr_risk = document.add_paragraph(a_df.loc['descrinst_p1_b2',lang],style='List Bullet')
                        paragh_ehtdescr_risk.paragraph_format.left_indent = Inches(0.5)
                    if reason_maint == 'y':
                        paragh_ehtdescr_maintain = document.add_paragraph(a_df.loc['descrinst_p1_b3',lang],style='List Bullet')
                        paragh_ehtdescr_maintain.paragraph_format.left_indent = Inches(0.5)
                    if reason_retro == 'y':
                        paragh_ehtdescr_retro = document.add_paragraph(a_df.loc['descrinst_p1_b4',lang],style='List Bullet')
                        paragh_ehtdescr_retro.paragraph_format.left_indent = Inches(0.5)
                    if reason_optym == 'y':
                        paragh_ehtdescr_optym = document.add_paragraph(a_df.loc['descrinst_p1_b5',lang],style='List Bullet')
                        paragh_ehtdescr_optym.paragraph_format.left_indent = Inches(0.5)

                    if (audit_lvl_panel != 'None' and audit_lvl_cct != 'None'): 
                        paragh_eht_descr2 = document.add_paragraph()
                        paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc',lang].format(EHT_cct_no))
                        if EHT_cct_no == '1':
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc_1a',lang])
                        else:
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc_1b',lang])
                        paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc_2',lang].format(EHT_pnl_no))
                        if EHT_pnl_no == '1':
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc_2a',lang])
                        else:
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2pc_2b',lang])
                        paragh_eht_descr2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if (audit_lvl_panel == 'None' and audit_lvl_cct != 'None'): 
                        paragh_eht_descr2 = document.add_paragraph()
                        paragh_eht_descr2.add_run(a_df.loc['descrinst_p2c',lang].format(EHT_cct_no))
                        if EHT_cct_no == '1':
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2c_1a',lang])
                        else:
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2c_1b',lang])
                        paragh_eht_descr2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if (audit_lvl_panel != 'None' and audit_lvl_cct == 'None'): 
                        paragh_eht_descr2 = document.add_paragraph()
                        paragh_eht_descr2.add_run(a_df.loc['descrinst_p2p',lang].format(EHT_pnl_no))
                        if EHT_pnl_no == '1':
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2p_1a',lang])
                        else:
                            paragh_eht_descr2.add_run(a_df.loc['descrinst_p2p_1b',lang])
                        paragh_eht_descr2.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if audit_lvl_cct != 'None':
                        #ctrl_loc = 'panel', 'field', 'both'
                        if (ctrl_el == 'y' and ctrl_mech == 'y'):
                            paragh_eht_descr3 = document.add_paragraph()
                            paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1em',lang])
                            if ctrl_loc == 'panel':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1em_p',lang])
                            if ctrl_loc == 'field':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1em_f',lang])
                            if ctrl_loc == 'both':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1em_pf',lang])
                            if ctrl_not == 'y':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1em_un',lang])
                            paragh_eht_descr3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (ctrl_el == 'y' and ctrl_mech == 'n'):
                            paragh_eht_descr3 = document.add_paragraph()
                            paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1e',lang])
                            if ctrl_loc == 'panel':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1e_p',lang])
                            if ctrl_loc == 'field':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1e_f',lang])
                            if ctrl_loc == 'both':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1e_pf',lang])
                            if ctrl_not == 'y':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1e_un',lang])
                            paragh_eht_descr3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (ctrl_el == 'n' and ctrl_mech == 'y'):
                            paragh_eht_descr3 = document.add_paragraph()
                            paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1m',lang])
                            if ctrl_loc == 'panel':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1m_p',lang])
                            if ctrl_loc == 'field':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1m_f',lang])
                            if ctrl_loc == 'both':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1m_pf',lang])
                            if ctrl_not == 'y':
                                paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1m_un',lang])
                            paragh_eht_descr3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            paragh_eht_descr3 = document.add_paragraph()
                            paragh_eht_descr3.add_run(a_df.loc['descrctrl_p1un',lang])
                            paragh_eht_descr3.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if (pid_avail == 'y' or iso_avail == 'y' or pnl_avail == 'y' or lay_avail == 'y' or calc_avail == 'y'):
                        paragh_eht_descr4 = document.add_paragraph()
                        paragh_eht_descr4.add_run(a_df.loc['descrdoc_p1_piplc',lang])
                        paragh_eht_descr4.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragh_eht_descr4.paragraph_format.space_after = Pt(5)

                        if pid_avail == 'y':
                            paragh_pid_avail = document.add_paragraph(a_df.loc['descrdoc_p1_pid',lang],style='List Bullet')
                            paragh_pid_avail.paragraph_format.left_indent = Inches(0.5)
                        if iso_avail == 'y':
                            paragh_iso_avail = document.add_paragraph(a_df.loc['descrdoc_p1_iso',lang],style='List Bullet')
                            paragh_iso_avail.paragraph_format.left_indent = Inches(0.5)
                        if pnl_avail == 'y':
                            paragh_pnl_avail = document.add_paragraph(a_df.loc['descrdoc_p1_pnl',lang],style='List Bullet')
                            paragh_pnl_avail.paragraph_format.left_indent = Inches(0.5)
                        if lay_avail == 'y':
                            paragh_lay_avail = document.add_paragraph(a_df.loc['descrdoc_p1_lay',lang],style='List Bullet')
                            paragh_lay_avail.paragraph_format.left_indent = Inches(0.5)
                        if calc_avail == 'y':
                            paragh_calc_avail = document.add_paragraph(a_df.loc['descrdoc_p1_calc',lang],style='List Bullet')
                            paragh_calc_avail.paragraph_format.left_indent = Inches(0.5)
                    else:
                        paragh_eht_descr4 = document.add_paragraph()
                        paragh_eht_descr4.add_run(a_df.loc['descrdoc_p1_no',lang])
                        paragh_eht_descr4.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragh_eht_descr4.paragraph_format.space_after = Pt(5)
                    #################################
                    ## SUB-SECTION - EHT PLOT PLAN ##
                    #################################

                    section_plotplan = document.add_section()

                    plotplan_heading = document.add_heading('',level=2)
                    section_plotplan.orientation = WD_ORIENT.LANDSCAPE
                    section_plotplan.page_width = 10058400
                    section_plotplan.page_height = 7772400

                    plotplan_heading_run = plotplan_heading.add_run(f"3.1 {a_df.loc['descr_subsection',lang]}")
                    plotplan_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    plotplan_heading_run.font.name = default_font
                    plotplan_heading_run.font.size = Pt(16)
                    plotplan_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    plotplan_heading.paragraph_format.space_after = Pt(10)

                    if lay_avail == 'y':
                        plotplan_img = document.add_picture(plotplan_img_rep,width=Inches(8.5),height=Inches(5.8))
                        plotplan_img_par = document.paragraphs[-1]
                        plotplan_img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        #plotplan_img_par.paragraph_format.space_after = Pt(24)
                    else:
                        plotplan_info = document.add_paragraph()
                        plotplan_info_run = plotplan_info.add_run(str(a_df.loc['descr_noplotplan',lang]).replace('\\n', '\n'))
                        #company_details_run_1.bold = True
                        plotplan_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        plotplan_info_run.font.name = default_font
                        plotplan_info_run.font.size = Pt(24)
                        plotplan_info_run.font.color.rgb = RGBColor(196, 38, 46)
                        plotplan_info.paragraph_format.space_after = Pt(36)

                    ####################################
                    ### SECTION - VISUAL INSPECTIONS ###
                    ####################################

                    vis_sec_no = 4

                    s_visual = document.add_section()
                    s_visual.orientation = WD_ORIENT.PORTRAIT
                    s_visual.page_width = 7772400
                    s_visual.page_height = 10058400
                    h_visual = document.add_heading('',level=1)
                    h_visual_r = h_visual.add_run(a_df.loc['vis_section',lang].format(vis_sec_no))
                    h_visual.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    h_visual_r.font.name = default_font
                    h_visual_r.font.size = Pt(18)
                    h_visual_r.font.color.rgb = RGBColor(0, 0, 0)
                    h_visual.paragraph_format.space_after = Pt(10)

                    p_vis1 = document.add_paragraph()
                    p_vis1.add_run(a_df.loc['vis_p1',lang])
                    p_vis1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_vis1.paragraph_format.space_after = Pt(10)

                    p_vis2 = document.add_paragraph()
                    p_vis2.add_run(a_df.loc['vis_p2',lang])
                    p_vis2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_vis2.paragraph_format.space_after = Pt(5)

                    if audit_lvl_panel != 'None':
                        p_vispanel_std = document.add_paragraph(a_df.loc['vis_p2_pnls',lang],style='List Bullet')
                        p_vispanel_std.paragraph_format.left_indent = Inches(0.5)
                    if audit_lvl_panel == 'Advanced':
                        p_vispanel_ifr = document.add_paragraph(a_df.loc['vis_p2_pnlir',lang],style='List Bullet')
                        p_vispanel_ifr.paragraph_format.left_indent = Inches(0.5)
                    if audit_lvl_cct != 'None':
                        p_viscct_std = document.add_paragraph(a_df.loc['vis_p2_cct',lang],style='List Bullet')
                        p_viscct_std.paragraph_format.left_indent = Inches(0.5)
                    if audit_lvl_insul != 'None':
                        p_visinsul_std = document.add_paragraph(a_df.loc['vis_p2_inss',lang],style='List Bullet')
                        p_visinsul_std.paragraph_format.left_indent = Inches(0.5)
                    if (audit_lvl_insul == 'Standard' or audit_lvl_insul == 'Advanced'):
                        p_visinsul_ifr = document.add_paragraph(a_df.loc['vis_p2_insir',lang],style='List Bullet')
                        p_visinsul_ifr.paragraph_format.left_indent = Inches(0.5)

                    p_vis3 = document.add_paragraph()
                    p_vis3.paragraph_format.space_after = Pt(30)
                        
                    vis_subsec_no = 1

                    #############################################
                    ## SUB- SECTION - PANEL VISUAL INSPECTIONS ##
                    #############################################

                    if audit_lvl_panel != 'None':
                        s_vispanel = document.add_section()
                        s_vispanel.orientation = WD_ORIENT.PORTRAIT
                        s_vispanel.page_width = 7772400
                        s_vispanel.page_height = 10058400
                        h_vispanel = document.add_heading('',level=2)
                        if (audit_lvl_panel == 'Basic' or audit_lvl_panel == 'Standard'):
                            h_vispanel_r = h_vispanel.add_run(a_df.loc['vispnl_subsection_s',lang].format(vis_sec_no, vis_subsec_no))
                        else:
                            h_vispanel_r = h_vispanel.add_run(a_df.loc['vispnl_subsection_ir',lang].format(vis_sec_no, vis_subsec_no))    
                        h_vispanel.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_vispanel_r.font.name = default_font
                        h_vispanel_r.font.size = Pt(16)
                        h_vispanel_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_vispanel.paragraph_format.space_after = Pt(10)

                        p_vispanel1 = document.add_paragraph()
                        p_vispanel1.add_run(
                                "Each panel is examined considering i.a. below categories:")
                        p_vispanel1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_vispanel1.paragraph_format.space_after = Pt(5)   
                        
                        b1_vispanel1 = document.add_paragraph(style='List Bullet')
                        b1_vispanel1.add_run(str(a_df.loc['vispnl_cat_acc',lang]).replace('\\n', '\n'))
                        b1_vispanel1.paragraph_format.left_indent = Inches(0.5)
                        b2_vispanel1 = document.add_paragraph(style='List Bullet')
                        b2_vispanel1.add_run(str(a_df.loc['vispnl_cat_ex',lang]).replace('\\n', '\n'))
                        b2_vispanel1.paragraph_format.left_indent = Inches(0.5)
                        b3_vispanel1 = document.add_paragraph(style='List Bullet')
                        b3_vispanel1.add_run(str(a_df.loc['vispnl_cat_in',lang]).replace('\\n', '\n'))
                        if audit_lvl_panel == 'Advanced':
                            b3_vispanel1.add_run(str(a_df.loc['vispnl_cat_in_ir',lang]).replace('\\n', '\n'))        
                        b3_vispanel1.paragraph_format.left_indent = Inches(0.5)

                        b3_vispanel1.paragraph_format.space_after = Pt(30)

                    ###################################################
                    # SUB2-SECTION - PANEL VISUAL INSPECTION FINDINGS #
                    ###################################################
                        
                        h_vispanelfind = document.add_heading('',level=3)
                        h_vispanelfind_r = h_vispanelfind.add_run(a_df.loc['vispnl_subsection2_find',lang].format(vis_sec_no, vis_subsec_no))
                        h_vispanelfind.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_vispanelfind_r.font.name = default_font
                        h_vispanelfind_r.font.size = Pt(14)
                        h_vispanelfind_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_vispanelfind.paragraph_format.space_after = Pt(10)
                        
                        vispanel_faults = False
                        
                        if (audit_lvl_panel == 'Basic' or audit_lvl_panel == 'Standard'):
                            if vispanel_majority == 'correct':
                                if vispanelfail_gen == 'n':
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_finds_corrn',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                if vispanelfail_gen == 'y':
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_finds_corry',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                    
                                    vispanel_faults = True
                                    p_vispanelfind2 = document.add_paragraph()
                                    p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                    p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind2.paragraph_format.space_after = Pt(10)
                                    
                            if vispanel_majority == 'incorrect':
                                p_vispanelfind1 = document.add_paragraph()
                                p_vispanelfind1.add_run(a_df.loc['vispnl_finds_incorr',lang])
                                p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                
                                vispanel_faults = True
                                p_vispanelfind2 = document.add_paragraph()
                                p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_vispanelfind2.paragraph_format.space_after = Pt(10)
                                
                        if audit_lvl_panel == 'Advanced':
                            if vispanel_majority == 'correct':
                                if (vispanelfail_gen == 'n' and vispanelfail_ifr == 'n'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_corrnn',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                if (vispanelfail_gen == 'y' and vispanelfail_ifr == 'n'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_corryn',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                    
                                    vispanel_faults = True
                                    p_vispanelfind2 = document.add_paragraph()
                                    p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                    p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (vispanelfail_gen == 'n' and vispanelfail_ifr == 'y'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_corrny',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                    
                                    vispanel_faults = True
                                    p_vispanelfind2 = document.add_paragraph()
                                    p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                    p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (vispanelfail_gen == 'y' and vispanelfail_ifr == 'y'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_corryy',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                    
                                    vispanel_faults = True
                                    p_vispanelfind2 = document.add_paragraph()
                                    p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                    p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind2.paragraph_format.space_after = Pt(10)
                                    
                            if vispanel_majority == 'incorrect':
                                if (vispanelfail_gen == 'y' and vispanelfail_ifr == 'n'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_incorryn',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                if (vispanelfail_gen == 'n' and vispanelfail_ifr == 'y'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_incorrny',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                if (vispanelfail_gen == 'y' and vispanelfail_ifr == 'y'):
                                    p_vispanelfind1 = document.add_paragraph()
                                    p_vispanelfind1.add_run(a_df.loc['vispnl_findir_incorryy',lang])
                                    p_vispanelfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind1.paragraph_format.space_after = Pt(10)
                                    
                                    vispanel_faults = True
                                    p_vispanelfind2 = document.add_paragraph()
                                    p_vispanelfind2.add_run(a_df.loc['vispnl_finds_ref',lang])
                                    p_vispanelfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_vispanelfind2.paragraph_format.space_after = Pt(10)

                    ############################################################
                    ## SUB2-SECTION - PANEL VISUAL INSPECTION RECORDED FAULTS ##
                    ############################################################

                        if vispanel_faults == True:
                            s_vispanelres = document.add_section()

                            s_vispanelres.orientation = WD_ORIENT.LANDSCAPE
                            s_vispanelres.page_width = 10058400
                            s_vispanelres.page_height = 7772400

                            h_vispanelres = document.add_heading('',level=3)
                            h_vispanelres_r = h_vispanelres.add_run(a_df.loc['vispnl_subsection2_fault',lang].format(vis_sec_no, vis_subsec_no))
                            h_vispanelres.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            h_vispanelres_r.font.name = default_font
                            h_vispanelres_r.font.size = Pt(14)
                            h_vispanelres_r.font.color.rgb = RGBColor(0, 0, 0)
                            h_vispanelres.paragraph_format.space_after = Pt(10)

                            # vispanelres_rec = [
                            #     ('HP-1', '12/10/2023', '5', 'cloudy', 'Is the panel mechanically damaged (e.g. loose parts, deformations)?', '' )
                            # ]

                            vispanelres_t = document.add_table(rows=1, cols=6,style='GridTable4-Accent3')
                            vispanelres_c = vispanelres_t.rows[0].cells
                            vispanelres_c[0].text = a_df.loc['vispnl_fault_t1',lang]
                            vispanelres_c[1].text = a_df.loc['vispnl_fault_t2',lang]
                            vispanelres_c[2].text = a_df.loc['vispnl_fault_t3',lang]
                            vispanelres_c[3].text = a_df.loc['vispnl_fault_t4',lang]
                            vispanelres_c[4].text = a_df.loc['vispnl_fault_t5',lang]
                            vispanelres_c[5].text = a_df.loc['vispnl_fault_t6',lang]

                            #for panid, tdate, otemp, weather, quest, remark in vispanelres_rec:
                            for row in fdf_pnl_vis_rep.itertuples():
                                vispanelres_c = vispanelres_t.add_row().cells
                                vispanelres_c[0].text = row[4]#panid
                                vispanelres_c[1].text = row[1]#tdate
                                vispanelres_c[2].text = row[2]#otemp
                                vispanelres_c[3].text = row[3]#weather
                                vispanelres_c[4].text = row[5]#quest
                                vispanelres_c[5].text = row[6]#remark
                            for cell in vispanelres_t.columns[0].cells:
                                cell.width = Cm(2)
                                if cell.text != a_df.loc['vispnl_fault_t1',lang]:
                                    cell.paragraphs[0].runs[0].font.bold = False
                            for cell in vispanelres_t.columns[1].cells:
                                cell.width = Cm(2)
                            for cell in vispanelres_t.columns[2].cells:
                                cell.width = Cm(2)
                            for cell in vispanelres_t.columns[3].cells:
                                cell.width = Cm(2)
                            for cell in vispanelres_t.columns[4].cells:
                                cell.width = Cm(9)
                            for cell in vispanelres_t.columns[5].cells:
                                cell.width = Cm(5)
                            
                        vis_subsec_no += 1

                    ###############################################
                    ## SUB- SECTION - CIRCUIT VISUAL INSPECTIONS ##
                    ###############################################

                    #if (audit_lvl_panel == 'Basic' or audit_lvl_panel == 'Standard'):

                    if audit_lvl_cct != 'None':
                        s_viscct = document.add_section()
                        s_viscct.orientation = WD_ORIENT.PORTRAIT
                        s_viscct.page_width = 7772400
                        s_viscct.page_height = 10058400
                        h_viscct = document.add_heading('',level=2)
                        h_viscct_r = h_viscct.add_run(a_df.loc['viscct_subsection',lang].format(vis_sec_no, vis_subsec_no))        
                        h_viscct.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_viscct_r.font.name = default_font
                        h_viscct_r.font.size = Pt(16)
                        h_viscct_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_viscct.paragraph_format.space_after = Pt(10)

                        p_viscct1 = document.add_paragraph()
                        p_viscct1.add_run(a_df.loc['viscct_p1',lang])
                        p_viscct1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_viscct1.paragraph_format.space_after = Pt(5)

                        p_viscct2 = document.add_paragraph()
                        p_viscct2.add_run(a_df.loc['viscct_audit',lang])
                        p_viscct2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_viscct2.paragraph_format.space_after = Pt(5)
                        
                        b1_viscct1 = document.add_paragraph(style='List Bullet')
                        b1_viscct1.add_run(a_df.loc['viscct_audit_b1',lang])
                        b1_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b2_viscct1 = document.add_paragraph(style='List Bullet')
                        b2_viscct1.add_run(a_df.loc['viscct_audit_b2',lang])
                        b2_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b3_viscct1 = document.add_paragraph(style='List Bullet')
                        b3_viscct1.add_run(a_df.loc['viscct_audit_b3',lang]) 
                        b3_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b3_viscct1 = document.add_paragraph(style='List Bullet')
                        b3_viscct1.add_run(a_df.loc['viscct_audit_b4',lang]) 
                        b3_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b4_viscct1 = document.add_paragraph(style='List Bullet')
                        b4_viscct1.add_run(a_df.loc['viscct_audit_b5',lang]) 
                        b4_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b5_viscct1 = document.add_paragraph(style='List Bullet')
                        b5_viscct1.add_run(a_df.loc['viscct_audit_b6',lang]) 
                        b5_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b6_viscct1 = document.add_paragraph(style='List Bullet')
                        b6_viscct1.add_run(a_df.loc['viscct_audit_b7',lang]) 
                        b6_viscct1.paragraph_format.left_indent = Inches(0.5)
                        b6_viscct1.paragraph_format.space_after = Pt(30)    

                    #####################################################
                    # SUB2-SECTION - CIRCUIT VISUAL INSPECTION FINDINGS #
                    #####################################################
                        
                        h_viscctfind = document.add_heading('',level=3)
                        h_viscctfind_r = h_viscctfind.add_run(a_df.loc['viscct_subsection2_find',lang].format(vis_sec_no, vis_subsec_no))
                        h_viscctfind.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_viscctfind_r.font.name = default_font
                        h_viscctfind_r.font.size = Pt(14)
                        h_viscctfind_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_viscctfind.paragraph_format.space_after = Pt(10)
                        
                        viscct_faults = False
                        
                        if viscct_majority == 'correct':
                            if viscctfail_gen == 'n':
                                p_viscctfind1 = document.add_paragraph()
                                p_viscctfind1.add_run(a_df.loc['viscct_finds_corrn',lang])
                                p_viscctfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_viscctfind1.paragraph_format.space_after = Pt(10)
                            if viscctfail_gen == 'y':
                                p_viscctfind1 = document.add_paragraph()
                                p_viscctfind1.add_run(a_df.loc['viscct_finds_corry',lang])
                                p_viscctfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_viscctfind1.paragraph_format.space_after = Pt(10)

                                viscct_faults = True
                                p_viscctfind2 = document.add_paragraph()
                                p_viscctfind2.add_run(a_df.loc['viscct_finds_ref',lang])
                                p_viscctfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_viscctfind2.paragraph_format.space_after = Pt(10)    
                                
                        if viscct_majority == 'incorrect':
                            p_viscctfind1 = document.add_paragraph()
                            p_viscctfind1.add_run(a_df.loc['viscct_finds_incorr',lang])
                            p_viscctfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_viscctfind1.paragraph_format.space_after = Pt(10)

                            viscct_faults = True
                            p_viscctfind2 = document.add_paragraph()
                            p_viscctfind2.add_run(a_df.loc['viscct_finds_ref',lang])
                            p_viscctfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_viscctfind2.paragraph_format.space_after = Pt(10)

                    ###################################################################
                    ## SUB2-SECTION -  EHT CIRCUIT VISUAL INSPECTION RECORDED FAULTS ##
                    ###################################################################

                        if viscct_faults == True:
                            s_viscctres = document.add_section()

                            s_viscctres.orientation = WD_ORIENT.LANDSCAPE
                            s_viscctres.page_width = 10058400
                            s_viscctres.page_height = 7772400

                            h_viscctres = document.add_heading('',level=3)
                            h_viscctres_r = h_viscctres.add_run(a_df.loc['viscct_subsection2_fault',lang].format(vis_sec_no, vis_subsec_no))
                            h_viscctres.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            h_viscctres_r.font.name = default_font
                            h_viscctres_r.font.size = Pt(14)
                            h_viscctres_r.font.color.rgb = RGBColor(0, 0, 0)
                            h_viscctres.paragraph_format.space_after = Pt(10)

                            # viscctres_rec = [
                            #     ('10001', '12/10/2023', '5', 'cloudy', 
                            #     'Is Field Temperature Sensor Junction Box lid sealing in acceptable condition?', '' )
                            # ]

                            viscctres_t = document.add_table(rows=1, cols=6,style='GridTable4-Accent3')
                            viscctres_c = viscctres_t.rows[0].cells
                            viscctres_c[0].text = a_df.loc['viscct_fault_t1',lang]
                            viscctres_c[1].text = a_df.loc['viscct_fault_t2',lang]
                            viscctres_c[2].text = a_df.loc['viscct_fault_t3',lang]
                            viscctres_c[3].text = a_df.loc['viscct_fault_t4',lang]
                            viscctres_c[4].text = a_df.loc['viscct_fault_t5',lang]
                            viscctres_c[5].text = a_df.loc['viscct_fault_t6',lang]

                            #for cctid, tdate, otemp, weather, quest, remark in fdf_cct_vis_rep:
                            #for tdate, otemp, weather, cctid, quest, remark in fdf_cct_vis_rep:
                            for row in fdf_cct_vis_rep.itertuples():
                                viscctres_c = viscctres_t.add_row().cells
                                viscctres_c[0].text = row[4]
                                viscctres_c[1].text = row[1]
                                viscctres_c[2].text = row[2]
                                viscctres_c[3].text = row[3]
                                viscctres_c[4].text = row[5]
                                viscctres_c[5].text = row[6]
                            for cell in viscctres_t.columns[0].cells:
                                cell.width = Cm(2)
                                if cell.text != a_df.loc['viscct_fault_t1',lang]:
                                    cell.paragraphs[0].runs[0].font.bold = False
                            for cell in viscctres_t.columns[1].cells:
                                cell.width = Cm(2.5)
                            for cell in viscctres_t.columns[2].cells:
                                cell.width = Cm(2)
                            for cell in viscctres_t.columns[3].cells:
                                cell.width = Cm(2)
                            for cell in viscctres_t.columns[4].cells:
                                cell.width = Cm(9)
                            for cell in viscctres_t.columns[5].cells:
                                cell.width = Cm(4.5)
                            
                            
                        vis_subsec_no += 1
                        
                    ############################################################
                    ## SUB- SECTION - CONTROLLER EXAMINATION ##
                    ############################################################
                        
                        if (ctrl_loc == 'field' or ctrl_loc == 'both'):  
                            s_ctrl = document.add_section()
                            s_ctrl.orientation = WD_ORIENT.PORTRAIT
                            s_ctrl.page_width = 7772400
                            s_ctrl.page_height = 10058400
                            h_ctrl = document.add_heading('',level=2)
                            h_ctrl_r = h_ctrl.add_run(a_df.loc['visctrl_subsection',lang].format(vis_sec_no, vis_subsec_no))        
                            h_ctrl.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            h_ctrl_r.font.name = default_font
                            h_ctrl_r.font.size = Pt(16)
                            h_ctrl_r.font.color.rgb = RGBColor(0, 0, 0)
                            h_ctrl.paragraph_format.space_after = Pt(10)        
                            
                            p_ctrl1 = document.add_paragraph()
                            p_ctrl1.add_run(a_df.loc['visctrl_p1',lang])
                            p_ctrl1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_ctrl1.paragraph_format.space_after = Pt(5)    
                        
                            p_ctrl2 = document.add_paragraph()
                            p_ctrl2.add_run(a_df.loc['visctrl_audit',lang])
                            p_ctrl2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_ctrl2.paragraph_format.space_after = Pt(5)

                            b1_ctrl2 = document.add_paragraph(style='List Bullet')
                            b1_ctrl2.add_run(a_df.loc['visctrl_audit_b1',lang])
                            b1_ctrl2.paragraph_format.left_indent = Inches(0.5)
                            b2_ctrl2 = document.add_paragraph(style='List Bullet')
                            b2_ctrl2.add_run(a_df.loc['visctrl_audit_b2',lang])
                            b2_ctrl2.paragraph_format.left_indent = Inches(0.5)
                            b3_ctrl2 = document.add_paragraph(style='List Bullet')
                            b3_ctrl2.add_run(a_df.loc['visctrl_audit_b3',lang])
                            b3_ctrl2.paragraph_format.left_indent = Inches(0.5)
                            b3_ctrl2.paragraph_format.space_after = Pt(30)   
                        
                    ########################################################
                    # SUB2-SECTION - CONTROLLER EXAMINATION FINDINGS #
                    ########################################################
                            
                            h_ctrlfind = document.add_heading('',level=3)
                            h_ctrlfind_r = h_ctrlfind.add_run(a_df.loc['visctrl_subsection2_find',lang].format(vis_sec_no, vis_subsec_no))
                            h_ctrlfind.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            h_ctrlfind_r.font.name = default_font
                            h_ctrlfind_r.font.size = Pt(14)
                            h_ctrlfind_r.font.color.rgb = RGBColor(0, 0, 0)
                            h_ctrlfind.paragraph_format.space_after = Pt(10)

                            ctrl_faults = False

                            if ctrl_majority == 'correct':

                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'n' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corrnnn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'n' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corrynn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'y' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corrnyn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'n' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corrnny',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'y' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corryyn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'n' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corryny',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'y' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corrnyy',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'y' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_corryyy',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                    ctrl_faults = True
                                    p_ctrlfind2 = document.add_paragraph()
                                    p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                    p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind2.paragraph_format.space_after = Pt(10)
                    
                            if ctrl_majority == 'incorrect':

                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'n' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorrynn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'y' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorrnyn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)    

                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'n' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorrnny',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'y' and ctrlfail_sens == 'n':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorryyn',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'n' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorryny',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                if ctrlfail_temp == 'n' and ctrlfail_pw == 'y' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorrnyy',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)

                                if ctrlfail_temp == 'y' and ctrlfail_pw == 'y' and ctrlfail_sens == 'y':
                                    p_ctrlfind1 = document.add_paragraph()
                                    p_ctrlfind1.add_run(a_df.loc['visctrl_finds_incorryyy',lang])
                                    p_ctrlfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_ctrlfind1.paragraph_format.space_after = Pt(10)
                                    
                                ctrl_faults = True
                                p_ctrlfind2 = document.add_paragraph()
                                p_ctrlfind2.add_run(a_df.loc['visctrl_finds_ref',lang])
                                p_ctrlfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_ctrlfind2.paragraph_format.space_after = Pt(10)
                                
                    ########################################################
                    # SUB2-SECTION - CONTROLLER EXAMINATION RESULTS #
                    ########################################################
                        
                        
                            if ctrl_faults == True:
                                s_ctrlres = document.add_section()

                                s_ctrlres.orientation = WD_ORIENT.LANDSCAPE
                                s_ctrlres.page_width = 10058400
                                s_ctrlres.page_height = 7772400

                                h_ctrlres = document.add_heading('',level=3)
                                h_ctrlres_r = h_ctrlres.add_run(a_df.loc['visctrl_subsection2_fault',lang].format(vis_sec_no, vis_subsec_no))
                                h_ctrlres.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                h_ctrlres_r.font.name = default_font
                                h_ctrlres_r.font.size = Pt(14)
                                h_ctrlres_r.font.color.rgb = RGBColor(0, 0, 0)
                                h_ctrlres.paragraph_format.space_after = Pt(10)

                                # ctrlres_rec = [
                                #     ('10001', '12/10/2023', '5', 'cloudy', 
                                #     'Does Limiter Temperature Setpoint match controlling document?', '' )
                                # ]

                                ctrlres_t = document.add_table(rows=1, cols=6,style='GridTable4-Accent3')
                                ctrlres_c = ctrlres_t.rows[0].cells
                                ctrlres_c[0].text = a_df.loc['visctrl_fault_t1',lang]
                                ctrlres_c[1].text = a_df.loc['visctrl_fault_t2',lang]
                                ctrlres_c[2].text = a_df.loc['visctrl_fault_t3',lang]
                                ctrlres_c[3].text = a_df.loc['visctrl_fault_t4',lang]
                                ctrlres_c[4].text = a_df.loc['visctrl_fault_t5',lang]
                                ctrlres_c[5].text = a_df.loc['visctrl_fault_t6',lang]

                                #for cctid, tdate, otemp, weather, quest, remark in ctrlres_rec:
                                for row in fdf_cct_th_rep.itertuples():
                                    ctrlres_c = ctrlres_t.add_row().cells
                                    ctrlres_c[0].text = row[4]
                                    ctrlres_c[1].text = row[1]
                                    ctrlres_c[2].text = row[2]
                                    ctrlres_c[3].text = row[3]
                                    ctrlres_c[4].text = row[5]
                                    ctrlres_c[5].text = row[6]
                                for cell in ctrlres_t.columns[0].cells:
                                    cell.width = Cm(2)
                                    if cell.text != a_df.loc['visctrl_fault_t1',lang]:
                                        cell.paragraphs[0].runs[0].font.bold = False
                                for cell in ctrlres_t.columns[1].cells:
                                    cell.width = Cm(2.5)
                                for cell in ctrlres_t.columns[2].cells:
                                    cell.width = Cm(2)
                                for cell in ctrlres_t.columns[3].cells:
                                    cell.width = Cm(2)
                                for cell in ctrlres_t.columns[4].cells:
                                    cell.width = Cm(9)
                                for cell in ctrlres_t.columns[5].cells:
                                    cell.width = Cm(4.5)
                            
                            
                                vis_subsec_no += 1    

                            
                            
                    ##################################################
                    ## SUB- SECTION - INSULATION VISUAL INSPECTIONS ##
                    ##################################################

                    if audit_lvl_insul != 'None':
                        s_visinsul = document.add_section()
                        s_visinsul.orientation = WD_ORIENT.PORTRAIT
                        s_visinsul.page_width = 7772400
                        s_visinsul.page_height = 10058400
                        h_visinsul = document.add_heading('',level=2)
                        if audit_lvl_insul == 'Basic':
                            h_visinsul_r = h_visinsul.add_run(a_df.loc['visins_subsection_s',lang].format(vis_sec_no, vis_subsec_no))
                        else:
                            h_visinsul_r = h_visinsul.add_run(a_df.loc['visins_subsection_ir',lang].format(vis_sec_no, vis_subsec_no))        
                        h_visinsul.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_visinsul_r.font.name = default_font
                        h_visinsul_r.font.size = Pt(16)
                        h_visinsul_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_visinsul.paragraph_format.space_after = Pt(10)

                        
                        p_visinsul1 = document.add_paragraph()
                        p_visinsul1.add_run(a_df.loc['visins_p1',lang])
                        p_visinsul1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_visinsul1.paragraph_format.space_after = Pt(5)

                        p_visinsul2 = document.add_paragraph()
                        p_visinsul2.add_run(a_df.loc['visins_audit',lang])
                        p_visinsul2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_visinsul2.paragraph_format.space_after = Pt(5)
                        
                        b1_visinsul1 = document.add_paragraph(style='List Bullet')
                        b1_visinsul1.add_run(a_df.loc['visins_audit_b1',lang])
                        b1_visinsul1.paragraph_format.left_indent = Inches(0.5)
                        b2_visinsul1 = document.add_paragraph(style='List Bullet')
                        b2_visinsul1.add_run(a_df.loc['visins_audit_b2',lang])
                        b2_visinsul1.paragraph_format.left_indent = Inches(0.5)
                        b3_visinsul1 = document.add_paragraph(style='List Bullet')
                        b3_visinsul1.add_run(a_df.loc['visins_audit_b3',lang])
                        b3_visinsul1.paragraph_format.left_indent = Inches(0.5)
                        if (audit_lvl_insul == 'Standard' or audit_lvl_insul == 'Advanced'):
                            b4_visinsul1 = document.add_paragraph(style='List Bullet')
                            b4_visinsul1.add_run(a_df.loc['visins_audit_b4',lang])
                            b4_visinsul1.paragraph_format.left_indent = Inches(0.5)   
                            b4_visinsul1.paragraph_format.space_after = Pt(30)
                        else:
                            b3_visinsul1.paragraph_format.space_after = Pt(30)
                        
                    ########################################################
                    # SUB2-SECTION - INSULATION VISUAL INSPECTION FINDINGS #
                    ########################################################
                        
                        h_visinsulfind = document.add_heading('',level=3)
                        h_visinsulfind_r = h_visinsulfind.add_run(a_df.loc['visins_subsection2_find',lang].format(vis_sec_no, vis_subsec_no))
                        h_visinsulfind.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_visinsulfind_r.font.name = default_font
                        h_visinsulfind_r.font.size = Pt(14)
                        h_visinsulfind_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_visinsulfind.paragraph_format.space_after = Pt(10)
                        
                        visinsul_faults = False
                       
                        if audit_lvl_insul == 'Basic':
                            if visinsul_majority == 'correct':
                                if visinsulfail_gen == 'n':
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_finds_corrn',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                if visinsulfail_gen == 'y':
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_finds_corry',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)

                                    visinsul_faults = True
                                    p_visinsulfind2 = document.add_paragraph()
                                    p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                    p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind2.paragraph_format.space_after = Pt(10)
                            if visinsul_majority == 'incorrect':
                                p_visinsulfind1 = document.add_paragraph()
                                p_visinsulfind1.add_run(a_df.loc['visins_finds_incorr',lang])
                                p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_visinsulfind1.paragraph_format.space_after = Pt(10)

                                visinsul_faults = True
                                p_visinsulfind2 = document.add_paragraph()
                                p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_visinsulfind2.paragraph_format.space_after = Pt(10)
                        
                        if (audit_lvl_insul == 'Standard' or audit_lvl_insul == 'Advanced'):
                            if visinsul_majority == 'correct':
                                if (visinsulfail_gen == 'n' and visinsulfail_ifr == 'n'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_corrnn',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                    
                                if (visinsulfail_gen == 'y' and visinsulfail_ifr == 'n'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_corryn',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                    
                                    visinsul_faults = True
                                    p_visinsulfind2 = document.add_paragraph()
                                    p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                    p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (visinsulfail_gen == 'n' and visinsulfail_ifr == 'y'):
                                    p_visinsulfind1 = document.add_paragraph()                
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_corrny',lang])         
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                    
                                    visinsul_faults = True
                                    p_visinsulfind2 = document.add_paragraph()
                                    p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                    p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (visinsulfail_gen == 'y' and visinsulfail_ifr == 'y'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_corryy',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)

                                    visinsul_faults = True
                                    p_visinsulfind2 = document.add_paragraph()
                                    p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                    p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind2.paragraph_format.space_after = Pt(10)
                                    
                            if visinsul_majority == 'incorrect':
                                if (visinsulfail_gen == 'y' and visinsulfail_ifr == 'n'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_incorryn',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                
                                if (visinsulfail_gen == 'n' and visinsulfail_ifr == 'y'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_incorrny',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)
                                    
                                if (visinsulfail_gen == 'y' and visinsulfail_ifr == 'y'):
                                    p_visinsulfind1 = document.add_paragraph()
                                    p_visinsulfind1.add_run(a_df.loc['visins_findir_incorryy',lang])
                                    p_visinsulfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_visinsulfind1.paragraph_format.space_after = Pt(10)

                                visinsul_faults = True
                                p_visinsulfind2 = document.add_paragraph()
                                p_visinsulfind2.add_run(a_df.loc['visins_finds_ref',lang])
                                p_visinsulfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_visinsulfind2.paragraph_format.space_after = Pt(10)

                    ######################################################################
                    ## SUB2-SECTION -  EHT INSULATION VISUAL INSPECTION RECORDED FAULTS ##
                    ######################################################################

                        if visinsul_faults == True:
                            s_visinsulres = document.add_section()

                            s_visinsulres.orientation = WD_ORIENT.LANDSCAPE
                            s_visinsulres.page_width = 10058400
                            s_visinsulres.page_height = 7772400

                            h_visinsulres = document.add_heading('',level=3)
                            h_visinsulres_r = h_visinsulres.add_run(a_df.loc['visins_subsection2_fault',lang].format(vis_sec_no, vis_subsec_no))
                            h_visinsulres.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            h_visinsulres_r.font.name = default_font
                            h_visinsulres_r.font.size = Pt(14)
                            h_visinsulres_r.font.color.rgb = RGBColor(0, 0, 0)
                            h_visinsulres.paragraph_format.space_after = Pt(10)

                            # visinsulres_rec = [
                            #     ('1', '10001', '12/10/2023', '5', 'cloudy', 
                            #     'Missing Insulation and Cladding and/or Blankets', '' )
                            # ]

                            visinsulres_t = document.add_table(rows=1, cols=7,style='GridTable4-Accent3')
                            visinsulres_c = visinsulres_t.rows[0].cells
                            visinsulres_c[0].text = a_df.loc['visins_fault_t1',lang]                          
                            visinsulres_c[1].text = a_df.loc['visins_fault_t2',lang]
                            visinsulres_c[2].text = a_df.loc['visins_fault_t3',lang]
                            visinsulres_c[3].text = a_df.loc['visins_fault_t4',lang]
                            visinsulres_c[4].text = a_df.loc['visins_fault_t5',lang]
                            visinsulres_c[5].text = a_df.loc['visins_fault_t6',lang]
                            visinsulres_c[6].text = a_df.loc['visins_fault_t7',lang]

                            fdf_ins_vis_rep['Section'] = fdf_ins_vis_rep['Section'].astype(str)

                            #for sec, cctid, tdate, otemp, weather, quest, remark in visinsulres_rec:
                            for row in fdf_ins_vis_rep.itertuples():
                                visinsulres_c = visinsulres_t.add_row().cells
                                visinsulres_c[0].text = row[1]#cctid
                                visinsulres_c[1].text = row[2]#sec
                                visinsulres_c[2].text = row[3]#tdate
                                visinsulres_c[3].text = row[4]#otemp
                                visinsulres_c[4].text = row[5]#weather
                                visinsulres_c[5].text = row[6]#quest
                                visinsulres_c[6].text = row[7]#remark
                            for cell in visinsulres_t.columns[0].cells:
                                cell.width = Cm(2.3)
                                if cell.text != a_df.loc['visins_fault_t1',lang]:
                                    cell.paragraphs[0].runs[0].font.bold = False
                            for cell in visinsulres_t.columns[1].cells:
                                cell.width = Cm(2)
                            for cell in visinsulres_t.columns[2].cells:
                                cell.width = Cm(2.5)
                            for cell in visinsulres_t.columns[3].cells:
                                cell.width = Cm(2.2)
                            for cell in visinsulres_t.columns[4].cells:
                                cell.width = Cm(2)
                            for cell in visinsulres_t.columns[5].cells:
                                cell.width = Cm(8.5)
                            for cell in visinsulres_t.columns[6].cells:
                                cell.width = Cm(5)
                                
                                
                        vis_subsec_no += 1
                        
                        
                        
                    ##############################
                    ### SECTION - MEASUREMENTS ###
                    ##############################

                    if (audit_lvl_panel != 'None' or
                        audit_lvl_cct != 'None' or
                        audit_lvl_insul == 'Advanced'
                        ):

                        mes_sec_no = vis_sec_no + 1

                        section_measure = document.add_section()
                        section_measure.orientation = WD_ORIENT.PORTRAIT
                        section_measure.page_width = 7772400
                        section_measure.page_height = 10058400
                        measure_heading = document.add_heading('',level=1)
                        measure_heading_run = measure_heading.add_run(a_df.loc['mes_section',lang].format(mes_sec_no))
                        measure_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        measure_heading_run.font.name = default_font
                        measure_heading_run.font.size = Pt(18)
                        measure_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        measure_heading.paragraph_format.space_after = Pt(10)

                        paragh_measure1 = document.add_paragraph()
                        paragh_measure1.add_run(a_df.loc['mes_p1',lang])
                        paragh_measure1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_measure1.paragraph_format.space_after = Pt(10)

                    # if (audit_lvl_panel == 'Standard' or
                    #     audit_lvl_panel == 'Advanced' or
                    #     audit_lvl_cct != 'None' or
                    #     audit_lvl_insul == 'Advanced'
                    #     ):
                        paragh_measure2 = document.add_paragraph()
                        paragh_measure2.add_run(a_df.loc['mes_audit',lang])
                        paragh_measure2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_measure2.paragraph_format.space_after = Pt(5)

                        if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                            paragh_measurepanel_rcd = document.add_paragraph(a_df.loc['mes_audit_rcda',lang],style='List Bullet')
                            paragh_measurepanel_rcd.paragraph_format.left_indent = Inches(0.5)
                        #     paragh_measurepanel_temp = document.add_paragraph('Temperature of panel components',style='List Bullet')
                        #     paragh_measurepanel_temp.paragraph_format.left_indent = Inches(0.5)
                        if audit_lvl_cct != 'None':
                            paragh_measurecct_ir = document.add_paragraph(a_df.loc['mes_audit_cctir',lang],style='List Bullet')
                            paragh_measurecct_ir.paragraph_format.left_indent = Inches(0.5)  
                            paragh_measurecct_ohm = document.add_paragraph(a_df.loc['mes_audit_cctcont',lang],style='List Bullet')
                            paragh_measurecct_ohm.paragraph_format.left_indent = Inches(0.5)
                            paragh_measurecct_va = document.add_paragraph(a_df.loc['mes_audit_cctvolt',lang],style='List Bullet')
                            paragh_measurecct_va.paragraph_format.left_indent = Inches(0.5)
                            paragh_measurecct_ia = document.add_paragraph(a_df.loc['mes_audit_cctcurr',lang],style='List Bullet')
                            paragh_measurecct_ia.paragraph_format.left_indent = Inches(0.5)
                        #     paragh_measurecct_temp = document.add_paragraph('Temperature along the EHT circuit route',style='List Bullet')
                        #     paragh_measurecct_temp.paragraph_format.left_indent = Inches(0.5)    
                        if audit_lvl_insul == 'Advanced':
                            paragh_measureinsul_thick = document.add_paragraph(a_df.loc['mes_audit_inspar',lang],style='List Bullet')
                            paragh_measureinsul_thick.paragraph_format.left_indent = Inches(0.5)  

                        if (audit_lvl_panel == 'Standard' or
                            audit_lvl_panel == 'Advanced' or
                            audit_lvl_cct != 'None' or
                            audit_lvl_insul == 'Advanced'
                            ):
                            paragh_measure3 = document.add_paragraph()
                            paragh_measure3.add_run(a_df.loc['mes_scope',lang])
                            if (audit_lvl_cct != 'None' or audit_lvl_insul != "None"):
                                if audit_lvl_panel == 'None':
                                    paragh_measure3.add_run(a_df.loc['mes_scope_1a',lang])
                                elif audit_lvl_panel != 'None':
                                    paragh_measure3.add_run(a_df.loc['mes_scope_1b',lang])
                            else:
                                paragh_measure3.add_run(a_df.loc['mes_scope_1c',lang])
                            paragh_measure3.add_run(a_df.loc['mes_scope_2',lang])
                            paragh_measure3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measure3.paragraph_format.space_after = Pt(10)

                        if audit_lvl_panel == 'Basic':
                            paragh_measurepan = document.add_paragraph()
                            paragh_measurepan.add_run(a_df.loc['mes_pnlb',lang])
                            paragh_measurepan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurepan.paragraph_format.space_after = Pt(5) 

                        if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                            paragh_measurepan = document.add_paragraph()
                            paragh_measurepan.add_run(a_df.loc['mes_pnlsa',lang])
                            paragh_measurepan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurepan.paragraph_format.space_after = Pt(5) 
                            
                        if audit_lvl_cct != 'None':

                            paragh_measurecctir = document.add_paragraph()
                            paragh_measurecctir.add_run(a_df.loc['mes_cctir',lang])
                            paragh_measurecctir.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctir.paragraph_format.space_after = Pt(5) 

                            paragh_measurecctcont = document.add_paragraph()
                            paragh_measurecctcont.add_run(a_df.loc['mes_cctcont',lang])
                            paragh_measurecctcont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctcont.paragraph_format.space_after = Pt(5)

                            paragh_measurecctvolt = document.add_paragraph()
                            paragh_measurecctvolt.add_run(a_df.loc['mes_cctvolt',lang])
                            paragh_measurecctvolt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctvolt.paragraph_format.space_after = Pt(5)

                            paragh_measurecctcurr = document.add_paragraph()
                            paragh_measurecctcurr.add_run(a_df.loc['mes_cctcurr',lang])
                            paragh_measurecctcurr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctcurr.paragraph_format.space_after = Pt(5)
                            
                            paragh_measurecctth = document.add_paragraph()
                            paragh_measurecctth.add_run(a_df.loc['mes_cctth',lang])
                            paragh_measurecctth.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctth.paragraph_format.space_after = Pt(5)

                        if audit_lvl_insul == 'Advanced':
                            paragh_measureins = document.add_paragraph()
                            paragh_measureins.add_run(a_df.loc['mes_insparam',lang])
                            paragh_measureins.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measureins.paragraph_format.space_after = Pt(5)
                        
                        if audit_lvl_cct != 'None':
                            if audit_lvl_panel == 'None':
                                paragh_measurecctfin = document.add_paragraph()
                                paragh_measurecctfin_run1 = paragh_measurecctfin.add_run(a_df.loc['mes_risk_c1',lang])
                                paragh_measurecctfin_run1.font.underline = True
                                paragh_measurecctfin.add_run(a_df.loc['mes_risk_c2',lang])
                                paragh_measurecctfin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_measurecctfin.paragraph_format.space_after = Pt(10)
                            else:
                                paragh_measurecctfin = document.add_paragraph()
                                paragh_measurecctfin_run1 = paragh_measurecctfin.add_run(a_df.loc['mes_risk_cp1',lang])
                                paragh_measurecctfin_run1.font.underline = True
                                paragh_measurecctfin.add_run(a_df.loc['mes_risk_cp2',lang])
                                paragh_measurecctfin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_measurecctfin.paragraph_format.space_after = Pt(10)
                        else:
                            if audit_lvl_panel != 'None':
                                paragh_measurecctfin = document.add_paragraph()
                                paragh_measurecctfin_run1 = paragh_measurecctfin.add_run(a_df.loc['mes_risk_p1',lang])
                                paragh_measurecctfin_run1.font.underline = True
                                paragh_measurecctfin.add_run(a_df.loc['mes_risk_p2',lang])
                                paragh_measurecctfin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_measurecctfin.paragraph_format.space_after = Pt(10)
                            
                        if audit_lvl_insul == 'Advanced':
                            paragh_measurecctfin2 = document.add_paragraph()
                            paragh_measurecctfin2.add_run(a_df.loc['mes_risk_i',lang])
                            paragh_measurecctfin2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_measurecctfin2.paragraph_format.space_after = Pt(10)

                        mes_subsec_no = 1

                        if audit_lvl_panel == 'Basic':
                            
                            section_rcdtest = document.add_section()
                            
                            rcdtest_heading = document.add_heading('',level=2)
                            rcdtest_heading_run = rcdtest_heading.add_run(a_df.loc['mespnlb_subsection',lang].format(mes_sec_no, mes_subsec_no))
                            rcdtest_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdtest_heading_run.font.name = default_font
                            rcdtest_heading_run.font.size = Pt(16)
                            rcdtest_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                            rcdtest_heading.paragraph_format.space_after = Pt(10)

                            p_rcdtest1 = document.add_paragraph()
                            p_rcdtest1.add_run(a_df.loc['mespnlb_p1',lang])
                            p_rcdtest1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_rcdtest1.paragraph_format.space_after = Pt(30)
                            
                            ## SUB2-SECTION - RCD PUSH BUTTON TEST FINDINGS ##
                            
                            rcdfind_heading = document.add_heading('',level=3)
                            rcdfind_heading_run = rcdfind_heading.add_run(a_df.loc['mespnlb_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                            rcdfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdfind_heading_run.font.name = default_font
                            rcdfind_heading_run.font.size = Pt(14)
                            rcdfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                            rcdfind_heading.paragraph_format.space_after = Pt(10)
                            
                            if rcdbut_test_majority == 'correct':
                                if rcdbut_test_fail == 'n':
                                    p_rcdbutfind1 = document.add_paragraph()
                                    p_rcdbutfind1.add_run(a_df.loc['mespnlb_corrn1',lang])
                                    p_rcdbutfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    
                                    p_rcdbutfind2 = document.add_paragraph()
                                    p_rcdbutfind2.add_run(a_df.loc['mespnlb_corrn2',lang])
                                    p_rcdbutfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdbutfind2.paragraph_format.space_after = Pt(10)
                                    
                                if rcdbut_test_fail == 'y':
                                    p_rcdbutfind1 = document.add_paragraph()
                                    p_rcdbutfind1.add_run(a_df.loc['mespnlb_corry',lang])
                                    p_rcdbutfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdbutfind1.paragraph_format.space_after = Pt(10)
                                    
                            if rcdbut_test_majority == 'incorrect':
                                    p_rcdbutfind1 = document.add_paragraph()
                                    p_rcdbutfind1.add_run(a_df.loc['mespnlb_incorr',lang])
                                    p_rcdbutfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdbutfind1.paragraph_format.space_after = Pt(10)      

                                    
                            p_rcdbutfindref = document.add_paragraph()
                            p_rcdbutfindref.add_run(a_df.loc['mespnlb_ref',lang])
                            p_rcdbutfindref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_rcdbutfindref.paragraph_format.space_after = Pt(10)
                                    
                            ## SUB2-SECTION - RCD PUSH BUTTON TEST RESULTS ##

                            s_rcdbutres = document.add_section()

                            s_rcdbutres.orientation = WD_ORIENT.LANDSCAPE
                            s_rcdbutres.page_width = 10058400
                            s_rcdbutres.page_height = 7772400

                            rcdbutres_head = document.add_heading('',level=3)
                            rcdbutres_head_r = rcdbutres_head.add_run(a_df.loc['mespnlb_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                            rcdbutres_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdbutres_head_r.font.name = default_font
                            rcdbutres_head_r.font.size = Pt(14)
                            rcdbutres_head_r.font.color.rgb = RGBColor(0, 0, 0)
                            rcdbutres_head.paragraph_format.space_after = Pt(10)

                            # rcdbutres_rec = [
                            #     ('HP-1', "Q1", '12/10/2023', '5', 'cloudy', 'Pass', '' )
                            # ]

                            rcdbutres_t = document.add_table(rows=1, cols=7,style='GridTable4-Accent3')
                            rcdbutres_c = rcdbutres_t.rows[0].cells
                            rcdbutres_c[0].text = a_df.loc['mespnlb_res_t1',lang]
                            rcdbutres_c[1].text = a_df.loc['mespnlb_res_t2',lang]
                            rcdbutres_c[2].text = a_df.loc['mespnlb_res_t3',lang]
                            rcdbutres_c[3].text = a_df.loc['mespnlb_res_t4',lang]
                            rcdbutres_c[4].text = a_df.loc['mespnlb_res_t5',lang]
                            rcdbutres_c[5].text = a_df.loc['mespnlb_res_t6',lang]
                            rcdbutres_c[6].text = a_df.loc['mespnlb_res_t7',lang]


                            #for panid, rcdid, tdate, otemp, weather, stat, remark in rcdbutres_rec:
                            for row in fdf_rcdbut_rep.itertuples():
                                rcdbutres_c = rcdbutres_t.add_row().cells
                                rcdbutres_c[0].text = row[1]#panid
                                rcdbutres_c[1].text = row[2]#rcdid
                                rcdbutres_c[2].text = row[3]#tdate
                                rcdbutres_c[3].text = row[4]#otemp
                                rcdbutres_c[4].text = row[5]#weather
                                rcdbutres_c[5].text = row[6]#stat
                                rcdbutres_c[6].text = row[7]#remark
                            for cell in rcdbutres_t.columns[0].cells:
                                cell.width = Cm(2)
                                if cell.text != a_df.loc['mespnlb_res_t1',lang]:
                                    cell.paragraphs[0].runs[0].font.bold = False
                            for cell in rcdbutres_t.columns[1].cells:
                                cell.width = Cm(3)
                            for cell in rcdbutres_t.columns[2].cells:
                                cell.width = Cm(3)
                            for cell in rcdbutres_t.columns[3].cells:
                                cell.width = Cm(3)
                            for cell in rcdbutres_t.columns[4].cells:
                                cell.width = Cm(3)
                            for cell in rcdbutres_t.columns[5].cells:
                                cell.width = Cm(3)
                            for cell in rcdbutres_t.columns[6].cells:
                                cell.width = Cm(4)
                            rcdbutres_t_space = document.add_paragraph()
                            #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                            #rcdbutres_t_space.paragraph_format.space_after = Pt(30)
                                    
                                    
                            mes_subsec_no += 1

                        ##########################################################################################################
                        ##########################################################################################################
                        ##########################################################################################################


                        if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                            
                            section_rcdmes = document.add_section()
                            
                            rcdmes_heading = document.add_heading('',level=2)
                            rcdmes_heading_run = rcdmes_heading.add_run(a_df.loc['mespnlsa_subsection',lang].format(mes_sec_no, mes_subsec_no))
                            rcdmes_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdmes_heading_run.font.name = default_font
                            rcdmes_heading_run.font.size = Pt(16)
                            rcdmes_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                            rcdmes_heading.paragraph_format.space_after = Pt(10)  
                            
                            p_rcdmes1 = document.add_paragraph()
                            p_rcdmes1.add_run(a_df.loc['mespnlsa_p1',lang])
                            p_rcdmes1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
                            p_rcdmes2 = document.add_paragraph()
                            p_rcdmes2.add_run(a_df.loc['mespnlsa_p2',lang])
                            p_rcdmes2.paragraph_format.space_after = Pt(10)    
                            

                            ## SUB2-SECTION - RCD MEASUREMENTS TEST FINDINGS ##
                            
                            rcdfind_heading = document.add_heading('',level=3)
                            rcdfind_heading_run = rcdfind_heading.add_run(a_df.loc['mespnlsa_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                            rcdfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdfind_heading_run.font.name = default_font
                            rcdfind_heading_run.font.size = Pt(14)
                            rcdfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                            rcdfind_heading.paragraph_format.space_after = Pt(10)    

                            
                            if (rcdbut_test_majority == 'correct' and rcdmes_test_majority == 'correct'):
                                if (rcdbut_test_fail == 'n' and rcdmes_test_fail == 'n'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_corrnn',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (rcdbut_test_fail == 'y' and rcdmes_test_fail == 'n'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_corryn',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (rcdbut_test_fail == 'n' and rcdmes_test_fail == 'y'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_corrny',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (rcdbut_test_fail == 'y' and rcdmes_test_fail == 'y'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_corryy',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)

                            if (rcdbut_test_majority == 'incorrect' or rcdmes_test_majority == 'incorrect'):
                                if (rcdbut_test_fail == 'y' and rcdmes_test_fail == 'n'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_incorryn',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)
                                    
                                if (rcdbut_test_fail == 'n' and rcdmes_test_fail == 'y'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_incorrny',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)   

                                if (rcdbut_test_fail == 'y' and rcdmes_test_fail == 'y'):
                                    p_rcdmesfind1 = document.add_paragraph()
                                    p_rcdmesfind1.add_run(a_df.loc['mespnlsa_incorryy',lang])
                                    p_rcdmesfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    #p_rcdbutfind1.paragraph_format.space_after = Pt(1)
                                    
                                    p_rcdmesfind2 = document.add_paragraph()
                                    p_rcdmesfind2.add_run(a_df.loc['mespnlsa_advise',lang])
                                    p_rcdmesfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_rcdmesfind2.paragraph_format.space_after = Pt(10)
                                    
                            p_rcdmesfindref = document.add_paragraph()
                            p_rcdmesfindref.add_run(a_df.loc['mespnlsa_ref',lang])
                            p_rcdmesfindref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_rcdmesfindref.paragraph_format.space_after = Pt(10)

                            ## SUB2-SECTION - RCD PUSH BUTTON TEST & MEASUREMENTS RESULTS ##

                            s_rcdmesres = document.add_section()

                            s_rcdmesres.orientation = WD_ORIENT.LANDSCAPE
                            s_rcdmesres.page_width = 10058400
                            s_rcdmesres.page_height = 7772400

                            rcdmesres_head = document.add_heading('',level=3)
                            rcdmesres_head_r = rcdmesres_head.add_run(a_df.loc['mespnlsa_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                            rcdmesres_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            rcdmesres_head_r.font.name = default_font
                            rcdmesres_head_r.font.size = Pt(14)
                            rcdmesres_head_r.font.color.rgb = RGBColor(0, 0, 0)
                            rcdmesres_head.paragraph_format.space_after = Pt(10)

                            # rcdmesres_rec = [
                            #     ('HP-1', "Q1", '12/10/2023', '5', 'cloudy', 'Pass', '250', '27', 'Pass', 'Pass', '' )
                            # ]

                            rcdmesres_t = document.add_table(rows=1, cols=11,style='GridTable4-Accent3')
                            rcdmesres_c = rcdmesres_t.rows[0].cells
                            rcdmesres_c[0].text = a_df.loc['mespnlsa_res_t1',lang]
                            rcdmesres_c[1].text = a_df.loc['mespnlsa_res_t2',lang]
                            rcdmesres_c[2].text = a_df.loc['mespnlsa_res_t3',lang]
                            rcdmesres_c[3].text = a_df.loc['mespnlsa_res_t4',lang]
                            rcdmesres_c[4].text = a_df.loc['mespnlsa_res_t5',lang]
                            rcdmesres_c[5].text = a_df.loc['mespnlsa_res_t6',lang]
                            rcdmesres_c[6].text = a_df.loc['mespnlsa_res_t7',lang]
                            rcdmesres_c[7].text = a_df.loc['mespnlsa_res_t8',lang]
                            #rcdmesres_c[8].text = 'Trip Status'
                            rcdmesres_c[8].text = a_df.loc['mespnlsa_res_t9',lang]
                            rcdmesres_c[9].text = a_df.loc['mespnlsa_res_t10',lang]


                            for col in ['Trip_time', 'Trip_current']:
                                fdf_rcdmes_rep[col] = fdf_rcdmes_rep[col].astype(str)

                            #for panid, rcdid, tdate, otemp, weather, butstat, ttime, tamp, tstat, ostat, remark in rcdmesres_rec:
                            for row in fdf_rcdmes_rep.itertuples():
                                rcdmesres_c = rcdmesres_t.add_row().cells
                                rcdmesres_c[0].text = row[1]#panid
                                rcdmesres_c[1].text = row[2]#rcdid
                                rcdmesres_c[2].text = row[3]#tdate
                                rcdmesres_c[3].text = row[4]#otemp
                                rcdmesres_c[4].text = row[5]#weather
                                rcdmesres_c[5].text = row[6]#butstat
                                rcdmesres_c[6].text = row[7]#ttime
                                rcdmesres_c[7].text = row[8]#tamp
                                #rcdmesres_c[8].text = row[9]#tstat
                                rcdmesres_c[8].text = row[9]#ostat
                                rcdmesres_c[9].text = row[10]#remark
                            for cell in rcdmesres_t.columns[0].cells:
                                cell.width = Cm(2)
                                if cell.text != a_df.loc['mespnlsa_res_t1',lang]:
                                    cell.paragraphs[0].runs[0].font.bold = False
                            for cell in rcdmesres_t.columns[1].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[2].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[3].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[4].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[5].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[6].cells:
                                cell.width = Cm(2.5)
                            for cell in rcdmesres_t.columns[7].cells:
                                cell.width = Cm(2.8)
                            for cell in rcdmesres_t.columns[8].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[9].cells:
                                cell.width = Cm(2)
                            for cell in rcdmesres_t.columns[10].cells:
                                cell.width = Cm(4)
                            #rcdmesres_t_space = document.add_paragraph()
                            #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                            #rcdbutres_t_space.paragraph_format.space_after = Pt(30)
                            
                            mes_subsec_no += 1

                    if audit_lvl_cct != 'None':

                    ## SUB-SECTION - IR TEST ##

                        section_irtest = document.add_section()
                        section_irtest.orientation = WD_ORIENT.PORTRAIT
                        section_irtest.page_width = 7772400
                        section_irtest.page_height = 10058400
                        
                        irtest_heading = document.add_heading('',level=2)
                        irtest_heading_run = irtest_heading.add_run(a_df.loc['mescctir_subsection',lang].format(mes_sec_no, mes_subsec_no))
                        irtest_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        irtest_heading_run.font.name = default_font
                        irtest_heading_run.font.size = Pt(16)
                        irtest_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        irtest_heading.paragraph_format.space_after = Pt(10)

                        paragh_irtest1 = document.add_paragraph()
                        paragh_irtest1.add_run(a_df.loc['mescctir_p1',lang])
                        paragh_irtest1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_irtest1.paragraph_format.space_after = Pt(10)

                        irtest_records = (
                            (a_df.loc['mescctir_t11',lang], a_df.loc['mescctir_t12',lang], a_df.loc['mescctir_t13',lang]),
                            (a_df.loc['mescctir_t21',lang], a_df.loc['mescctir_t22',lang], a_df.loc['mescctir_t23',lang]),
                            (a_df.loc['mescctir_t31',lang], a_df.loc['mescctir_t32',lang], a_df.loc['mescctir_t33',lang]),
                            (a_df.loc['mescctir_t41',lang], a_df.loc['mescctir_t42',lang], a_df.loc['mescctir_t43',lang])
                        )

                        irtest_table_desc = document.add_paragraph(a_df.loc['mescctir_tdesc',lang])
                        irtest_table_desc.paragraph_format.space_after = Pt(1)
                        irtest_table = document.add_table(rows=1, cols=3,style='GridTable4-Accent3')
                        irtest_cells = irtest_table.rows[0].cells
                        irtest_cells[0].text = a_df.loc['mescctir_tcol1',lang]
                        irtest_cells[1].text = a_df.loc['mescctir_tcol2',lang]
                        irtest_cells[2].text = a_df.loc['mescctir_tcol3',lang]
                        for descr, detail, comm in irtest_records:
                            irtest_cells = irtest_table.add_row().cells
                            irtest_cells[0].text = descr
                            irtest_cells[1].text = detail
                            irtest_cells[2].text = comm
                        for cell in irtest_table.columns[0].cells:
                            cell.width = Cm(6)
                            if cell.text != a_df.loc['mescctir_tcol1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in irtest_table.columns[1].cells:
                            cell.width = Cm(2)
                        for cell in irtest_table.columns[2].cells:
                            cell.width = Cm(10.5)
                        irtest_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        irtest_table_space.paragraph_format.space_after = Pt(30)

                        irtest_img = document.add_picture('ir_test.png',width=Inches(4))
                        irtest_img_par = document.paragraphs[-1]
                        irtest_img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        ## SUB2-SECTION - IR TEST FINDINGS ##

                        section_irfind = document.add_section()

                        irfind_heading = document.add_heading('',level=3)
                        irfind_heading_run = irfind_heading.add_run(a_df.loc['mescctir_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                        irfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        irfind_heading_run.font.name = default_font
                        irfind_heading_run.font.size = Pt(14)
                        irfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        irfind_heading.paragraph_format.space_after = Pt(10)

                        if ir_test_majority == 'perfect':

                            if ir_test_below10 == 'n' and ir_test_fault == 'n':
                                paragh_irfind1 = document.add_paragraph()
                                paragh_irfind1.add_run(a_df.loc['mescctir_perfnn',lang])
                                paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind1.paragraph_format.space_after = Pt(10)

                            else:

                                if ir_test_below10 == 'y' and ir_test_fault == 'n':
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_perfyn',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(5)

                                if ir_test_below10 == 'n' and ir_test_fault == 'y':
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_perfny',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(5)

                                if ir_test_below10 == 'y' and ir_test_fault == 'y': 
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_perfyy',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(5)

                                paragh_irfind2 = document.add_paragraph()
                                paragh_irfind2_run = paragh_irfind2.add_run(a_df.loc['mescctir_coz',lang])
                                paragh_irfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind2.paragraph_format.space_after = Pt(5)

                                bullet1_irfind2 = document.add_paragraph(style='List Bullet')
                                bullet1_irfind2.add_run(str(a_df.loc['mescctir_coz_b1',lang]).replace('\\n', '\n'))
                                bullet1_irfind2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_irfind2 = document.add_paragraph(style='List Bullet')
                                bullet2_irfind2.add_run(str(a_df.loc['mescctir_coz_b2',lang]).replace('\\n', '\n'))
                                bullet2_irfind2.paragraph_format.left_indent = Inches(0.5)

                                paragh_irfindout = document.add_paragraph()
                                paragh_irfindout.add_run(a_df.loc['mescctir_note',lang])
                                paragh_irfindout.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfindout.paragraph_format.space_after = Pt(10)

                        if ir_test_majority == 'good':

                            if ir_test_below10 == 'n' and ir_test_fault == 'n':
                                paragh_irfind1 = document.add_paragraph()
                                paragh_irfind1.add_run(a_df.loc['mescctir_goodnn',lang])
                                paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind1.paragraph_format.space_after = Pt(10)

                            else:

                                if ir_test_below10 == 'y' and ir_test_fault == 'n':
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_goodyn',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(10)

                                if ir_test_below10 == 'n' and ir_test_fault == 'y':
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_goodny',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(10)

                                if ir_test_below10 == 'y' and ir_test_fault == 'y':    
                                    paragh_irfind1 = document.add_paragraph()
                                    paragh_irfind1.add_run(a_df.loc['mescctir_goodyy',lang])
                                    paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_irfind1.paragraph_format.space_after = Pt(10)

                                paragh_irfind2 = document.add_paragraph()
                                paragh_irfind2_run = paragh_irfind2.add_run(a_df.loc['mescctir_coz',lang])
                                paragh_irfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind2.paragraph_format.space_after = Pt(5)

                                bullet1_irfind2 = document.add_paragraph(style='List Bullet')
                                bullet1_irfind2.add_run(str(a_df.loc['mescctir_coz_b1',lang]).replace('\\n', '\n'))
                                bullet1_irfind2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_irfind2 = document.add_paragraph(style='List Bullet')
                                bullet2_irfind2.add_run(str(a_df.loc['mescctir_coz_b2',lang]).replace('\\n', '\n'))
                                bullet2_irfind2.paragraph_format.left_indent = Inches(0.5)

                                paragh_irfindout = document.add_paragraph()
                                paragh_irfindout.add_run(a_df.loc['mescctir_note',lang])
                                paragh_irfindout.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfindout.paragraph_format.space_after = Pt(10)

                        if ir_test_majority == 'poor':       

                            if ir_test_fault == 'n':
                                paragh_irfind1 = document.add_paragraph()
                                paragh_irfind1.add_run(a_df.loc['mescctir_poorn',lang])
                                paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind1.paragraph_format.space_after = Pt(10)

                            if ir_test_fault == 'y':
                                paragh_irfind1 = document.add_paragraph()
                                paragh_irfind1.add_run(a_df.loc['mescctir_poory',lang])
                                paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_irfind1.paragraph_format.space_after = Pt(10)

                            paragh_irfind2 = document.add_paragraph()
                            paragh_irfind2_run = paragh_irfind2.add_run(a_df.loc['mescctir_coz',lang])
                            paragh_irfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_irfind2.paragraph_format.space_after = Pt(5)

                            bullet1_irfind2 = document.add_paragraph(style='List Bullet')
                            bullet1_irfind2.add_run(str(a_df.loc['mescctir_coz_b1',lang]).replace('\\n', '\n'))
                            bullet1_irfind2.paragraph_format.left_indent = Inches(0.5)
                            bullet2_irfind2 = document.add_paragraph(style='List Bullet')
                            bullet2_irfind2.add_run(str(a_df.loc['mescctir_coz_b2',lang]).replace('\\n', '\n'))
                            bullet2_irfind2.paragraph_format.left_indent = Inches(0.5)

                            paragh_irfindout = document.add_paragraph()
                            paragh_irfindout.add_run(a_df.loc['mescctir_note',lang])
                            paragh_irfindout.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_irfindout.paragraph_format.space_after = Pt(10)

                        if ir_test_majority == 'fault':

                            paragh_irfind1 = document.add_paragraph()
                            paragh_irfind1.add_run(a_df.loc['mescctir_fault',lang])
                            paragh_irfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_irfind1.paragraph_format.space_after = Pt(10)

                            paragh_irfind2 = document.add_paragraph()
                            paragh_irfind2_run = paragh_irfind2.add_run(a_df.loc['mescctir_coz',lang])
                            paragh_irfind2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_irfind2.paragraph_format.space_after = Pt(5)

                            bullet1_irfind2 = document.add_paragraph(style='List Bullet')
                            bullet1_irfind2.add_run(str(a_df.loc['mescctir_coz_b1',lang]).replace('\\n', '\n'))
                            bullet1_irfind2.paragraph_format.left_indent = Inches(0.5)
                            bullet2_irfind2 = document.add_paragraph(style='List Bullet')
                            bullet2_irfind2.add_run(str(a_df.loc['mescctir_coz_b2',lang]).replace('\\n', '\n'))
                            bullet2_irfind2.paragraph_format.left_indent = Inches(0.5)

                            paragh_irfindout = document.add_paragraph()
                            paragh_irfindout.add_run(a_df.loc['mescctir_note',lang])
                            paragh_irfindout.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_irfindout.paragraph_format.space_after = Pt(10)

                        paragh_irfindref = document.add_paragraph()
                        paragh_irfindref.add_run(a_df.loc['mescctir_ref',lang])
                        paragh_irfindref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_irfindref.paragraph_format.space_after = Pt(10)

                        ## SUB2-SECTION - IR TEST RESULTS ##

                        section_irresult = document.add_section()

                        section_irresult.orientation = WD_ORIENT.LANDSCAPE
                        section_irresult.page_width = 10058400
                        section_irresult.page_height = 7772400

                        irresult_heading = document.add_heading('',level=3)
                        irresult_heading_run = irresult_heading.add_run(a_df.loc['mescctir_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                        irresult_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        irresult_heading_run.font.name = default_font
                        irresult_heading_run.font.size = Pt(14)
                        irresult_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        irresult_heading.paragraph_format.space_after = Pt(10)

                        # irresult_records = [
                        #     ('1001', "2500", '12/10/2023', '5', 'sunny', '15', 'N/A', 'N/A', 'Good', 'Trend analysis' )
                        # ]

                        irresult_table = document.add_table(rows=1, cols=10,style='GridTable4-Accent3')
                        irresult_cells = irresult_table.rows[0].cells
                        irresult_cells[0].text = a_df.loc['mescctir_t1',lang]
                        irresult_cells[1].text = a_df.loc['mescctir_t2',lang]
                        irresult_cells[2].text = a_df.loc['mescctir_t3',lang]
                        irresult_cells[3].text = a_df.loc['mescctir_t4',lang]
                        irresult_cells[4].text = a_df.loc['mescctir_t5',lang]
                        irresult_cells[5].text = a_df.loc['mescctir_t6',lang]
                        irresult_cells[6].text = a_df.loc['mescctir_t7',lang]
                        irresult_cells[7].text = a_df.loc['mescctir_t8',lang]
                        irresult_cells[8].text = a_df.loc['mescctir_t9',lang]
                        irresult_cells[9].text = a_df.loc['mescctir_t10',lang]

                        fdf_irtest_rep['Voltage'] = fdf_irtest_rep['Voltage'].astype(int)
                        for col in ['Voltage', 'L1-PE', 'L2-PE', 'L3-PE']:
                            fdf_irtest_rep[col] = fdf_irtest_rep[col].astype(str)

                        #for cctid, volt, tdate, otemp, weather, irl1, irl2, irl3, irstat, remark in irresult_records:
                        for row in fdf_irtest_rep.itertuples():
                            irresult_cells = irresult_table.add_row().cells
                            irresult_cells[0].text = row[2]
                            irresult_cells[1].text = row[7]
                            irresult_cells[2].text = row[3]
                            irresult_cells[3].text = row[4]
                            irresult_cells[4].text = row[5]
                            irresult_cells[5].text = row[8]
                            irresult_cells[6].text = row[9]
                            irresult_cells[7].text = row[10]
                            irresult_cells[8].text = row[11]
                            irresult_cells[9].text = row[12]
                        for cell in irresult_table.columns[0].cells:
                            cell.width = Cm(2)
                            if cell.text != a_df.loc['mescctir_t1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in irresult_table.columns[1].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[2].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[3].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[4].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[5].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[6].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[7].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[8].cells:
                            cell.width = Cm(3)
                        for cell in irresult_table.columns[9].cells:
                            cell.width = Cm(4)
                        irresult_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        irresult_table_space.paragraph_format.space_after = Pt(30)

                        mes_subsec_no += 1
                        
                        ## SUB-SECTION - RESISTANCE TEST ##

                        section_resist = document.add_section()

                        section_resist.orientation = WD_ORIENT.PORTRAIT

                        section_resist.page_width = 7772400
                        section_resist.page_height = 10058400

                        resist_heading = document.add_heading('',level=2)
                        resist_heading_run = resist_heading.add_run(a_df.loc['mescctcont_subsection',lang].format(mes_sec_no, mes_subsec_no))
                        resist_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        resist_heading_run.font.name = default_font
                        resist_heading_run.font.size = Pt(16)
                        resist_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        resist_heading.paragraph_format.space_after = Pt(10)

                        paragh_resist = document.add_paragraph()
                        paragh_resist.add_run(a_df.loc['mescctcont_p1',lang])
                        paragh_resist.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_resist.paragraph_format.space_after = Pt(30)

                        resist_img = document.add_picture('cont_test.png',width=Inches(4))
                        resist_img_par = document.paragraphs[-1]
                        resist_img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        ## SUB2-SECTION - CONTINUITY FINDINGS ##

                        section_resistfind = document.add_section()

                        resistfind_heading = document.add_heading('',level=3)
                        resistfind_heading_run = resistfind_heading.add_run(a_df.loc['mescctcont_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                        resistfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        resistfind_heading_run.font.name = default_font
                        resistfind_heading_run.font.size = Pt(14)
                        resistfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        resistfind_heading.paragraph_format.space_after = Pt(10)

                        if cont_major == 'good':

                            if cont_short == 'n' and cont_broke == 'n':
                                paragh_resist1 = document.add_paragraph()
                                paragh_resist1.add_run(a_df.loc['mescctcont_goodnn',lang])
                                paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_resist1.paragraph_format.space_after = Pt(10)

                            else:

                                if cont_short == 'y' and cont_broke == 'n':
                                    paragh_resist1 = document.add_paragraph()
                                    paragh_resist1.add_run(a_df.loc['mescctcont_goodyn',lang])
                                    paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_resist1.paragraph_format.space_after = Pt(10)

                                if cont_short == 'n' and cont_broke == 'y':
                                    paragh_resist1 = document.add_paragraph()
                                    paragh_resist1.add_run(a_df.loc['mescctcont_goodny',lang])
                                    paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_resist1.paragraph_format.space_after = Pt(10)

                                if cont_short == 'y' and cont_broke == 'y':
                                    paragh_resist1 = document.add_paragraph()
                                    paragh_resist1.add_run(a_df.loc['mescctcont_goodyy',lang])
                                    paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    paragh_resist1.paragraph_format.space_after = Pt(10)

                                paragh_resist2 = document.add_paragraph()
                                paragh_resist2_run = paragh_resist2.add_run(a_df.loc['mescctcont_coz',lang])
                                paragh_resist2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_resist2.paragraph_format.space_after = Pt(5)

                                bullet1_resistfind2 = document.add_paragraph(style='List Bullet')
                                bullet1_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b1',lang]).replace('\\n', '\n'))
                                bullet1_resistfind2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_resistfind2 = document.add_paragraph(style='List Bullet')
                                bullet2_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b2',lang]).replace('\\n', '\n'))
                                bullet2_resistfind2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_resistfind2 = document.add_paragraph(style='List Bullet')
                                bullet3_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b3',lang]).replace('\\n', '\n'))
                                bullet3_resistfind2.paragraph_format.left_indent = Inches(0.5)
                                bullet4_resistfind2 = document.add_paragraph(style='List Bullet')
                                bullet4_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b4',lang]).replace('\\n', '\n'))
                                bullet4_resistfind2.paragraph_format.left_indent = Inches(0.5)         

                        if cont_major == 'poor':

                            if cont_short == 'y' and cont_broke == 'n':
                                paragh_resist1 = document.add_paragraph()
                                paragh_resist1.add_run(a_df.loc['mescctcont_pooryn',lang])
                                paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_resist1.paragraph_format.space_after = Pt(10)

                            if cont_short == 'n' and cont_broke == 'y':
                                paragh_resist1 = document.add_paragraph()
                                paragh_resist1.add_run(a_df.loc['mescctcont_poorny',lang])
                                paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_resist1.paragraph_format.space_after = Pt(10)

                            if cont_short == 'y' and cont_broke == 'y':
                                paragh_resist1 = document.add_paragraph()
                                paragh_resist1.add_run(a_df.loc['mescctcont_pooryy',lang])
                                paragh_resist1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_resist1.paragraph_format.space_after = Pt(10)

                            paragh_resist2 = document.add_paragraph()
                            paragh_resist2_run = paragh_resist2.add_run(a_df.loc['mescctcont_coz',lang])
                            paragh_resist2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragh_resist2.paragraph_format.space_after = Pt(5)

                            bullet1_resistfind2 = document.add_paragraph(style='List Bullet')
                            bullet1_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b1',lang]).replace('\\n', '\n'))
                            bullet1_resistfind2.paragraph_format.left_indent = Inches(0.5)
                            bullet2_resistfind2 = document.add_paragraph(style='List Bullet')
                            bullet2_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b2',lang]).replace('\\n', '\n'))
                            bullet2_resistfind2.paragraph_format.left_indent = Inches(0.5)
                            bullet3_resistfind2 = document.add_paragraph(style='List Bullet')
                            bullet3_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b3',lang]).replace('\\n', '\n'))
                            bullet3_resistfind2.paragraph_format.left_indent = Inches(0.5)
                            bullet4_resistfind2 = document.add_paragraph(style='List Bullet')
                            bullet4_resistfind2.add_run(str(a_df.loc['mescctcont_coz_b4',lang]).replace('\\n', '\n'))
                            bullet4_resistfind2.paragraph_format.left_indent = Inches(0.5)    

                        paragh_resistfindref = document.add_paragraph()
                        paragh_resistfindref.add_run(a_df.loc['mescctcont_ref',lang])
                        paragh_resistfindref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_resistfindref.paragraph_format.space_after = Pt(10)

                        ## SUB2-SECTION - CONTINUITY TEST RESULTS ##

                        section_contresult = document.add_section()

                        section_contresult.orientation = WD_ORIENT.LANDSCAPE
                        section_contresult.page_width = 10058400
                        section_contresult.page_height = 7772400

                        contresult_heading = document.add_heading('',level=3)
                        contresult_heading_run = contresult_heading.add_run(a_df.loc['mescctcont_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                        contresult_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        contresult_heading_run.font.name = default_font
                        contresult_heading_run.font.size = Pt(14)
                        contresult_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        contresult_heading.paragraph_format.space_after = Pt(10)

                        # contresult_records = [
                        #     ('1001', '12/10/2023', '5', 'sunny', '100', 'N/A', 'N/A', 'N/A', 'Good', 'Trend analysis' )
                        # ]

                        contresult_table = document.add_table(rows=1, cols=10,style='GridTable4-Accent3')
                        contresult_cells = contresult_table.rows[0].cells
                        contresult_cells[0].text = a_df.loc['mescctcont_t1',lang]
                        contresult_cells[1].text = a_df.loc['mescctcont_t2',lang]
                        contresult_cells[2].text = a_df.loc['mescctcont_t3',lang]
                        contresult_cells[3].text = a_df.loc['mescctcont_t4',lang]
                        contresult_cells[4].text = a_df.loc['mescctcont_t5',lang]
                        contresult_cells[5].text = a_df.loc['mescctcont_t6',lang]
                        contresult_cells[6].text = a_df.loc['mescctcont_t7',lang]
                        contresult_cells[7].text = a_df.loc['mescctcont_t8',lang]
                        contresult_cells[8].text = a_df.loc['mescctcont_t9',lang]
                        contresult_cells[9].text = a_df.loc['mescctcont_t10',lang]


                        for col in ['L-N', 'L1-L2', 'L1-L3', 'L2-L3']:
                            fdf_conttest_rep[col] = fdf_conttest_rep[col].astype(str)

                        #for cctid, tdate, otemp, weather, l1n, l1l2, l1l3, l2l3, stat, remark in contresult_records:
                        for row in fdf_conttest_rep.itertuples():
                            contresult_cells = contresult_table.add_row().cells
                            contresult_cells[0].text = row[1]
                            contresult_cells[1].text = row[2]
                            contresult_cells[2].text = row[3]
                            contresult_cells[3].text = row[4]
                            contresult_cells[4].text = row[6]
                            contresult_cells[5].text = row[7]
                            contresult_cells[6].text = row[8]
                            contresult_cells[7].text = row[9]
                            contresult_cells[8].text = row[10]
                            contresult_cells[9].text = row[11]
                        for cell in contresult_table.columns[0].cells:
                            cell.width = Cm(2)
                            if cell.text != a_df.loc['mescctcont_t1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in contresult_table.columns[1].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[2].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[3].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[4].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[5].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[6].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[7].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[8].cells:
                            cell.width = Cm(3)
                        for cell in contresult_table.columns[9].cells:
                            cell.width = Cm(4)
                        contresult_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        contresult_table_space.paragraph_format.space_after = Pt(30)

                        mes_subsec_no += 1
                        
                        ## SUB-SECTION - VOLTAGE MEASUREMENT ##

                        section_volt = document.add_section()

                        section_volt.orientation = WD_ORIENT.PORTRAIT

                        section_volt.page_width = 7772400
                        section_volt.page_height = 10058400

                        volt_heading = document.add_heading('',level=2)
                        volt_heading_run = volt_heading.add_run(a_df.loc['mescctvolt_subsection',lang].format(mes_sec_no, mes_subsec_no))
                        volt_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        volt_heading_run.font.name = default_font
                        volt_heading_run.font.size = Pt(16)
                        volt_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        volt_heading.paragraph_format.space_after = Pt(10)

                        paragh_volt = document.add_paragraph()
                        paragh_volt.add_run(a_df.loc['mescctvolt_p1',lang])
                        paragh_volt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_volt.paragraph_format.space_after = Pt(30)

                        ## 2SUB-SECTION - VOLTAGE FINDINGS ##

                        voltfind_heading = document.add_heading('',level=3)
                        voltfind_heading_run = voltfind_heading.add_run(a_df.loc['mescctvolt_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                        voltfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        voltfind_heading_run.font.name = default_font
                        voltfind_heading_run.font.size = Pt(14)
                        voltfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        voltfind_heading.paragraph_format.space_after = Pt(10)

                        if volt_major == 'in-range':

                            if volt_out == 'n' and volt_zero == 'n':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_corrnn',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                            if volt_out == 'y' and volt_zero == 'n':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_corryn',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)

                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)         

                            if volt_out == 'n' and volt_zero == 'y':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_corrny',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)

                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)

                            if volt_out == 'y' and volt_zero == 'y':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_corryy',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)

                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)         
                                bullet5_volt2 = document.add_paragraph(style='List Bullet')
                                bullet5_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b5',lang]).replace('\\n', '\n'))
                                bullet5_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet6_volt2 = document.add_paragraph(style='List Bullet')
                                bullet6_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b6',lang]).replace('\\n', '\n'))
                                bullet6_volt2.paragraph_format.left_indent = Inches(0.5)       

                        if volt_major == 'out-of-range':

                            if volt_out == 'y' and volt_zero == 'n':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_incorryn',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)

                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)         

                            if volt_out == 'n' and volt_zero == 'y':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_incorrny',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)
   
                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozny_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)

                            if volt_out == 'y' and volt_zero == 'y':
                                paragh_volt1 = document.add_paragraph()
                                paragh_volt1.add_run(a_df.loc['mescctvolt_incorryy',lang])
                                paragh_volt1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt1.paragraph_format.space_after = Pt(10)

                                paragh_volt2 = document.add_paragraph()
                                paragh_volt2_run = paragh_volt2.add_run(a_df.loc['mescctvolt_coz',lang])
                                paragh_volt2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_volt2.paragraph_format.space_after = Pt(5)

                                bullet1_volt2 = document.add_paragraph(style='List Bullet')
                                bullet1_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b1',lang]).replace('\\n', '\n'))
                                bullet1_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_volt2 = document.add_paragraph(style='List Bullet')
                                bullet2_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b2',lang]).replace('\\n', '\n'))
                                bullet2_volt2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_volt2 = document.add_paragraph(style='List Bullet')
                                bullet3_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b3',lang]).replace('\\n', '\n'))
                                bullet3_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet4_volt2 = document.add_paragraph(style='List Bullet')
                                bullet4_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b4',lang]).replace('\\n', '\n'))
                                bullet4_volt2.paragraph_format.left_indent = Inches(0.5)         
                                bullet5_volt2 = document.add_paragraph(style='List Bullet')
                                bullet5_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b5',lang]).replace('\\n', '\n'))
                                bullet5_volt2.paragraph_format.left_indent = Inches(0.5)     
                                bullet6_volt2 = document.add_paragraph(style='List Bullet')
                                bullet6_volt2.add_run(str(a_df.loc['mescctvolt_cozyy_b6',lang]).replace('\\n', '\n'))
                                bullet6_volt2.paragraph_format.left_indent = Inches(0.5)            

                        paragh_voltref = document.add_paragraph()
                        paragh_voltref.add_run(a_df.loc['mescctvolt_ref',lang])
                        paragh_voltref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_voltref.paragraph_format.space_after = Pt(10)


                        ## SUB2-SECTION - VOLTAGE TEST RESULTS ##

                        section_voltresult = document.add_section()

                        section_voltresult.orientation = WD_ORIENT.LANDSCAPE
                        section_voltresult.page_width = 10058400
                        section_voltresult.page_height = 7772400

                        voltresult_heading = document.add_heading('',level=3)
                        voltresult_heading_run = voltresult_heading.add_run(a_df.loc['mescctvolt_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                        voltresult_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        voltresult_heading_run.font.name = default_font
                        voltresult_heading_run.font.size = Pt(14)
                        voltresult_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        voltresult_heading.paragraph_format.space_after = Pt(10)

                        # voltresult_records = [
                        #     ('1001', '12/10/2023', '5', 'sunny', '230', 'N/A', 'N/A', 'N/A', 'Good', 'In range' )
                        # ]

                        for col in ['L-N', 'L1-L2', 'L1-L3', 'L2-L3']:
                            fdf_volttest_rep[col] = fdf_volttest_rep[col].astype(str)

                        voltresult_table = document.add_table(rows=1, cols=10,style='GridTable4-Accent3')
                        voltresult_cells = voltresult_table.rows[0].cells
                        voltresult_cells[0].text = a_df.loc['mescctvolt_t1',lang]
                        voltresult_cells[1].text = a_df.loc['mescctvolt_t2',lang]
                        voltresult_cells[2].text = a_df.loc['mescctvolt_t3',lang]
                        voltresult_cells[3].text = a_df.loc['mescctvolt_t4',lang]
                        voltresult_cells[4].text = a_df.loc['mescctvolt_t5',lang]
                        voltresult_cells[5].text = a_df.loc['mescctvolt_t6',lang]
                        voltresult_cells[6].text = a_df.loc['mescctvolt_t7',lang]
                        voltresult_cells[7].text = a_df.loc['mescctvolt_t8',lang]
                        voltresult_cells[8].text = a_df.loc['mescctvolt_t9',lang]
                        voltresult_cells[9].text = a_df.loc['mescctvolt_t10',lang]

                        for row in fdf_volttest_rep.itertuples():
                            voltresult_cells = voltresult_table.add_row().cells
                            voltresult_cells[0].text = row[1]
                            voltresult_cells[1].text = row[2]
                            voltresult_cells[2].text = row[3]
                            voltresult_cells[3].text = row[4]
                            voltresult_cells[4].text = row[6]
                            voltresult_cells[5].text = row[7]
                            voltresult_cells[6].text = row[8]
                            voltresult_cells[7].text = row[9]
                            voltresult_cells[8].text = row[10]
                            voltresult_cells[9].text = row[11]
                        for cell in voltresult_table.columns[0].cells:
                            cell.width = Cm(2)
                            if cell.text != a_df.loc['mescctvolt_t1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in voltresult_table.columns[1].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[2].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[3].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[4].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[5].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[6].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[7].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[8].cells:
                            cell.width = Cm(3)
                        for cell in voltresult_table.columns[9].cells:
                            cell.width = Cm(4)
                        voltresult_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        voltresult_table_space.paragraph_format.space_after = Pt(30)

                        mes_subsec_no += 1
                        
                        ## SUB-SECTION - CURRENT MEASUREMENT ##

                        section_curr = document.add_section()

                        section_curr.orientation = WD_ORIENT.PORTRAIT

                        section_curr.page_width = 7772400
                        section_curr.page_height = 10058400

                        curr_heading = document.add_heading('',level=2)
                        curr_heading_run = curr_heading.add_run(a_df.loc['mescctcurr_subsection',lang].format(mes_sec_no, mes_subsec_no))
                        curr_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        curr_heading_run.font.name = default_font
                        curr_heading_run.font.size = Pt(16)
                        curr_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        curr_heading.paragraph_format.space_after = Pt(10)

                        paragh_curr = document.add_paragraph()
                        paragh_curr.add_run(a_df.loc['mescctcurr_p1',lang])
                        paragh_curr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_curr.paragraph_format.space_after = Pt(30)

                        ## 2SUB-SECTION - CURRENT FINDINGS ##

                        currfind_heading = document.add_heading('',level=3)
                        currfind_heading_run = currfind_heading.add_run(a_df.loc['mescctcurr_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                        currfind_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        currfind_heading_run.font.name = default_font
                        currfind_heading_run.font.size = Pt(14)
                        currfind_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        currfind_heading.paragraph_format.space_after = Pt(10)

                        if curr_major == 'in-range':
                            if curr_trip == 'n' and curr_zero == 'n':
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_corrnn',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                            if curr_trip == 'y' and curr_zero == 'n':
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_corryn',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozyn',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet4_curr2 = document.add_paragraph(style='List Bullet')
                                bullet4_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet5_curr2 = document.add_paragraph(style='List Bullet')
                                bullet5_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b5',lang]).replace('\\n', '\n'))
                                bullet5_curr2.paragraph_format.left_indent = Inches(0.5)

                            if curr_trip == 'n' and curr_zero == 'y':   
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_corrny',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozny',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)

                            if curr_trip == 'y' and curr_zero == 'y':   
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_corryy',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)


                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozyn',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet4_curr2 = document.add_paragraph(style='List Bullet')
                                bullet4_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet5_curr2 = document.add_paragraph(style='List Bullet')
                                bullet5_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b5',lang]).replace('\\n', '\n'))
                                bullet5_curr2.paragraph_format.left_indent = Inches(0.5)

                                paragh_curr3 = document.add_paragraph()
                                paragh_curr3_run = paragh_curr3.add_run(a_df.loc['mescctcurr_cozny',lang])
                                paragh_curr3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr3.paragraph_format.space_after = Pt(5)

                                bullet6_curr2 = document.add_paragraph(style='List Bullet')
                                bullet6_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet6_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet7_curr2 = document.add_paragraph(style='List Bullet')
                                bullet7_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet7_curr2.paragraph_format.left_indent = Inches(0.5)

                        if curr_major == 'out-of-range':
                            if curr_trip == 'y' and curr_zero == 'n':
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_incorryn',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozyn',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet4_curr2 = document.add_paragraph(style='List Bullet')
                                bullet4_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet5_curr2 = document.add_paragraph(style='List Bullet')
                                bullet5_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b5',lang]).replace('\\n', '\n'))
                                bullet5_curr2.paragraph_format.left_indent = Inches(0.5)

                            if curr_trip == 'n' and curr_zero == 'y':
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_incorrny',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozny',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)     

                            if curr_trip == 'y' and curr_zero == 'y':   
                                paragh_curr1 = document.add_paragraph()
                                paragh_curr1.add_run(a_df.loc['mescctcurr_incorryy',lang])
                                paragh_curr1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr1.paragraph_format.space_after = Pt(10)

                                paragh_curr2 = document.add_paragraph()
                                paragh_curr2_run = paragh_curr2.add_run(a_df.loc['mescctcurr_cozyn',lang])
                                paragh_curr2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr2.paragraph_format.space_after = Pt(5)

                                bullet1_curr2 = document.add_paragraph(style='List Bullet')
                                bullet1_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b1',lang]).replace('\\n', '\n'))
                                bullet1_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet2_curr2 = document.add_paragraph(style='List Bullet')
                                bullet2_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b2',lang]).replace('\\n', '\n'))
                                bullet2_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet3_curr2 = document.add_paragraph(style='List Bullet')
                                bullet3_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b3',lang]).replace('\\n', '\n'))
                                bullet3_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet4_curr2 = document.add_paragraph(style='List Bullet')
                                bullet4_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b4',lang]).replace('\\n', '\n'))
                                bullet4_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet5_curr2 = document.add_paragraph(style='List Bullet')
                                bullet5_curr2.add_run(str(a_df.loc['mescctcurr_cozyn_b5',lang]).replace('\\n', '\n'))
                                bullet5_curr2.paragraph_format.left_indent = Inches(0.5)

                                paragh_curr3 = document.add_paragraph()
                                paragh_curr3_run = paragh_curr3.add_run(a_df.loc['mescctcurr_cozny',lang])
                                paragh_curr3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                paragh_curr3.paragraph_format.space_after = Pt(5)

                                bullet6_curr2 = document.add_paragraph(style='List Bullet')
                                bullet6_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b2',lang]).replace('\\n', '\n'))
                                bullet6_curr2.paragraph_format.left_indent = Inches(0.5)
                                bullet7_curr2 = document.add_paragraph(style='List Bullet')
                                bullet7_curr2.add_run(str(a_df.loc['mescctcurr_cozny_b3',lang]).replace('\\n', '\n'))
                                bullet7_curr2.paragraph_format.left_indent = Inches(0.5)

                        paragh_currref = document.add_paragraph()
                        paragh_currref.add_run(a_df.loc['mescctcurr_ref',lang])
                        paragh_currref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragh_currref.paragraph_format.space_after = Pt(10)


                        ## SUB2-SECTION - CURRENT TEST RESULTS ##

                        section_currresult = document.add_section()

                        section_currresult.orientation = WD_ORIENT.LANDSCAPE
                        section_currresult.page_width = 10058400
                        section_currresult.page_height = 7772400

                        currresult_heading = document.add_heading('',level=3)
                        currresult_heading_run = currresult_heading.add_run(a_df.loc['mescctcurr_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                        currresult_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        currresult_heading_run.font.name = default_font
                        currresult_heading_run.font.size = Pt(14)
                        currresult_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                        currresult_heading.paragraph_format.space_after = Pt(10)

                        # currresult_records = [
                        #     ('1001', '12/10/2023', '5', 'sunny', '10', 'N/A', 'N/A', 'Good', 'In range' )
                        # ]

                        for col in ['L1', 'L2', 'L3']:
                            fdf_currtest_rep[col] = fdf_currtest_rep[col].astype(str)

                        currresult_table = document.add_table(rows=1, cols=9,style='GridTable4-Accent3')
                        currresult_cells = currresult_table.rows[0].cells
                        currresult_cells[0].text = a_df.loc['mescctcurr_t1',lang]
                        currresult_cells[1].text = a_df.loc['mescctcurr_t2',lang]
                        currresult_cells[2].text = a_df.loc['mescctcurr_t3',lang]
                        currresult_cells[3].text = a_df.loc['mescctcurr_t4',lang]
                        currresult_cells[4].text = a_df.loc['mescctcurr_t5',lang]
                        currresult_cells[5].text = a_df.loc['mescctcurr_t6',lang]
                        currresult_cells[6].text = a_df.loc['mescctcurr_t7',lang]
                        currresult_cells[7].text = a_df.loc['mescctcurr_t8',lang]
                        currresult_cells[8].text = a_df.loc['mescctcurr_t9',lang]

                        for row in fdf_currtest_rep.itertuples():
                            currresult_cells = currresult_table.add_row().cells
                            currresult_cells[0].text = row[1]
                            currresult_cells[1].text = row[2]
                            currresult_cells[2].text = row[3]
                            currresult_cells[3].text = row[4]
                            currresult_cells[4].text = row[6]
                            currresult_cells[5].text = row[7]
                            currresult_cells[6].text = row[8]
                            currresult_cells[7].text = row[9]
                            currresult_cells[8].text = row[10]
                        for cell in currresult_table.columns[0].cells:
                            cell.width = Cm(2)
                            if cell.text != a_df.loc['mescctcurr_t1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in currresult_table.columns[1].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[2].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[3].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[4].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[5].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[6].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[7].cells:
                            cell.width = Cm(3)
                        for cell in currresult_table.columns[8].cells:
                            cell.width = Cm(4)
                        currresult_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        currresult_table_space.paragraph_format.space_after = Pt(30)

                        mes_subsec_no += 1

                    ##########################################################################################################
                    ##########################################################################################################
                    ##########################################################################################################

                    if audit_lvl_insul == 'Advanced':
                        
                        s_insmes = document.add_section()
                        
                        s_insmes.orientation = WD_ORIENT.PORTRAIT
                        s_insmes.page_width = 7772400
                        s_insmes.page_height = 10058400
                        
                        h_insmes = document.add_heading('',level=2)
                        h_insmes_r = h_insmes.add_run(a_df.loc['mesins_subsection',lang].format(mes_sec_no, mes_subsec_no))
                        h_insmes.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_insmes_r.font.name = default_font
                        h_insmes_r.font.size = Pt(16)
                        h_insmes_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_insmes.paragraph_format.space_after = Pt(10)  
                        
                        p_insmes1 = document.add_paragraph()
                        p_insmes1.add_run(a_df.loc['mesins_p1',lang])
                        p_insmes1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_insmes1.paragraph_format.space_after = Pt(30)  

                        ## SUB2-SECTION - RCD MEASUREMENTS TEST FINDINGS ##
                        
                        h_insfind = document.add_heading('',level=3)
                        h_insfind_r = h_insfind.add_run(a_df.loc['mesins_subsection2_find',lang].format(mes_sec_no, mes_subsec_no))
                        h_insfind.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_insfind_r.font.name = default_font
                        h_insfind_r.font.size = Pt(14)
                        h_insfind_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_insfind.paragraph_format.space_after = Pt(10)    

                        
                        if inscheck_majority == 'correct':
                            #if (pipedia_fail == 'n' and insthick_fail == 'n' and instype_fail == 'n'):
                            if incheck_some == 'n':
                                p_insfind1 = document.add_paragraph()
                                p_insfind1.add_run(a_df.loc['mesins_corrn',lang])
                                p_insfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                
                            if incheck_some == 'y':
                                p_insfind1 = document.add_paragraph()
                                p_insfind1.add_run(a_df.loc['mesins_corry',lang])
                                p_insfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        if inscheck_majority == 'incorrect':
                            p_insfind1 = document.add_paragraph()
                            p_insfind1.add_run(a_df.loc['mesins_incorr',lang])
                            p_insfind1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        p_insrref = document.add_paragraph()
                        p_insrref.add_run(a_df.loc['mesins_ref',lang])
                        p_insrref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_insrref.paragraph_format.space_after = Pt(10)

                    ################################################
                    ## SUB2-SECTION - PIPING & INSULATION RESULTS ##
                    ################################################

                        s_insresult = document.add_section()

                        s_insresult.orientation = WD_ORIENT.LANDSCAPE
                        s_insresult.page_width = 10058400
                        s_insresult.page_height = 7772400

                        h_insresult = document.add_heading('',level=3)
                        h_insresult_r = h_insresult.add_run(a_df.loc['mesins_subsection2_res',lang].format(mes_sec_no, mes_subsec_no))
                        h_insresult.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        h_insresult_r.font.name = default_font
                        h_insresult_r.font.size = Pt(14)
                        h_insresult_r.font.color.rgb = RGBColor(0, 0, 0)
                        h_insresult.paragraph_format.space_after = Pt(10)

                        # insresult_records = [
                        #     ('1','10001', '50', '50', 'RW', 'RW', '30', '30', 'mm', 'Correct', '')
                        # ]
                        insresult_table = document.add_table(rows=1, cols=10,style='GridTable4-Accent3')
                        insresult_cells = insresult_table.rows[0].cells
                        insresult_cells[0].text = a_df.loc['mesins_t1',lang]
                        insresult_cells[1].text = a_df.loc['mesins_t2',lang]
                        insresult_cells[2].text = a_df.loc['mesins_t3',lang]
                        insresult_cells[3].text = a_df.loc['mesins_t4',lang]
                        insresult_cells[4].text = a_df.loc['mesins_t5',lang]
                        insresult_cells[5].text = a_df.loc['mesins_t6',lang]
                        #insresult_cells[3].text = 'Pipe DN Act.'
                        insresult_cells[6].text = a_df.loc['mesins_t7',lang]
                        #insresult_cells[5].text = 'Ins. Type Act.'
                        insresult_cells[7].text = a_df.loc['mesins_t8',lang]
                        #insresult_cells[7].text = 'Ins. Thk. Act.'
                        insresult_cells[8].text = a_df.loc['mesins_t9',lang]
                        insresult_cells[9].text = a_df.loc['mesins_t10',lang]
                        #insresult_cells[7].text = 'Remarks'


                        for col in ['Section', 'PipeDiameter', 'InsulationThickness']:
                            fdf_insparam_rep[col] = fdf_insparam_rep[col].astype(str)

                        #for sec, cctid, spipe, apipe, sinstype, ainstype, sinsthk, ainsthk, uom, stat, remark in insresult_records:
                        for row in fdf_insparam_rep.itertuples():
                            insresult_cells = insresult_table.add_row().cells
                            insresult_cells[0].text = row[1]#cctid
                            insresult_cells[1].text = row[2]#sec
                            insresult_cells[2].text = row[3]#date
                            insresult_cells[3].text = row[4]#temp
                            insresult_cells[4].text = row[5]#weather
                            #insresult_cells[2].text = spipe
                            insresult_cells[5].text = row[6]#apipe
                            #insresult_cells[4].text = sinstype
                            insresult_cells[6].text = row[7]#ainstype
                            #insresult_cells[6].text = sinsthk
                            insresult_cells[7].text = row[8]#ainsthk
                            insresult_cells[8].text = row[9]#uom
                            insresult_cells[9].text = row[10]#stat
                            #insresult_cells[10].text = remark
                        for cell in insresult_table.columns[0].cells:
                            cell.width = Cm(2)
                            if cell.text != a_df.loc['mesins_t1',lang]:
                                cell.paragraphs[0].runs[0].font.bold = False
                        for cell in insresult_table.columns[1].cells:
                            cell.width = Cm(1.8)
                        for cell in insresult_table.columns[2].cells:
                            cell.width = Cm(2.5)
                        for cell in insresult_table.columns[3].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[4].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[5].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[6].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[7].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[8].cells:
                            cell.width = Cm(2)
                        for cell in insresult_table.columns[9].cells:
                            cell.width = Cm(2)
                        # for cell in insresult_table.columns[10].cells:
                        #     cell.width = Cm(4)
                        #currresult_table_space = document.add_paragraph()
                        #ehtpanel_table_space.paragraph_format.space_before = Pt(0.1)
                        #currresult_table_space.paragraph_format.space_after = Pt(30)

                        mes_subsec_no += 1

                    ##################################
                    ### SECTION - SUMMARY ###
                    ##################################
                    try:
                        summ_sec_no = mes_sec_no + 1
                    except:
                        summ_sec_no = vis_sec_no + 1
                    # else:
                    #     summ_sec_no = mes_sec_no + 1

                    section_summary = document.add_section()
                    section_summary.orientation = WD_ORIENT.PORTRAIT
                    section_summary.page_width = 7772400
                    section_summary.page_height = 10058400
                        
                    summary_heading = document.add_heading('',level=1)
                    summary_heading_run = summary_heading.add_run(a_df.loc['sum_section',lang].format(summ_sec_no))
                    summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    summary_heading_run.font.name = default_font
                    summary_heading_run.font.size = Pt(18)
                    summary_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    summary_heading.paragraph_format.space_after = Pt(10)

                    if (audited_375 == 'y' and audited_374 == 'y' and audited_373 == 'y'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_yyy_p1',lang])
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                        p_summ2 = document.add_paragraph()                          
                        p_summ2.add_run(a_df.loc['sum_yyy_p2',lang].format(audit_lvl_panel, audit_lvl_cct, audit_lvl_insul))
                        p_summ2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ2.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'y' and audited_374 == 'n' and audited_373 == 'y'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_yny_p1',lang])
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                        p_summ2 = document.add_paragraph()                          
                        p_summ2.add_run(a_df.loc['sum_yny_p2',lang].format(audit_lvl_panel, audit_lvl_cct))
                        p_summ2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ2.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'y' and audited_374 == 'y' and audited_373 == 'n'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_yyn_p1',lang])
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                        p_summ2 = document.add_paragraph()                          
                        p_summ2.add_run(a_df.loc['sum_yyn_p2',lang].format(audit_lvl_panel, audit_lvl_insul))
                        p_summ2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ2.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'n' and audited_374 == 'y' and audited_373 == 'y'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_nyy_p1',lang])
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                        p_summ2 = document.add_paragraph()                          
                        p_summ2.add_run(a_df.loc['sum_nyy_p2',lang].format(audit_lvl_cct, audit_lvl_insul))
                        p_summ2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ2.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'y' and audited_374 == 'n' and audited_373 == 'n'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_ynn_p1',lang].format(audit_lvl_panel))
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'n' and audited_374 == 'y' and audited_373 == 'n'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_nyn_p1',lang].format(audit_lvl_insul))
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                    elif (audited_375 == 'n' and audited_374 == 'n' and audited_373 == 'y'):
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_nyn_p1',lang].format(audit_lvl_cct))
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)
                    else:
                        p_summ1 = document.add_paragraph()        
                        p_summ1.add_run(a_df.loc['sum_nnn_p1',lang])
                        p_summ1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ1.paragraph_format.space_after = Pt(5)

                    panel_score_lst = []
                    cct_score_lst = []
                    insul_score_lst = []

                    if audited_375 == 'y':
                        if pnl_cnt != 1:
                            p_summ_pnl1 = document.add_paragraph()        
                            p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c2',lang].format(pnl_cnt))
                            if pnl_cnt_fault != 1:
                                p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c2f2',lang].format(pnl_cnt_fault))
                            else:
                                p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c2f1',lang].format(pnl_cnt_fault))
                        else:
                            p_summ_pnl1 = document.add_paragraph()        
                            p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c1',lang].format(pnl_cnt))
                            if pnl_cnt_fault != 0:                   
                                p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c1f1',lang])
                            else:
                                p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_c1f0',lang])
                        if audit_lvl_panel == 'Advanced':                         
                            p_summ_pnl1.add_run(a_df.loc['sumpnl_p1_ca',lang])
                        p_summ_pnl1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ_pnl1.paragraph_format.space_after = Pt(5)

                        panel_score_lst.append(1-f_pnl_fault)

                        if audit_lvl_panel == 'Basic':
                            p_summ_pnl2 = document.add_paragraph() 
                            p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b',lang])
                            if cnt_butfault > 1:
                                # Many fault meaning = many RCDs
                                p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c2',lang].format(cnt_butfault, f_rcdbutfail*100, cnt_butall))
                            elif cnt_butfault == 1:
                                if cnt_butall > 1:
                                    # 1 fault, many RCDs
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c2f1',lang].format(cnt_butfault, f_rcdbutfail*100, cnt_butall))
                                else:
                                    # 1 fault, 1 RCD
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c1f1',lang])
                            elif cnt_butfault == 0:
                                if cnt_butall > 1:
                                    # Many RCDs, no faults
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c2f0',lang].format(cnt_butfault))
                                elif cnt_butall == 1:
                                    # 1 RCD, no faults
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c1f0',lang])
                                elif cnt_butall == 0:
                                    # No RCDs, no faults
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_b_c0f0',lang])

                            if cnt_butall != 0:
                                panel_score_lst.append(1-f_rcdbutfail)

                        else:
                            df_tot_failed_rcds = fdf_rcdmes_rep[(fdf_rcdmes_rep.Button_test=='Failed') | (fdf_rcdmes_rep.Status=='Failed')]
                            tot_failed_rcds = len(df_tot_failed_rcds.RCD_ID.unique().tolist())
                            if cnt_mesall != 0:
                                f_tot_failed_rcds = tot_failed_rcds / cnt_mesall
                            else:
                                f_tot_failed_rcds = 0
                            p_summ_pnl2 = document.add_paragraph() 
                            p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_sa',lang])
                            if tot_failed_rcds != 1:
                                if tot_failed_rcds != 0:
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_sa_c21',lang].format(tot_failed_rcds, f_tot_failed_rcds*100, cnt_mesall))
                                elif cnt_mesall != 0:
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_sa_c20',lang].format(cnt_mesall))
                                else:
                                    p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_sa_c00',lang])                                 
                            else:
                                p_summ_pnl2.add_run(a_df.loc['sumpnl_p1_sa_c1',lang].format(cnt_mesall, tot_failed_rcds))

                            panel_score_lst.append(1-f_tot_failed_rcds)

                        p_summ_pnl2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ_pnl2.paragraph_format.space_after = Pt(5)

                    if audited_373 == 'y':
                        if cct_cnt != 1:
                            p_summ_cct1 = document.add_paragraph()        
                            p_summ_cct1.add_run(a_df.loc['sumcct_p1_c2',lang].format(cct_cnt, cct_cnt_fault))
                        else:
                            p_summ_cct1 = document.add_paragraph()        
                            p_summ_cct1.add_run(a_df.loc['sumcct_p1_c1',lang].format(cct_cnt))
                            if cct_cnt_fault != 0:                   
                                p_summ_cct1.add_run(a_df.loc['sumcct_p1_c1f1',lang])
                            else:
                                p_summ_cct1.add_run(a_df.loc['sumcct_p1_c1f0',lang])

                        cct_score_lst.append(1-f_cct_fault)

                        p_summ_cct1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ_cct1.paragraph_format.space_after = Pt(5)

                        if th_cnt != 0:
                            if th_cnt != 1:
                                p_summ_th1 = document.add_paragraph()        
                                p_summ_th1.add_run(a_df.loc['sumcctth_p1_c2',lang].format(th_cnt, th_cnt_fault))
                                if th_cnt_fault == 1:
                                    p_summ_th1.add_run(a_df.loc['sumcctth_p1_c2f1',lang])
                                else:
                                    p_summ_th1.add_run(a_df.loc['sumcctth_p1_c2fe',lang])
                            else:
                                p_summ_th1 = document.add_paragraph()        
                                p_summ_th1.add_run(a_df.loc['sumcctth_p1_c1',lang].format(th_cnt))
                                if th_cnt_fault != 0:                   
                                    p_summ_th1.add_run(a_df.loc['sumcctth_p1_c1f1',lang])
                                else:
                                    p_summ_th1.add_run(a_df.loc['sumcctth_p1_c1f0',lang])
                            cct_score_lst.append(1-f_th_fault)
                            p_summ_th1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_summ_th1.paragraph_format.space_after = Pt(5)

                        p_summ_mes1 = document.add_paragraph()        
                        p_summ_mes1.add_run(a_df.loc['sumcct_m1',lang])
                        p_summ_mes1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ_mes1.paragraph_format.space_after = Pt(5)

                        #IR Test
                        p_summ_ir1 = document.add_paragraph()
                        p_summ_ir1.add_run(a_df.loc['sumcctir_p1',lang])
                        p_summ_ir1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_summ_ir1.paragraph_format.space_after = Pt(1)
                        fdf_irtest_rep_summ = fdf_irtest_rep.groupby('Status').agg({'ID':'count'}).reset_index()
                        map_irscore = {
                            'Perfect':1,
                            'Good':0.7,
                            'Poor':0.3,
                            'Fault':0
                            }
                        fdf_irtest_rep_score = fdf_irtest_rep_summ[['ID','Status']]
                        fdf_irtest_rep_score['Status'] = fdf_irtest_rep_score['Status'].replace(map_irscore)
                        fdf_irtest_rep_score['Total'] = fdf_irtest_rep_score['ID'] * fdf_irtest_rep_score['Status']
                        cct_score_lst.append(fdf_irtest_rep_score['Total'].sum()/fdf_irtest_rep_score['ID'].sum())                     
                        fdf_irtest_rep_summ['ID'] = fdf_irtest_rep_summ['ID'].astype(str)

                        # summir_table = document.add_table(rows=1, cols=2,style='GridTable4-Accent3')
                        # summir_cells = summir_table.rows[0].cells
                        # summir_cells[0].text = '#EHT Circuits'
                        # summir_cells[1].text = 'Status'

                        # for row in fdf_irtest_rep_summ.itertuples():
                        #     summir_cells = summir_table.add_row().cells
                        #     summir_cells[0].text = row[2]
                        #     summir_cells[1].text = row[1]

                        # for cell in summir_table.columns[0].cells:
                        #     cell.width = Cm(4)
                        #     if cell.text != '#EHT Circuits':
                        #         cell.paragraphs[0].runs[0].font.bold = False
                        # for cell in summir_table.columns[1].cells:
                        #     cell.width = Cm(4)

                        for index, row in fdf_irtest_rep_summ.iterrows():
                            p_summ_irx = document.add_paragraph(style='List Bullet')
                            p_summ_irx.add_run(f"{row['ID']} - {row['Status']} ")
                            p_summ_irx.paragraph_format.left_indent = Inches(0.5)

                        # p_summ_ir2 = document.add_paragraph()
                        # p_summ_ir2.paragraph_format.space_after = Pt(5)

                        #Continuity Test
                        p_summ_cont1 = document.add_paragraph()
                        p_summ_cont1.add_run(a_df.loc['sumcctcont_p1',lang])
                        p_summ_cont1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_summ_cont1.paragraph_format.space_after = Pt(1)
                        fdf_conttest_rep_summ = fdf_conttest_rep.groupby('Status').agg({'ID':'count'}).reset_index()

                        map_contscore = {
                            'Tentatively acceptable':0.8,
                            'Fault':0
                            }
                        fdf_conttest_rep_score = fdf_conttest_rep_summ[['ID','Status']]
                        fdf_conttest_rep_score['Status'] = fdf_conttest_rep_score['Status'].replace(map_contscore)
                        fdf_conttest_rep_score['Total'] = fdf_conttest_rep_score['ID'] * fdf_conttest_rep_score['Status']                          
                        cct_score_lst.append(fdf_conttest_rep_score['Total'].sum()/fdf_conttest_rep_score['ID'].sum())    
                        fdf_conttest_rep_summ['ID'] = fdf_conttest_rep_summ['ID'].astype(str)

                        for index, row in fdf_conttest_rep_summ.iterrows():
                            p_summ_contx = document.add_paragraph(style='List Bullet')
                            p_summ_contx.add_run(f"{row['ID']} - {row['Status']} ")
                            p_summ_contx.paragraph_format.left_indent = Inches(0.5)

                        # p_summ_cont2 = document.add_paragraph()
                        # p_summ_cont2.paragraph_format.space_after = Pt(5)

                        #Voltage Test
                        p_summ_volt1 = document.add_paragraph()
                        p_summ_volt1.add_run(a_df.loc['sumcctvolt_p1',lang])
                        p_summ_volt1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_summ_volt1.paragraph_format.space_after = Pt(1)
                        fdf_volttest_rep_summ = fdf_volttest_rep.groupby('Status').agg({'ID':'count'}).reset_index()
                        map_voltscore = {
                            'Correct':1,
                            'Fault':0
                            }
                        fdf_volttest_rep_score = fdf_volttest_rep_summ[['ID','Status']]
                        fdf_volttest_rep_score['Status'] = fdf_volttest_rep_score['Status'].replace(map_voltscore)
                        fdf_volttest_rep_score['Total'] = fdf_volttest_rep_score['ID'] * fdf_volttest_rep_score['Status']                        
                        cct_score_lst.append(fdf_volttest_rep_score['Total'].sum()/fdf_volttest_rep_score['ID'].sum())

                        fdf_volttest_rep_summ['ID'] = fdf_volttest_rep_summ['ID'].astype(str)

                        for index, row in fdf_volttest_rep_summ.iterrows():
                            p_summ_voltx = document.add_paragraph(style='List Bullet')
                            p_summ_voltx.add_run(f"{row['ID']} - {row['Status']} ")
                            p_summ_voltx.paragraph_format.left_indent = Inches(0.5)

                        # p_summ_volt2 = document.add_paragraph()
                        # p_summ_volt2.paragraph_format.space_after = Pt(5)

                        #Current Test
                        p_summ_curr1 = document.add_paragraph()
                        p_summ_curr1.add_run(a_df.loc['sumcctcurr_p1',lang])
                        p_summ_curr1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_summ_curr1.paragraph_format.space_after = Pt(1)
                        fdf_currtest_rep_summ = fdf_currtest_rep.groupby('Status').agg({'ID':'count'}).reset_index()
                        map_currscore = {
                            'Tentatively acceptable':0.8,
                            'Fault':0
                            }
                        fdf_currtest_rep_score = fdf_currtest_rep_summ[['ID','Status']]
                        fdf_currtest_rep_score['Status'] = fdf_currtest_rep_score['Status'].replace(map_currscore)
                        fdf_currtest_rep_score['Total'] = fdf_currtest_rep_score['ID'] * fdf_currtest_rep_score['Status']                        
                        cct_score_lst.append(fdf_currtest_rep_score['Total'].sum()/fdf_currtest_rep_score['ID'].sum())
                        fdf_currtest_rep_summ['ID'] = fdf_currtest_rep_summ['ID'].astype(str)
  
                        for index, row in fdf_currtest_rep_summ.iterrows():
                            p_summ_currx = document.add_paragraph(style='List Bullet')
                            p_summ_currx.add_run(f"{row['ID']} - {row['Status']} ")
                            p_summ_currx.paragraph_format.left_indent = Inches(0.5)

                        p_summ_curr2 = document.add_paragraph()
                        p_summ_curr2.paragraph_format.space_after = Pt(5)

                    if audited_374 == 'y':
                        if ins_cnt != 1:
                            p_summ_ins1 = document.add_paragraph()        
                            p_summ_ins1.add_run(a_df.loc['sumins_p1_c2',lang].format(ins_cnt, ins_cnt_fault))
                        else:
                            p_summ_ins1 = document.add_paragraph()        
                            p_summ_ins1.add_run(a_df.loc['sumins_p1_c1',lang].format(ins_cnt, ins_cnt_fault))
                            if ins_cnt_fault != 0:                   
                                p_summ_ins1.add_run(a_df.loc['sumins_p1_c1f1',lang])
                            else:
                                p_summ_ins1.add_run(a_df.loc['sumins_p1_c1f0',lang])

                        insul_score_lst.append(1-f_ins_fault)

                        if (audit_lvl_insul == 'Standard' or audit_lvl_insul == 'Advanced'):                         
                            p_summ_ins1.add_run(a_df.loc['sumins_p1_sa',lang])
                        p_summ_ins1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_summ_ins1.paragraph_format.space_after = Pt(5)

                        if audit_lvl_insul == 'Advanced':
                            p_summ_ins2 = document.add_paragraph() 
                            p_summ_ins2.add_run(a_df.loc['sumins_p2_sa',lang])
                            if cnt_insparam_fault > 1:
                                if cnt_insparam_all != 1:
                                    # Among many checkups there were many issues
                                    p_summ_ins2.add_run(a_df.loc['sumins_p2_sa_f2c2',lang].format(cnt_insparam_all, cnt_insparam_fault, int(round(f_insparam_fault,2)*100)))
                                else:
                                    # Among 1 checkup there was many issues, not likely
                                    p_summ_ins2.add_run(a_df.loc['sumins_p2_sa_f2c1',lang].format(cnt_insparam_all, cnt_insparam_fault))
                            elif cnt_insparam_fault == 0:
                                if cnt_insparam_all != 1:
                                    # Among many checkups there was 0 issues
                                    p_summ_ins2.add_run(a_df.loc['sumins_p2_sa_f0c2',lang].format(cnt_insparam_all, cnt_insparam_fault))
                                else:
                                    # Among 1 checkup there was 0 issues
                                    p_summ_ins2.add_run(a_df.loc['sumins_p2_sa_f0c1',lang].format(cnt_insparam_all, cnt_insparam_fault))                                  
                            else:
                                # 1 issue
                                if cnt_insparam_all != 1:
                                    # Among many checkups there was 1 issue
                                    p_summ_ins2.add_run(a_df.loc['sumins_p2_sa_f1c2',lang].format(cnt_insparam_all, cnt_insparam_fault, int(round(f_insparam_fault,2)*100)))
                                else:
                                    # Among 1 checkup there was 1 issue, 100%
                                    p_summ_ins2.add_run(a_df.loc['ssumins_p2_sa_f1c1',lang].format(cnt_insparam_all, cnt_insparam_fault, int(round(f_insparam_fault,2)*100)))
                        
                            insul_score_lst.append(1-f_insparam_fault)
                            p_summ_ins2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_summ_ins2.paragraph_format.space_after = Pt(5)

                    p_summ_score_final = document.add_paragraph() 
                    if audited_375 == 'y':
                        #panel_score_val = str(round(np.mean(panel_score_lst)*100,0))#[:2]
                        try:
                            panel_score_val = int(round(np.mean(panel_score_lst),2)*100)
                        except:
                            panel_score_val = 0
                        p_summ_score_panel_r = p_summ_score_final.add_run(a_df.loc['sumscore_pnl',lang].format(panel_score_val))
                        p_summ_score_panel_r.font.size = Pt(14)
                        p_summ_score_final.add_run(
                            "\n")
                    if audited_373 == 'y':
                        #cct_score_val = str(round(np.mean(cct_score_lst)*100,0))#[:2]
                        try:
                            cct_score_val = int(round(np.mean(cct_score_lst),2)*100)
                        except:
                            cct_score_val = 0
                        p_summ_score_cct_r = p_summ_score_final.add_run(a_df.loc['sumscore_cct',lang].format(cct_score_val))
                        p_summ_score_cct_r.font.size = Pt(14)
                        p_summ_score_final.add_run(
                            "\n")
                    if audited_374 == 'y':
                        #insul_score_val = str(round(np.mean(insul_score_lst)*100,0))#[:2]
                        try:
                            insul_score_val = int(round(np.mean(insul_score_lst),2)*100)
                        except:
                            insul_score_val = 0
                        p_summ_score_insul_r = p_summ_score_final.add_run(a_df.loc['sumscore_ins',lang].format(insul_score_val))
                        p_summ_score_insul_r.font.size = Pt(14)
                    p_summ_score_final.paragraph_format.keep_together      
                    ##################################
                    ### SECTION - ATTACHMENTS ###
                    ##################################
                    att_sec_no = summ_sec_no + 1
                    section_attach = document.add_section()
                    section_attach.orientation = WD_ORIENT.PORTRAIT
                    section_attach.page_width = 7772400
                    section_attach.page_height = 10058400
                        
                    attach_heading = document.add_heading('',level=1)
                    attach_heading_run = attach_heading.add_run(a_df.loc['att_section',lang].format(att_sec_no))
                    attach_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    attach_heading_run.font.name = default_font
                    attach_heading_run.font.size = Pt(18)
                    attach_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    attach_heading.paragraph_format.space_after = Pt(10)


                    # attlist_heading = document.add_heading('',level=2)
                    # attlist_heading_run = attlist_heading.add_run(f"{summ_sec_no}.1 List of documents")
                    # attlist_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # attlist_heading_run.font.name = default_font
                    # attlist_heading_run.font.size = Pt(16)
                    # attlist_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    # attlist_heading.paragraph_format.space_after = Pt(10)


                    if audited_375 == 'y':
                        # p_attach_pnl1 = document.add_paragraph()   
                        # p_attach_pnl1.add_run(
                        #     "EHT Panel Audit Reports:")
                        # p_attach_pnl1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # p_attach_pnl1.paragraph_format.space_after = Pt(5)                    

                        # df_pnl_att = form_df_pnl.ID.unique()

                        # for pnl_id in df_pnl_att:
                        #     p_attach_pnlx = document.add_paragraph(style='List Bullet')
                        #     p_attach_pnlx.add_run(f"{pnl_id} - EHT Panel - Audit Inspection Form")
                        #     p_attach_pnlx.paragraph_format.left_indent = Inches(0.5)

                        # p_attach_pnl1.add_run(
                        #     "EHT Panel Audit Reports:")
                        
                        df_pnl_att = form_df_pnl.ID.unique()

                        p_attach_pnl0 = document.add_paragraph()
                        p_attach_pnl0.add_run(a_df.loc['att_pnl_p0',lang])
                        p_attach_pnl1 = document.add_paragraph(style='List Bullet')
                        p_attach_pnl1.add_run(a_df.loc['att_pnl_p1',lang])
                        for pnl_id in df_pnl_att:
                            p_attach_pnl1.add_run(str(a_df.loc['att_pnl_p1_b1',lang].format(pnl_id)).replace('\\n', '\n'))                   
                        p_attach_pnl1.paragraph_format.left_indent = Inches(0.5)
                        p_attach_pnl2 = document.add_paragraph(style='List Bullet')
                        p_attach_pnl2.add_run(str(a_df.loc['att_pnl_p1_b2',lang]).replace('\\n', '\n'))
                        p_attach_pnl2.paragraph_format.left_indent = Inches(0.5)
                        if audit_lvl_panel == 'Standard':
                            if pnl_avail == 'y':
                                p_attach_pnl3 = document.add_paragraph(style='List Bullet')
                                p_attach_pnl3.add_run(str(a_df.loc['att_pnl_p1_b3',lang]).replace('\\n', '\n'))
                                p_attach_pnl3.paragraph_format.left_indent = Inches(0.5)
                        if (audit_lvl_panel == 'Standard' or audit_lvl_panel == 'Advanced'):
                            p_attach_pnl4 = document.add_paragraph(style='List Bullet')
                            p_attach_pnl4.add_run(str(a_df.loc['att_pnl_p1_b4',lang]).replace('\\n', '\n'))
                            p_attach_pnl4.paragraph_format.left_indent = Inches(0.5)
                        if audit_lvl_panel == 'Advanced':
                            if pnl_avail == 'n':
                                p_attach_pnl5 = document.add_paragraph(style='List Bullet')
                                p_attach_pnl5.add_run(str(a_df.loc['att_pnl_p1_b5',lang]).replace('\\n', '\n'))
                                p_attach_pnl5.paragraph_format.left_indent = Inches(0.5)
                            p_attach_pnl6 = document.add_paragraph(style='List Bullet')
                            p_attach_pnl6.add_run(str(a_df.loc['att_pnl_p1_b6',lang]).replace('\\n', '\n'))
                            p_attach_pnl6.paragraph_format.left_indent = Inches(0.5)                                                        

                    if audited_373 == 'y':

                        df_cct_att = form_df_cct.ID.unique()

                        p_attach_cct0 = document.add_paragraph()
                        p_attach_cct0.add_run(a_df.loc['att_cct_p0',lang])
                        p_attach_cct1 = document.add_paragraph(style='List Bullet')
                        p_attach_cct1.add_run(a_df.loc['att_cct_p1',lang])
                        for cct_id in df_cct_att:
                            p_attach_cct1.add_run(str(a_df.loc['att_cct_p1_b1',lang].format(cct_id)).replace('\\n', '\n'))                      
                        p_attach_cct1.paragraph_format.left_indent = Inches(0.5)
                        p_attach_cct2 = document.add_paragraph(style='List Bullet')
                        p_attach_cct2.add_run(str(a_df.loc['att_cct_p1_b2',lang]).replace('\\n', '\n'))
                        p_attach_cct2.paragraph_format.left_indent = Inches(0.5)
                        if audit_lvl_cct == 'Standard':
                            if iso_avail == 'y':
                                p_attach_cct3 = document.add_paragraph(style='List Bullet')
                                p_attach_cct3.add_run(str(a_df.loc['att_cct_p1_b3',lang]).replace('\\n', '\n'))
                                p_attach_cct3.paragraph_format.left_indent = Inches(0.5)
                            if calc_avail == 'y':
                                p_attach_cct4 = document.add_paragraph(style='List Bullet')
                                p_attach_cct4.add_run(str(a_df.loc['att_cct_p1_b4',lang]).replace('\\n', '\n'))
                                p_attach_cct4.paragraph_format.left_indent = Inches(0.5)
                            if pnl_avail == 'y':
                                p_attach_cct5 = document.add_paragraph(style='List Bullet')
                                p_attach_cct5.add_run(str(a_df.loc['att_cct_p1_b5',lang]).replace('\\n', '\n'))
                                p_attach_cct5.paragraph_format.left_indent = Inches(0.5)
                        if (audit_lvl_cct == 'Standard' or audit_lvl_cct == 'Advanced'):
                            p_attach_cct6 = document.add_paragraph(style='List Bullet')
                            p_attach_cct6.add_run(str(a_df.loc['att_cct_p1_b6',lang]).replace('\\n', '\n'))
                            p_attach_cct6.paragraph_format.left_indent = Inches(0.5)
                        if audit_lvl_cct == 'Advanced':
                            if (iso_avail == 'n' and calc_avail == 'n'):
                                p_attach_cct7 = document.add_paragraph(style='List Bullet')
                                p_attach_cct7.add_run(str(a_df.loc['att_cct_p1_b7',lang]).replace('\\n', '\n'))
                                p_attach_cct7.paragraph_format.left_indent = Inches(0.5)
                                p_attach_cct8 = document.add_paragraph(style='List Bullet')
                                p_attach_cct8.add_run(str(a_df.loc['att_cct_p1_b8',lang]).replace('\\n', '\n'))
                                p_attach_cct8.paragraph_format.left_indent = Inches(0.5)
                                p_attach_cct9 = document.add_paragraph(style='List Bullet')
                                p_attach_cct9.add_run(str(a_df.loc['att_cct_p1_b9',lang]).replace('\\n', '\n'))
                                p_attach_cct9.paragraph_format.left_indent = Inches(0.5)

                        # p_attach_cct1 = document.add_paragraph()        
                        # p_attach_cct1.add_run(
                        #     "EHT Circuit Audit Reports:")
                        # p_attach_cct1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # p_attach_cct1.paragraph_format.space_after = Pt(5)                    

                        # df_cct_att = form_df_cct.ID.unique()

                        # for cct_id in df_cct_att:
                        #     p_attach_cctx = document.add_paragraph(style='List Bullet')
                        #     p_attach_cctx.add_run(f"{cct_id} - EHT Circuit - Audit Inspection Form")
                        #     p_attach_cctx.paragraph_format.left_indent = Inches(0.5)                        

                    if audited_374 == 'y':

                        df_ins_att = form_df_ins.ID.unique()

                        p_attach_ins0 = document.add_paragraph()
                        p_attach_ins0.add_run(a_df.loc['att_ins_p0',lang])
                        p_attach_ins1 = document.add_paragraph(style='List Bullet')
                        p_attach_ins1.add_run(a_df.loc['att_ins_p1',lang])
                        for cct_id in df_ins_att:
                            p_attach_ins1.add_run(str(a_df.loc['att_ins_p1_b1',lang].format(cct_id)).replace('\\n', '\n'))               
                        p_attach_ins1.paragraph_format.left_indent = Inches(0.5)
                        p_attach_ins2 = document.add_paragraph(style='List Bullet')
                        p_attach_ins2.add_run(str(a_df.loc['att_ins_p1_b2',lang]).replace('\\n', '\n'))
                        p_attach_ins2.paragraph_format.left_indent = Inches(0.5)
                        if (audit_lvl_insul == 'Standard' or audit_lvl_insul == 'Advanced'):
                            p_attach_ins3 = document.add_paragraph(style='List Bullet')
                            p_attach_ins3.add_run(str(a_df.loc['att_ins_p1_b3',lang]).replace('\\n', '\n'))
                            p_attach_ins3.paragraph_format.left_indent = Inches(0.5)
                        # p_attach_ins1 = document.add_paragraph()        
                        # p_attach_ins1.add_run(
                        #     "EHT Insulation Audit Reports:")
                        # p_attach_ins1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # p_attach_ins1.paragraph_format.space_after = Pt(5)                    

                        # df_ins_att = form_df_ins.ID.unique()

                        # for cct_id in df_ins_att:
                        #     p_attach_insx = document.add_paragraph(style='List Bullet')
                        #     p_attach_insx.add_run(f"{cct_id} - EHT Insulation - Audit Inspection Form")
                        #     p_attach_insx.paragraph_format.left_indent = Inches(0.5)


                    att_subsec_no = 1
                    ##################################
                    ### SUB-SECTION - EHT PANEL ATTACHMENTS ###
                    ##################################

                    if audited_375 == 'y':
                        section_att_panel = document.add_section()
                        section_att_panel.orientation = WD_ORIENT.PORTRAIT
                        section_att_panel.page_width = 7772400
                        section_att_panel.page_height = 10058400
                            
                        attpanel_heading = document.add_heading('',level=2)
                        attpanel_heading_r = attpanel_heading.add_run(a_df.loc['att_subsection_pnl',lang].format(att_sec_no, att_subsec_no))
                        
                        attpanel_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        attpanel_heading_r.font.name = default_font
                        attpanel_heading_r.font.size = Pt(16)
                        attpanel_heading_r.font.color.rgb = RGBColor(0, 0, 0)
                        attpanel_heading.paragraph_format.space_after = Pt(10)

                        att_subsec_no += 1

                    if audited_373 == 'y':
                        section_att_cct = document.add_section()
                        section_att_cct.orientation = WD_ORIENT.PORTRAIT
                        section_att_cct.page_width = 7772400
                        section_att_cct.page_height = 10058400
                            
                        attcct_heading = document.add_heading('',level=2)
                        attcct_heading_r = attcct_heading.add_run(a_df.loc['att_subsection_cct',lang].format(att_sec_no, att_subsec_no))
                        attcct_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        attcct_heading_r.font.name = default_font
                        attcct_heading_r.font.size = Pt(16)
                        attcct_heading_r.font.color.rgb = RGBColor(0, 0, 0)
                        attcct_heading.paragraph_format.space_after = Pt(10)

                        att_subsec_no += 1

                    if audited_374 == 'y':
                        section_att_ins = document.add_section()
                        section_att_ins.orientation = WD_ORIENT.PORTRAIT
                        section_att_ins.page_width = 7772400
                        section_att_ins.page_height = 10058400
                            
                        attins_heading = document.add_heading('',level=2)
                        attins_heading_r = attins_heading.add_run(a_df.loc['att_subsection_ins',lang].format(att_sec_no, att_subsec_no))
                        attins_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        attins_heading_r.font.name = default_font
                        attins_heading_r.font.size = Pt(16)
                        attins_heading_r.font.color.rgb = RGBColor(0, 0, 0)
                        attins_heading.paragraph_format.space_after = Pt(10)

                        att_subsec_no += 1



                        # CHANGE FONTS IN ALL TABLES ABOVE
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs = cell.paragraphs
                                for paragraph in paragraphs:
                                    for run in paragraph.runs:
                                        font = run.font
                                        font.name= default_font
                                        font.size= Pt(9)


                    for p in document.paragraphs:
                        for run in p.runs:
                            font = run.font
                            font.name= default_font

                    #document.save('audit_report.docx')
                    io_val = BytesIO()
                    document.save(io_val)


                    ### DISPATCH TO S3 LOG ###
                    cet_timezone = pytz.timezone('Europe/Paris')
                    cet_time = datetime.now(pytz.utc).astimezone(cet_timezone)
                    
                    log_row = {'timestamp':cet_time, 'lang': lang,'reason': prescreen_selection,
                                'rep_date':repdate, 'proj_name': projname,
                                'proj_no': projpono, 'po_no': projsapno,
                                'comp_name': compname, 'client_name': clientname,
                                'comp_country': compcountry, 'comp_city': compcity,
                                'comp_street': compstreet, 'comp_zip': compzip,
                                'pid_avail': pid_avail, 'iso_avail': iso_avail, 
                                'pnl_avail': pnl_avail, 'lay_avail': lay_avail,
                                'calc_avail': calc_avail
                                }
                    
                    if audited_373 == 'y':
                        cct_row = {'cct_scope': audit_lvl_cct, 'cct_no': EHT_cct_no,
                                    'ctrl_el': ctrl_el, 'ctrl_mech': ctrl_mech,
                                    'ctrl_not': ctrl_not, 'ctrl_loc': ctrl_loc,
                                    'cct_major': viscct_majority, 'cct_fault': viscctfail_gen,
                                    'ctrl_major': ctrl_majority, 'ctrl_err_tm': ctrlfail_mtemp,
                                    'ctrl_err_lim': ctrlfail_ltemp, 'ctrl_err_pwr': ctrlfail_pw,
                                    'ctrl_err_rtd': ctrlfail_sens, 'ir_major': ir_test_majority,
                                    'ir_b10': ir_test_below10, 'ir_fault': ir_test_fault, 
                                    'cont_major': cont_major, 'cont_short': cont_short,
                                    'cont_break': cont_broke, 'volt_major': volt_major,
                                    'volt_out': volt_out, 'volt_zero': volt_zero,
                                    'curr_major': curr_major, 'curr_trip': curr_trip,
                                    'curr_zero': curr_zero
                                    }
                        
                        log_row.update(cct_row)

                    if audited_374 == 'y':
                        ins_row = {'ins_scope': audit_lvl_insul, 'ins_no': EHT_ins_no,
                                    'ins_major': visinsul_majority, 'ins_fault_s': visinsulfail_gen
                                    }
                        
                        if (insul_lvl_index == 1 or insul_lvl_index == 2):
                            ins_row.update({'ins_fault_ir': visinsulfail_ifr})

                        if insul_lvl_index == 2:
                            ins_row.update({'ins_param_major': inscheck_majority,
                                            'ins_param_err': incheck_some})

                        log_row.update(ins_row)

                    if audited_375 == 'y':
                        pnl_row = {'pnl_scope': audit_lvl_panel, 'pnl_no': EHT_pnl_no,
                         'pnl_major': vispanel_majority, 'pnl_fault_s': vispanelfail_gen,
                         'rcd_but_major': rcdbut_test_majority, 'rcd_but_fault': rcdbut_test_fail
                                }

                        if panel_lvl_index == 2:
                            pnl_row.update({'pnl_fault_if': vispanelfail_ifr})

                        if (panel_lvl_index == 1 or panel_lvl_index == 2):
                            pnl_row.update({'rcd_mes_major': rcdmes_test_majority,
                                            'rcd_mes_fault': rcdmes_test_fail})
                                    
                        log_row.update(pnl_row)

                    log_df = get_csvlog('s3-nvent-prontoforms-data/Logs/audit_report.csv')
                    log_df.loc[len(log_df)] = log_row
                    save_csv(log_df,'s3-nvent-prontoforms-data/Logs/audit_report.csv')

#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

                #xl_to_csv_full = convert_results(form_df)
                #xl_to_csv_fault = convert_results(fault_df)

                #st.sidebar.markdown("Complete Download")
        try:
            # document_ready = Document("audit_report.docx")   
            # io_val = BytesIO()
            # document_ready.save(io_val)
            user_input_full = st.sidebar.text_input("Name your file: ", max_chars = 30,value = projpono+"_Audit_report")
        except:
            pass

        try:
            io_getval = io_val.getvalue()
        except:
            pass
        else:
            #user_input_full = st.sidebar.text_input("Name your file: ", max_chars = 30,value = projpono+"_Audit_report")
            st.sidebar.download_button(
                label="Download Report",
                data=io_getval,
                file_name=user_input_full + ".docx"
                )            
